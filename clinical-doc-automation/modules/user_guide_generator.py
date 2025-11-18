"""
EDC/ePRO User Guide Generator Module

This module automatically generates comprehensive user guides for EDC/ePRO systems
based on Protocol and CRF design specifications.

Author: Clinical Documentation Automation Team
Date: 2025-11-18
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, List, Optional, Any
import os


class ScreenshotPlaceholder:
    """Class to manage screenshot placeholders in the document"""

    def __init__(self, section: str, step: str, description: str,
                 width: float = 6.0, height: float = 4.0):
        """
        Initialize a screenshot placeholder

        Args:
            section: Section name where the screenshot belongs
            step: Step number or identifier
            description: Description of what the screenshot should show
            width: Placeholder width in inches (default: 6.0)
            height: Placeholder height in inches (default: 4.0)
        """
        self.section = section
        self.step = step
        self.description = description
        self.width = width
        self.height = height
        self.inserted = False

    def __repr__(self):
        return f"Screenshot({self.section} - {self.step}): {self.description}"


class UserGuideGenerator:
    """
    Main class for generating EDC/ePRO User Guides
    """

    def __init__(self, protocol_info: Dict[str, Any], crf_design: Dict[str, Any],
                 system_name: str = "EDC/ePRO System"):
        """
        Initialize the User Guide Generator

        Args:
            protocol_info: Dictionary containing protocol information
                {
                    'protocol_id': str,
                    'protocol_title': str,
                    'sponsor': str,
                    'version': str,
                    'date': str
                }
            crf_design: Dictionary containing CRF design information
                {
                    'forms': [
                        {
                            'form_name': str,
                            'form_title': str,
                            'visit': str,
                            'fields': [
                                {
                                    'field_name': str,
                                    'field_label': str,
                                    'field_type': str,
                                    'required': bool,
                                    'validation': str
                                }
                            ]
                        }
                    ]
                }
            system_name: Name of the EDC/ePRO system
        """
        self.protocol_info = protocol_info
        self.crf_design = crf_design
        self.system_name = system_name
        self.doc = Document()
        self.screenshots = []

        # Initialize document settings
        self._setup_document_styles()

    def _setup_document_styles(self):
        """Setup custom styles for the document"""
        # Define custom styles
        styles = self.doc.styles

        # Heading 1 style
        if 'Custom Heading 1' not in styles:
            h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.size = Pt(18)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(0, 51, 102)

        # Heading 2 style
        if 'Custom Heading 2' not in styles:
            h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(0, 102, 204)

        # Step style
        if 'Step Style' not in styles:
            step_style = styles.add_style('Step Style', WD_STYLE_TYPE.PARAGRAPH)
            step_style.font.size = Pt(11)
            step_style.font.bold = True

        # Note style
        if 'Note Style' not in styles:
            note_style = styles.add_style('Note Style', WD_STYLE_TYPE.PARAGRAPH)
            note_style.font.size = Pt(10)
            note_style.font.italic = True
            note_style.font.color.rgb = RGBColor(128, 128, 128)

    def _add_heading(self, text: str, level: int = 1):
        """Add a heading to the document"""
        heading = self.doc.add_heading(text, level=level)
        return heading

    def _add_paragraph(self, text: str, style: Optional[str] = None):
        """Add a paragraph to the document"""
        para = self.doc.add_paragraph(text)
        if style:
            para.style = style
        return para

    def _add_screenshot_placeholder(self, section: str, step: str,
                                   description: str, width: float = 6.0):
        """
        Add a screenshot placeholder to the document

        Args:
            section: Section name
            step: Step identifier
            description: Screenshot description
            width: Width in inches
        """
        placeholder = ScreenshotPlaceholder(section, step, description, width)
        self.screenshots.append(placeholder)

        # Add placeholder box
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Light Grid Accent 1'
        cell = table.rows[0].cells[0]
        cell_para = cell.paragraphs[0]

        # Add placeholder text
        run = cell_para.add_run(f"[SCREENSHOT PLACEHOLDER]\n")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)

        run2 = cell_para.add_run(f"Section: {section}\n")
        run2.font.size = Pt(9)

        run3 = cell_para.add_run(f"Step: {step}\n")
        run3.font.size = Pt(9)

        run4 = cell_para.add_run(f"Description: {description}")
        run4.font.size = Pt(9)
        run4.italic = True

        self.doc.add_paragraph()  # Add spacing

        return placeholder

    def generate_cover_page(self):
        """Generate the cover page"""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{self.system_name}\nUser Guide")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 51, 102)

        self.doc.add_paragraph()

        # Protocol information
        protocol_para = self.doc.add_paragraph()
        protocol_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        protocol_run = protocol_para.add_run(
            f"Protocol: {self.protocol_info.get('protocol_id', 'N/A')}\n"
            f"{self.protocol_info.get('protocol_title', 'N/A')}"
        )
        protocol_run.font.size = Pt(14)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Version and date
        version_para = self.doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_para.add_run(
            f"Version: {self.protocol_info.get('version', '1.0')}\n"
            f"Date: {self.protocol_info.get('date', datetime.now().strftime('%Y-%m-%d'))}\n"
            f"Sponsor: {self.protocol_info.get('sponsor', 'N/A')}"
        )

        self.doc.add_page_break()

    def generate_introduction(self):
        """Generate the Introduction section"""
        self._add_heading("1. Introduction", level=1)

        self._add_paragraph(
            f"This user guide provides comprehensive instructions for using the "
            f"{self.system_name} for the {self.protocol_info.get('protocol_id', 'N/A')} study. "
            f"The guide covers all aspects of data entry, navigation, and system usage."
        )

        self.doc.add_paragraph()

        # Purpose
        self._add_heading("1.1 Purpose", level=2)
        self._add_paragraph(
            "The purpose of this guide is to:"
        )

        purposes = [
            "Provide step-by-step instructions for accessing and using the system",
            "Explain how to enter and manage clinical data",
            "Guide users through the query management process",
            "Demonstrate report generation and data review procedures",
            "Ensure data quality and compliance with study protocols"
        ]

        for purpose in purposes:
            self.doc.add_paragraph(purpose, style='List Bullet')

        self.doc.add_paragraph()

        # Intended audience
        self._add_heading("1.2 Intended Audience", level=2)
        self._add_paragraph(
            "This guide is intended for:"
        )

        audiences = [
            "Clinical Research Coordinators (CRCs)",
            "Site Personnel",
            "Data Managers",
            "Clinical Research Associates (CRAs)",
            "Principal Investigators"
        ]

        for audience in audiences:
            self.doc.add_paragraph(audience, style='List Bullet')

        self.doc.add_page_break()

    def generate_system_access(self):
        """Generate the System Access section"""
        self._add_heading("2. System Access", level=1)

        # Login procedure
        self._add_heading("2.1 Logging In", level=2)

        steps = [
            ("Navigate to the system URL",
             "Open your web browser and navigate to the system URL provided by your administrator."),
            ("Enter credentials",
             "Enter your username and password in the login fields."),
            ("Two-factor authentication (if applicable)",
             "If two-factor authentication is enabled, enter the verification code sent to your registered device."),
            ("Click Login",
             "Click the 'Login' button to access the system.")
        ]

        for i, (step_title, step_desc) in enumerate(steps, 1):
            para = self._add_paragraph(f"Step {i}: {step_title}", style='Step Style')
            self._add_paragraph(step_desc)
            self._add_screenshot_placeholder(
                "System Access",
                f"Step {i}",
                f"Screenshot showing {step_title.lower()}"
            )

        # Password requirements
        self._add_heading("2.2 Password Requirements", level=2)
        self._add_paragraph(
            "Your password must meet the following requirements:"
        )

        requirements = [
            "Minimum 8 characters in length",
            "At least one uppercase letter",
            "At least one lowercase letter",
            "At least one number",
            "At least one special character (!@#$%^&*)",
            "Cannot be the same as your username",
            "Must be changed every 90 days"
        ]

        for req in requirements:
            self.doc.add_paragraph(req, style='List Bullet')

        # Troubleshooting
        self._add_heading("2.3 Troubleshooting Login Issues", level=2)

        troubleshooting = [
            ("Forgot Password",
             "Click the 'Forgot Password' link on the login page and follow the instructions to reset your password."),
            ("Account Locked",
             "After 3 failed login attempts, your account will be locked. Contact your system administrator to unlock it."),
            ("Browser Compatibility",
             "Use the latest version of Chrome, Firefox, Safari, or Edge for optimal performance.")
        ]

        for issue, solution in troubleshooting:
            para = self._add_paragraph(f"{issue}:", style='Step Style')
            self._add_paragraph(solution)

        self.doc.add_page_break()

    def generate_navigation(self):
        """Generate the Navigation section"""
        self._add_heading("3. System Navigation", level=1)

        # Main interface
        self._add_heading("3.1 Main Interface Overview", level=2)
        self._add_paragraph(
            "After logging in, you will see the main interface with the following components:"
        )

        components = [
            ("Navigation Menu", "Located on the left side, provides access to all system functions"),
            ("Dashboard", "Displays study overview, pending tasks, and recent activities"),
            ("Patient List", "Shows all enrolled patients and their visit status"),
            ("Search Bar", "Allows quick search for patients, forms, or data"),
            ("User Profile", "Access to user settings and logout option")
        ]

        for component, description in components:
            para = self._add_paragraph(f"{component}:", style='Step Style')
            self._add_paragraph(description)

        self._add_screenshot_placeholder(
            "Navigation",
            "Main Interface",
            "Screenshot showing the main interface with all components labeled"
        )

        # Menu structure
        self._add_heading("3.2 Menu Structure", level=2)

        menu_items = [
            "Home - Return to dashboard",
            "Patients - Manage patient records",
            "Data Entry - Access CRFs for data entry",
            "Queries - View and respond to data queries",
            "Reports - Generate and view reports",
            "Administration - System settings (admin users only)"
        ]

        for item in menu_items:
            self.doc.add_paragraph(item, style='List Bullet')

        # Patient selection
        self._add_heading("3.3 Selecting a Patient", level=2)

        patient_steps = [
            ("Access Patient List", "Click on 'Patients' in the navigation menu"),
            ("Search or Browse", "Use the search bar or browse the patient list"),
            ("Select Patient", "Click on the patient ID or name to view patient details"),
            ("View Visit Schedule", "Review the patient's visit schedule and form status")
        ]

        for i, (step_title, step_desc) in enumerate(patient_steps, 1):
            para = self._add_paragraph(f"Step {i}: {step_title}", style='Step Style')
            self._add_paragraph(step_desc)

        self._add_screenshot_placeholder(
            "Navigation",
            "Patient Selection",
            "Screenshot showing patient list and selection process"
        )

        self.doc.add_page_break()

    def generate_data_entry_instructions(self):
        """Generate the Data Entry Instructions section for each CRF"""
        self._add_heading("4. Data Entry Instructions", level=1)

        self._add_paragraph(
            "This section provides detailed instructions for entering data into each CRF form. "
            "Please follow these instructions carefully to ensure data quality and compliance."
        )

        self.doc.add_paragraph()

        # General guidelines
        self._add_heading("4.1 General Data Entry Guidelines", level=2)

        guidelines = [
            "Always select the correct patient and visit before entering data",
            "Complete all required fields (marked with *)",
            "Follow field validation rules and format requirements",
            "Save your work frequently to prevent data loss",
            "Review all entered data before marking the form as complete",
            "Do not use abbreviations unless specified in the protocol",
            "Enter dates in the format specified (typically YYYY-MM-DD)",
            "For numeric fields, enter values within the specified range"
        ]

        for guideline in guidelines:
            self.doc.add_paragraph(guideline, style='List Bullet')

        self.doc.add_paragraph()

        # Field types reference
        self._add_heading("4.2 Field Types Reference", level=2)

        field_types = [
            ("Text Field", "Enter free text. Maximum length may be specified."),
            ("Number Field", "Enter numeric values only. May have min/max validation."),
            ("Date Field", "Enter date in specified format. Use date picker if available."),
            ("Dropdown", "Select one option from the dropdown list."),
            ("Radio Button", "Select one option from the available choices."),
            ("Checkbox", "Select all applicable options."),
            ("Text Area", "Enter longer text responses or comments.")
        ]

        for field_type, description in field_types:
            para = self._add_paragraph(f"{field_type}:", style='Step Style')
            self._add_paragraph(description)

        self.doc.add_paragraph()

        # Generate instructions for each CRF form
        if 'forms' in self.crf_design:
            for idx, form in enumerate(self.crf_design['forms'], 1):
                self._generate_form_instructions(form, idx)

        self.doc.add_page_break()

    def _generate_form_instructions(self, form: Dict[str, Any], form_number: int):
        """Generate instructions for a specific CRF form"""
        form_name = form.get('form_name', f'Form {form_number}')
        form_title = form.get('form_title', form_name)
        visit = form.get('visit', 'N/A')

        self._add_heading(f"4.{form_number + 2} {form_title}", level=2)

        self._add_paragraph(f"Form Name: {form_name}")
        self._add_paragraph(f"Visit: {visit}")

        self.doc.add_paragraph()

        # Accessing the form
        para = self._add_paragraph("Accessing the Form:", style='Step Style')
        access_steps = [
            f"Navigate to Data Entry > {visit}",
            f"Select the patient",
            f"Click on '{form_title}' in the form list"
        ]

        for step in access_steps:
            self.doc.add_paragraph(step, style='List Bullet')

        self._add_screenshot_placeholder(
            f"Data Entry - {form_title}",
            "Form Access",
            f"Screenshot showing how to access {form_title}"
        )

        # Field-by-field instructions
        if 'fields' in form and form['fields']:
            self._add_paragraph("Field Instructions:", style='Step Style')

            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Field Name'
            header_cells[1].text = 'Type'
            header_cells[2].text = 'Required'
            header_cells[3].text = 'Instructions'

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add field rows
            for field in form['fields']:
                row_cells = table.add_row().cells
                row_cells[0].text = field.get('field_label', field.get('field_name', ''))
                row_cells[1].text = field.get('field_type', 'Text')
                row_cells[2].text = 'Yes' if field.get('required', False) else 'No'

                # Generate instruction text
                instruction = self._generate_field_instruction(field)
                row_cells[3].text = instruction

            self.doc.add_paragraph()

        # Saving the form
        para = self._add_paragraph("Saving and Completing the Form:", style='Step Style')
        save_steps = [
            "Click 'Save' to save your progress without submitting",
            "Review all entered data for accuracy",
            "Click 'Save and Complete' when all required fields are filled",
            "Confirm that the form status changes to 'Complete'"
        ]

        for step in save_steps:
            self.doc.add_paragraph(step, style='List Bullet')

        self._add_screenshot_placeholder(
            f"Data Entry - {form_title}",
            "Completed Form",
            f"Screenshot showing completed {form_title} with all fields filled"
        )

        self.doc.add_paragraph()

    def _generate_field_instruction(self, field: Dict[str, Any]) -> str:
        """Generate instruction text for a specific field"""
        field_type = field.get('field_type', 'text').lower()
        validation = field.get('validation', '')

        instructions = []

        if field_type in ['text', 'string']:
            instructions.append("Enter text value")
        elif field_type in ['number', 'integer', 'decimal']:
            instructions.append("Enter numeric value")
        elif field_type == 'date':
            instructions.append("Enter date in YYYY-MM-DD format")
        elif field_type in ['dropdown', 'select']:
            instructions.append("Select one option from dropdown")
        elif field_type == 'radio':
            instructions.append("Select one option")
        elif field_type == 'checkbox':
            instructions.append("Select all applicable options")
        elif field_type == 'textarea':
            instructions.append("Enter detailed text response")

        if validation:
            instructions.append(f"Validation: {validation}")

        return ". ".join(instructions) + "."

    def generate_query_management(self):
        """Generate the Query Management section"""
        self._add_heading("5. Query Management", level=1)

        self._add_paragraph(
            "Queries are generated when there are questions or concerns about entered data. "
            "This section explains how to view, respond to, and resolve queries."
        )

        self.doc.add_paragraph()

        # Viewing queries
        self._add_heading("5.1 Viewing Queries", level=2)

        view_steps = [
            ("Access Query Module", "Click on 'Queries' in the navigation menu"),
            ("View Query List", "Review the list of open queries for your site"),
            ("Filter Queries", "Use filters to view queries by status, patient, or form"),
            ("Select Query", "Click on a query to view details")
        ]

        for i, (step_title, step_desc) in enumerate(view_steps, 1):
            para = self._add_paragraph(f"Step {i}: {step_title}", style='Step Style')
            self._add_paragraph(step_desc)

        self._add_screenshot_placeholder(
            "Query Management",
            "Query List",
            "Screenshot showing the query list with filters and status indicators"
        )

        # Query status
        self._add_heading("5.2 Query Status Types", level=2)

        statuses = [
            ("Open", "New query requiring response"),
            ("Answered", "Site has responded, awaiting review"),
            ("Closed", "Query has been resolved"),
            ("Cancelled", "Query has been cancelled by data manager")
        ]

        for status, description in statuses:
            para = self._add_paragraph(f"{status}:", style='Step Style')
            self._add_paragraph(description)

        # Responding to queries
        self._add_heading("5.3 Responding to Queries", level=2)

        response_steps = [
            ("Open Query", "Click on the query to view details and the question"),
            ("Review Data", "Review the queried data field and protocol requirements"),
            ("Verify or Correct", "Either verify the data is correct or make corrections"),
            ("Enter Response", "Enter your response in the query response field"),
            ("Attach Documentation", "If needed, attach supporting documentation"),
            ("Submit Response", "Click 'Submit Response' to send your answer"),
            ("Confirm Submission", "Verify the query status changes to 'Answered'")
        ]

        for i, (step_title, step_desc) in enumerate(response_steps, 1):
            para = self._add_paragraph(f"Step {i}: {step_title}", style='Step Style')
            self._add_paragraph(step_desc)

        self._add_screenshot_placeholder(
            "Query Management",
            "Query Response",
            "Screenshot showing the query response interface with all fields"
        )

        # Best practices
        self._add_heading("5.4 Query Response Best Practices", level=2)

        best_practices = [
            "Respond to queries promptly (within 48 hours if possible)",
            "Provide clear and complete explanations",
            "Reference source documents when applicable",
            "Correct any data errors before responding",
            "Do not close queries yourself - wait for data manager to close",
            "Document any protocol deviations or exceptions",
            "Maintain professional communication in all responses"
        ]

        for practice in best_practices:
            self.doc.add_paragraph(practice, style='List Bullet')

        self.doc.add_page_break()

    def generate_report_generation(self):
        """Generate the Report Generation section"""
        self._add_heading("6. Report Generation", level=1)

        self._add_paragraph(
            "The system provides various reports for data review, monitoring, and study management. "
            "This section explains how to generate and use different types of reports."
        )

        self.doc.add_paragraph()

        # Available reports
        self._add_heading("6.1 Available Reports", level=2)

        reports = [
            ("Patient Enrollment Report", "Shows enrollment status and demographics"),
            ("Visit Completion Report", "Displays visit completion status by patient"),
            ("Data Completion Report", "Shows CRF completion rates"),
            ("Query Report", "Lists all open and closed queries"),
            ("Audit Trail Report", "Shows all data changes and user activities"),
            ("Missing Data Report", "Identifies incomplete forms and missing data"),
            ("Protocol Deviation Report", "Lists any protocol deviations or violations")
        ]

        for report_name, description in reports:
            para = self._add_paragraph(f"{report_name}:", style='Step Style')
            self._add_paragraph(description)

        # Generating a report
        self._add_heading("6.2 Generating a Report", level=2)

        generation_steps = [
            ("Access Reports", "Click on 'Reports' in the navigation menu"),
            ("Select Report Type", "Choose the report type from the available options"),
            ("Set Parameters", "Select date range, patients, visits, or other filters"),
            ("Choose Format", "Select output format (PDF, Excel, CSV)"),
            ("Generate Report", "Click 'Generate Report' button"),
            ("Download or View", "Download the report or view it in the browser")
        ]

        for i, (step_title, step_desc) in enumerate(generation_steps, 1):
            para = self._add_paragraph(f"Step {i}: {step_title}", style='Step Style')
            self._add_paragraph(step_desc)

        self._add_screenshot_placeholder(
            "Report Generation",
            "Report Parameters",
            "Screenshot showing report generation interface with parameter selection"
        )

        # Scheduled reports
        self._add_heading("6.3 Scheduling Reports", level=2)

        self._add_paragraph(
            "You can schedule reports to be generated automatically at regular intervals:"
        )

        schedule_steps = [
            ("Select Report Type", "Choose the report you want to schedule"),
            ("Click Schedule", "Click the 'Schedule Report' button"),
            ("Set Frequency", "Choose daily, weekly, or monthly frequency"),
            ("Set Recipients", "Enter email addresses for report distribution"),
            ("Confirm Schedule", "Review and confirm the schedule settings")
        ]

        for i, (step_title, step_desc) in enumerate(schedule_steps, 1):
            para = self._add_paragraph(f"Step {i}: {step_title}", style='Step Style')
            self._add_paragraph(step_desc)

        # Report interpretation
        self._add_heading("6.4 Report Interpretation", level=2)

        interpretation_tips = [
            "Review report headers for generation date and parameters used",
            "Check for any warnings or alerts highlighted in red",
            "Compare current data with previous reports to track trends",
            "Use filters to drill down into specific issues or areas",
            "Export reports for sharing with study team members",
            "Archive important reports for regulatory compliance"
        ]

        for tip in interpretation_tips:
            self.doc.add_paragraph(tip, style='List Bullet')

        self._add_screenshot_placeholder(
            "Report Generation",
            "Sample Report",
            "Screenshot showing a sample generated report with key features highlighted"
        )

        self.doc.add_page_break()

    def generate_appendix(self):
        """Generate appendix with additional resources"""
        self._add_heading("7. Appendix", level=1)

        # Contact information
        self._add_heading("7.1 Support Contact Information", level=2)

        self._add_paragraph("For technical support or questions, contact:")

        contacts = [
            ("Help Desk", "helpdesk@example.com", "+1-800-XXX-XXXX"),
            ("Data Management", "datamanagement@example.com", "+1-800-XXX-XXXX"),
            ("System Administrator", "sysadmin@example.com", "+1-800-XXX-XXXX")
        ]

        for role, email, phone in contacts:
            para = self._add_paragraph(f"{role}:", style='Step Style')
            self._add_paragraph(f"Email: {email}")
            self._add_paragraph(f"Phone: {phone}")
            self.doc.add_paragraph()

        # Common abbreviations
        self._add_heading("7.2 Common Abbreviations", level=2)

        abbreviations = [
            ("AE", "Adverse Event"),
            ("CRF", "Case Report Form"),
            ("CRA", "Clinical Research Associate"),
            ("CRC", "Clinical Research Coordinator"),
            ("EDC", "Electronic Data Capture"),
            ("ePRO", "Electronic Patient Reported Outcome"),
            ("ICF", "Informed Consent Form"),
            ("PI", "Principal Investigator"),
            ("SAE", "Serious Adverse Event")
        ]

        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Abbreviation'
        header_cells[1].text = 'Definition'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for abbr, definition in abbreviations:
            row_cells = table.add_row().cells
            row_cells[0].text = abbr
            row_cells[1].text = definition

        self.doc.add_paragraph()

        # Revision history
        self._add_heading("7.3 Document Revision History", level=2)

        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Version'
        header_cells[1].text = 'Date'
        header_cells[2].text = 'Author'
        header_cells[3].text = 'Changes'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add initial version
        row_cells = table.add_row().cells
        row_cells[0].text = self.protocol_info.get('version', '1.0')
        row_cells[1].text = self.protocol_info.get('date', datetime.now().strftime('%Y-%m-%d'))
        row_cells[2].text = 'Clinical Documentation Automation'
        row_cells[3].text = 'Initial release'

    def get_screenshot_list(self) -> List[ScreenshotPlaceholder]:
        """
        Get list of all screenshot placeholders

        Returns:
            List of ScreenshotPlaceholder objects
        """
        return self.screenshots

    def export_screenshot_list(self, output_path: str):
        """
        Export screenshot list to a text file for reference

        Args:
            output_path: Path to save the screenshot list
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("SCREENSHOT REQUIREMENTS LIST\n")
            f.write("=" * 80 + "\n\n")
            f.write(f"Total Screenshots Needed: {len(self.screenshots)}\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

            for i, screenshot in enumerate(self.screenshots, 1):
                f.write(f"{i}. Section: {screenshot.section}\n")
                f.write(f"   Step: {screenshot.step}\n")
                f.write(f"   Description: {screenshot.description}\n")
                f.write(f"   Dimensions: {screenshot.width}\" x {screenshot.height}\"\n")
                f.write(f"   Status: {'Inserted' if screenshot.inserted else 'Pending'}\n")
                f.write("-" * 80 + "\n")

    def generate(self, output_path: str, include_appendix: bool = True) -> str:
        """
        Generate the complete user guide

        Args:
            output_path: Path to save the generated document
            include_appendix: Whether to include the appendix section

        Returns:
            Path to the generated document
        """
        print("Generating User Guide...")

        # Generate all sections
        print("  - Cover Page")
        self.generate_cover_page()

        print("  - Introduction")
        self.generate_introduction()

        print("  - System Access")
        self.generate_system_access()

        print("  - Navigation")
        self.generate_navigation()

        print("  - Data Entry Instructions")
        self.generate_data_entry_instructions()

        print("  - Query Management")
        self.generate_query_management()

        print("  - Report Generation")
        self.generate_report_generation()

        if include_appendix:
            print("  - Appendix")
            self.generate_appendix()

        # Save document
        print(f"\nSaving document to {output_path}")
        self.doc.save(output_path)

        # Export screenshot list
        screenshot_list_path = output_path.replace('.docx', '_screenshots.txt')
        print(f"Saving screenshot list to {screenshot_list_path}")
        self.export_screenshot_list(screenshot_list_path)

        print(f"\n✓ User Guide generated successfully!")
        print(f"✓ Total screenshots needed: {len(self.screenshots)}")

        return output_path


# Utility functions

def create_sample_protocol_info() -> Dict[str, Any]:
    """Create sample protocol information for testing"""
    return {
        'protocol_id': 'PROTO-2025-001',
        'protocol_title': 'A Phase III, Randomized, Double-Blind Study of Treatment X in Patients with Condition Y',
        'sponsor': 'Example Pharmaceuticals Inc.',
        'version': '1.0',
        'date': datetime.now().strftime('%Y-%m-%d')
    }


def create_sample_crf_design() -> Dict[str, Any]:
    """Create sample CRF design for testing"""
    return {
        'forms': [
            {
                'form_name': 'demographics',
                'form_title': 'Demographics',
                'visit': 'Screening',
                'fields': [
                    {
                        'field_name': 'subject_initials',
                        'field_label': 'Subject Initials',
                        'field_type': 'text',
                        'required': True,
                        'validation': '3 letters only'
                    },
                    {
                        'field_name': 'date_of_birth',
                        'field_label': 'Date of Birth',
                        'field_type': 'date',
                        'required': True,
                        'validation': 'Must be at least 18 years ago'
                    },
                    {
                        'field_name': 'gender',
                        'field_label': 'Gender',
                        'field_type': 'radio',
                        'required': True,
                        'validation': 'Male/Female/Other'
                    },
                    {
                        'field_name': 'race',
                        'field_label': 'Race',
                        'field_type': 'dropdown',
                        'required': True,
                        'validation': 'Select from predefined list'
                    }
                ]
            },
            {
                'form_name': 'vital_signs',
                'form_title': 'Vital Signs',
                'visit': 'All Visits',
                'fields': [
                    {
                        'field_name': 'systolic_bp',
                        'field_label': 'Systolic Blood Pressure (mmHg)',
                        'field_type': 'number',
                        'required': True,
                        'validation': 'Range: 70-200'
                    },
                    {
                        'field_name': 'diastolic_bp',
                        'field_label': 'Diastolic Blood Pressure (mmHg)',
                        'field_type': 'number',
                        'required': True,
                        'validation': 'Range: 40-130'
                    },
                    {
                        'field_name': 'heart_rate',
                        'field_label': 'Heart Rate (bpm)',
                        'field_type': 'number',
                        'required': True,
                        'validation': 'Range: 40-200'
                    },
                    {
                        'field_name': 'temperature',
                        'field_label': 'Temperature (°C)',
                        'field_type': 'decimal',
                        'required': True,
                        'validation': 'Range: 35.0-42.0'
                    },
                    {
                        'field_name': 'comments',
                        'field_label': 'Comments',
                        'field_type': 'textarea',
                        'required': False,
                        'validation': 'Max 500 characters'
                    }
                ]
            },
            {
                'form_name': 'adverse_events',
                'form_title': 'Adverse Events',
                'visit': 'All Visits',
                'fields': [
                    {
                        'field_name': 'ae_occurred',
                        'field_label': 'Did any adverse event occur?',
                        'field_type': 'radio',
                        'required': True,
                        'validation': 'Yes/No'
                    },
                    {
                        'field_name': 'ae_description',
                        'field_label': 'Event Description',
                        'field_type': 'textarea',
                        'required': False,
                        'validation': 'Required if AE occurred'
                    },
                    {
                        'field_name': 'ae_severity',
                        'field_label': 'Severity',
                        'field_type': 'dropdown',
                        'required': False,
                        'validation': 'Mild/Moderate/Severe'
                    },
                    {
                        'field_name': 'ae_related',
                        'field_label': 'Related to Study Drug?',
                        'field_type': 'radio',
                        'required': False,
                        'validation': 'Yes/No/Possibly'
                    }
                ]
            }
        ]
    }


if __name__ == "__main__":
    # Example usage
    print("EDC/ePRO User Guide Generator - Example")
    print("=" * 60)

    # Create sample data
    protocol_info = create_sample_protocol_info()
    crf_design = create_sample_crf_design()

    # Initialize generator
    generator = UserGuideGenerator(
        protocol_info=protocol_info,
        crf_design=crf_design,
        system_name="MyClinicalEDC System"
    )

    # Generate user guide
    output_path = "user_guide_example.docx"
    generator.generate(output_path)

    print("\nScreenshot placeholders:")
    for i, screenshot in enumerate(generator.get_screenshot_list(), 1):
        print(f"{i}. {screenshot}")
