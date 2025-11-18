"""
Bestat樣式分析器測試套件

測試所有核心功能：
- 樣式提取
- 樣式套用
- 規範驗證
- 樣式比較
- 配置儲存與載入

作者：Clinical Document Automation Team
版本：1.0.0
"""

import unittest
import os
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, RGBColor, Pt

from bestat_style_analyzer import (
    BestatStyleAnalyzer,
    analyze_document_style,
    apply_bestat_style_to_document,
    validate_document_compliance
)


class TestBestatStyleAnalyzer(unittest.TestCase):
    """Bestat樣式分析器測試類別"""

    @classmethod
    def setUpClass(cls):
        """設定測試環境"""
        cls.test_dir = Path(__file__).parent.parent / "output" / "test_bestat"
        cls.test_dir.mkdir(parents=True, exist_ok=True)

        # 建立測試文件
        cls.test_doc_path = cls.test_dir / "test_document.docx"
        cls._create_test_document(cls.test_doc_path)

        # 建立第二個測試文件用於比較
        cls.test_doc2_path = cls.test_dir / "test_document2.docx"
        cls._create_test_document2(cls.test_doc2_path)

        print(f"\n測試環境已建立: {cls.test_dir}")

    @staticmethod
    def _create_test_document(path: Path):
        """建立測試用Word文件"""
        doc = Document()

        # 設定頁面
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # 添加標題
        title = doc.add_paragraph("測試文件標題")
        title.runs[0].font.name = "Arial"
        title.runs[0].font.size = Pt(18)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(0, 51, 153)

        # 添加一級標題
        h1 = doc.add_paragraph("第一章 簡介")
        h1.runs[0].font.name = "Arial"
        h1.runs[0].font.size = Pt(16)
        h1.runs[0].font.bold = True

        # 添加內文
        body = doc.add_paragraph("這是測試內文段落。")
        body.runs[0].font.name = "Calibri"
        body.runs[0].font.size = Pt(11)

        # 添加表格
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Light Grid Accent 1'

        # 填入表格內容
        header_cells = table.rows[0].cells
        header_cells[0].text = "項目"
        header_cells[1].text = "數值"
        header_cells[2].text = "備註"

        data_cells = table.rows[1].cells
        data_cells[0].text = "測試項目"
        data_cells[1].text = "100"
        data_cells[2].text = "正常"

        # 儲存文件
        doc.save(str(path))

    @staticmethod
    def _create_test_document2(path: Path):
        """建立第二個測試文件（樣式稍有不同）"""
        doc = Document()

        # 設定頁面（稍微不同的邊距）
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.2)  # 不同
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # 添加標題（不同顏色）
        title = doc.add_paragraph("測試文件標題2")
        title.runs[0].font.name = "Arial"
        title.runs[0].font.size = Pt(18)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # 不同顏色

        # 添加內文
        body = doc.add_paragraph("這是第二個測試文件。")
        body.runs[0].font.name = "Calibri"
        body.runs[0].font.size = Pt(11)

        # 儲存文件
        doc.save(str(path))

    def setUp(self):
        """每個測試前的設定"""
        self.analyzer = BestatStyleAnalyzer()

    def test_01_initialization(self):
        """測試初始化"""
        print("\n[測試] 初始化...")

        self.assertIsNotNone(self.analyzer)
        self.assertIsNotNone(self.analyzer.style_config)
        self.assertIn('company_info', self.analyzer.style_config)
        self.assertIn('page_setup', self.analyzer.style_config)
        self.assertIn('fonts', self.analyzer.style_config)

        print("  ✓ 初始化成功")
        print(f"  ✓ 公司名稱: {self.analyzer.style_config['company_info']['name']}")

    def test_02_default_config(self):
        """測試預設配置"""
        print("\n[測試] 預設配置...")

        config = self.analyzer.style_config

        # 檢查公司資訊
        self.assertEqual(config['company_info']['name'], 'Bestat Inc.')

        # 檢查頁面設定
        self.assertEqual(config['page_setup']['size'], 'A4')
        self.assertEqual(config['page_setup']['orientation'], 'portrait')

        # 檢查字體設定
        self.assertIn('title', config['fonts'])
        self.assertIn('heading1', config['fonts'])
        self.assertIn('body', config['fonts'])

        # 檢查顏色方案
        self.assertIn('primary', config['colors'])
        primary = config['colors']['primary']
        self.assertEqual(primary['r'], 0)
        self.assertEqual(primary['g'], 51)
        self.assertEqual(primary['b'], 153)

        print("  ✓ 預設配置驗證通過")

    def test_03_extract_styles(self):
        """測試樣式提取"""
        print("\n[測試] 樣式提取...")

        styles = self.analyzer.extract_styles_from_document(str(self.test_doc_path))

        # 檢查提取的樣式結構
        self.assertIn('page_setup', styles)
        self.assertIn('fonts', styles)
        self.assertIn('colors', styles)
        self.assertIn('header', styles)
        self.assertIn('footer', styles)
        self.assertIn('table_style', styles)
        self.assertIn('metadata', styles)

        # 檢查頁面設定
        page_setup = styles['page_setup']
        self.assertEqual(page_setup['size'], 'A4')
        self.assertEqual(page_setup['width_inches'], 8.27)

        # 檢查元數據
        metadata = styles['metadata']
        self.assertIn('extracted_from', metadata)
        self.assertIn('total_paragraphs', metadata)
        self.assertGreater(metadata['total_paragraphs'], 0)

        print("  ✓ 樣式提取成功")
        print(f"  ✓ 提取了 {len(styles)} 個主要類別")
        print(f"  ✓ 總段落數: {metadata['total_paragraphs']}")

    def test_04_save_load_config(self):
        """測試配置儲存與載入"""
        print("\n[測試] 配置儲存與載入...")

        # 儲存配置
        config_path = self.test_dir / "test_config.json"
        self.analyzer.save_style_config(str(config_path))

        self.assertTrue(config_path.exists())
        print(f"  ✓ 配置已儲存至: {config_path}")

        # 載入配置
        new_analyzer = BestatStyleAnalyzer()
        loaded_config = new_analyzer.load_style_config(str(config_path))

        self.assertIsNotNone(loaded_config)
        self.assertIn('company_info', loaded_config)

        print("  ✓ 配置載入成功")

        # 驗證JSON格式
        with open(config_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
            self.assertIsInstance(json_data, dict)
            print("  ✓ JSON格式驗證通過")

    def test_05_apply_bestat_style(self):
        """測試套用Bestat樣式"""
        print("\n[測試] 套用Bestat樣式...")

        # 建立新文件
        doc = Document()
        doc.add_paragraph("測試內容")

        # 套用Bestat樣式
        styled_doc = self.analyzer.apply_bestat_style(
            doc,
            document_title="測試文件",
            protocol_number="PRO-2025-TEST",
            version="1.0"
        )

        self.assertIsNotNone(styled_doc)

        # 檢查頁面設定
        section = styled_doc.sections[0]
        self.assertAlmostEqual(section.page_width.inches, 8.27, places=1)
        self.assertAlmostEqual(section.top_margin.inches, 1.0, places=1)

        # 儲存測試文件
        output_path = self.test_dir / "bestat_styled.docx"
        styled_doc.save(str(output_path))

        print(f"  ✓ Bestat樣式已套用")
        print(f"  ✓ 測試文件已儲存: {output_path}")

    def test_06_validate_compliance(self):
        """測試規範驗證"""
        print("\n[測試] 規範驗證...")

        # 建立符合規範的文件
        doc = Document()
        styled_doc = self.analyzer.apply_bestat_style(
            doc,
            document_title="驗證測試",
            protocol_number="PRO-2025-VAL",
            version="1.0"
        )

        # 驗證
        result = self.analyzer.validate_bestat_compliance(styled_doc)

        self.assertIn('compliant', result)
        self.assertIn('issues', result)
        self.assertIn('warnings', result)
        self.assertIsInstance(result['compliant'], bool)

        print(f"  ✓ 驗證完成")
        print(f"  ✓ 符合規範: {result['compliant']}")
        print(f"  ✓ 問題數: {result['total_issues']}")
        print(f"  ✓ 警告數: {result['total_warnings']}")

        if result['issues']:
            print("\n  問題清單:")
            for issue in result['issues']:
                print(f"    - {issue}")

        if result['warnings']:
            print("\n  警告清單:")
            for warning in result['warnings']:
                print(f"    - {warning}")

    def test_07_compare_styles(self):
        """測試樣式比較"""
        print("\n[測試] 樣式比較...")

        # 比較兩個文件
        comparison = self.analyzer.compare_styles(
            str(self.test_doc_path),
            str(self.test_doc2_path)
        )

        self.assertIn('differences', comparison)
        self.assertIn('total_differences', comparison)

        print(f"  ✓ 比較完成")
        print(f"  ✓ 總差異數: {comparison['total_differences']}")

        if comparison['differences']['page_setup']:
            print("\n  頁面設定差異:")
            for diff in comparison['differences']['page_setup']:
                print(f"    - {diff}")

    def test_08_generate_style_report(self):
        """測試生成樣式報告"""
        print("\n[測試] 生成樣式報告...")

        report_path = self.test_dir / "style_report.json"
        self.analyzer.generate_style_report(
            str(self.test_doc_path),
            str(report_path)
        )

        self.assertTrue(report_path.exists())

        # 讀取並驗證報告
        with open(report_path, 'r', encoding='utf-8') as f:
            report = json.load(f)

        self.assertIn('document', report)
        self.assertIn('extracted_styles', report)
        self.assertIn('validation_results', report)
        self.assertIn('recommendations', report)

        print(f"  ✓ 報告已生成: {report_path}")
        print(f"  ✓ 建議數: {len(report['recommendations'])}")

    def test_09_convenience_functions(self):
        """測試便利函數"""
        print("\n[測試] 便利函數...")

        # 測試 analyze_document_style
        output_json = self.test_dir / "quick_analysis.json"
        styles = analyze_document_style(
            str(self.test_doc_path),
            str(output_json)
        )

        self.assertIsNotNone(styles)
        self.assertTrue(output_json.exists())
        print("  ✓ analyze_document_style() 測試通過")

        # 測試 validate_document_compliance
        validation = validate_document_compliance(str(self.test_doc_path))
        self.assertIn('compliant', validation)
        print("  ✓ validate_document_compliance() 測試通過")

    def test_10_custom_config(self):
        """測試自訂配置"""
        print("\n[測試] 自訂配置...")

        custom_config = {
            "company_info": {
                "name": "Custom Company"
            },
            "fonts": {
                "title": {
                    "size": 20
                }
            }
        }

        analyzer = BestatStyleAnalyzer(config=custom_config)

        # 檢查配置合併
        self.assertEqual(analyzer.style_config['company_info']['name'], 'Custom Company')
        self.assertEqual(analyzer.style_config['fonts']['title']['size'], 20)

        # 檢查其他預設值保持不變
        self.assertEqual(analyzer.style_config['page_setup']['size'], 'A4')

        print("  ✓ 自訂配置合併成功")

    def test_11_error_handling(self):
        """測試錯誤處理"""
        print("\n[測試] 錯誤處理...")

        # 測試不存在的檔案
        with self.assertRaises(FileNotFoundError):
            self.analyzer.extract_styles_from_document("nonexistent.docx")
        print("  ✓ 檔案不存在錯誤處理正確")

        # 測試不存在的配置檔案
        with self.assertRaises(FileNotFoundError):
            self.analyzer.load_style_config("nonexistent.json")
        print("  ✓ 配置檔案不存在錯誤處理正確")

    @classmethod
    def tearDownClass(cls):
        """清理測試環境"""
        print(f"\n測試完成！測試檔案保留在: {cls.test_dir}")


class TestWordFormatterIntegration(unittest.TestCase):
    """測試與Word Formatter的整合"""

    def setUp(self):
        """設定測試環境"""
        self.test_dir = Path(__file__).parent.parent / "output" / "test_integration"
        self.test_dir.mkdir(parents=True, exist_ok=True)

    def test_word_formatter_bestat_integration(self):
        """測試Word Formatter整合"""
        print("\n[測試] Word Formatter 整合...")

        try:
            from word_formatter import WordFormatter

            # 建立Word Formatter實例
            formatter = WordFormatter()
            formatter.create_document()

            # 檢查Bestat分析器是否可用
            self.assertIsNotNone(formatter.bestat_analyzer)

            # 套用Bestat樣式
            formatter.apply_bestat_style(
                document_title="整合測試",
                protocol_number="PRO-2025-INT",
                version="1.0"
            )

            # 驗證規範
            validation = formatter.validate_bestat_compliance()
            self.assertIn('compliant', validation)

            # 儲存文件
            output_path = self.test_dir / "integrated_document.docx"
            formatter.save_document(str(output_path))

            self.assertTrue(output_path.exists())
            print(f"  ✓ 整合測試通過")
            print(f"  ✓ 文件已儲存: {output_path}")

        except ImportError:
            self.skipTest("Word Formatter 不可用")


def run_tests(verbosity=2):
    """執行所有測試"""
    print("=" * 70)
    print("Bestat樣式分析器測試套件")
    print("=" * 70)

    # 建立測試套件
    suite = unittest.TestSuite()

    # 添加測試
    suite.addTests(unittest.TestLoader().loadTestsFromTestCase(TestBestatStyleAnalyzer))
    suite.addTests(unittest.TestLoader().loadTestsFromTestCase(TestWordFormatterIntegration))

    # 執行測試
    runner = unittest.TextTestRunner(verbosity=verbosity)
    result = runner.run(suite)

    # 顯示總結
    print("\n" + "=" * 70)
    print("測試總結")
    print("=" * 70)
    print(f"執行測試: {result.testsRun}")
    print(f"成功: {result.testsRun - len(result.failures) - len(result.errors)}")
    print(f"失敗: {len(result.failures)}")
    print(f"錯誤: {len(result.errors)}")
    print(f"跳過: {len(result.skipped)}")

    return result


if __name__ == "__main__":
    run_tests(verbosity=2)
