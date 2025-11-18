"""
DVP (Data Validation Plan) Generator Module

This module provides automated generation of Data Validation Plans for clinical trials.
It creates comprehensive validation rules based on Protocol information and CRF design.

Author: Clinical Doc Automation Team
Date: 2025-11-18
"""

from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise ImportError(
        "python-docx is required. Install it with: pip install python-docx"
    )


class Severity(Enum):
    """Validation rule severity levels"""
    CRITICAL = "Critical"
    MAJOR = "Major"
    MINOR = "Minor"


class ValidationType(Enum):
    """Types of validation checks"""
    RANGE_CHECK = "Range Check"
    REQUIRED_FIELD = "Required Field"
    LOGICAL_CHECK = "Logical Check"
    CROSS_FORM = "Cross-Form Validation"
    DATE_CONSISTENCY = "Date Consistency"
    PROTOCOL_DEVIATION = "Protocol Deviation"
    CUSTOM = "Custom"


@dataclass
class ValidationRule:
    """Represents a single validation rule"""
    rule_id: str
    description: str
    severity: Severity
    query_text: str
    validation_type: ValidationType
    form_name: Optional[str] = None
    field_name: Optional[str] = None
    details: Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        """Validate rule data"""
        if not self.rule_id:
            raise ValueError("Rule ID cannot be empty")
        if not self.description:
            raise ValueError("Description cannot be empty")
        if not self.query_text:
            raise ValueError("Query text cannot be empty")


@dataclass
class ProtocolInfo:
    """Protocol information for DVP generation"""
    protocol_number: str
    protocol_title: str
    sponsor: str
    indication: str
    phase: str
    version: str = "1.0"
    date: str = ""

    def __post_init__(self):
        if not self.date:
            self.date = datetime.now().strftime("%Y-%m-%d")


@dataclass
class CRFField:
    """CRF field definition"""
    field_name: str
    field_label: str
    form_name: str
    data_type: str  # e.g., 'numeric', 'date', 'text', 'dropdown'
    required: bool = False
    min_value: Optional[float] = None
    max_value: Optional[float] = None
    valid_values: Optional[List[str]] = None
    date_format: Optional[str] = None
    units: Optional[str] = None


class DVPGenerator:
    """
    Data Validation Plan Generator

    Generates comprehensive DVP documents based on Protocol and CRF information.
    """

    def __init__(self, protocol_info: ProtocolInfo):
        """
        Initialize DVP Generator

        Args:
            protocol_info: Protocol information for the study
        """
        self.protocol_info = protocol_info
        self.validation_rules: List[ValidationRule] = []
        self.rule_counter = 0
        self.crf_fields: List[CRFField] = []

    def add_crf_fields(self, fields: List[CRFField]) -> None:
        """
        Add CRF field definitions

        Args:
            fields: List of CRF field definitions
        """
        self.crf_fields.extend(fields)

    def _generate_rule_id(self, validation_type: ValidationType) -> str:
        """Generate unique rule ID"""
        self.rule_counter += 1
        prefix_map = {
            ValidationType.RANGE_CHECK: "RNG",
            ValidationType.REQUIRED_FIELD: "REQ",
            ValidationType.LOGICAL_CHECK: "LOG",
            ValidationType.CROSS_FORM: "CRS",
            ValidationType.DATE_CONSISTENCY: "DAT",
            ValidationType.PROTOCOL_DEVIATION: "PRO",
            ValidationType.CUSTOM: "CUS"
        }
        prefix = prefix_map.get(validation_type, "VAL")
        return f"{prefix}-{self.rule_counter:04d}"

    def generate_range_checks(self) -> List[ValidationRule]:
        """Generate range check validations for numeric fields"""
        rules = []

        for field in self.crf_fields:
            if field.data_type == 'numeric' and (field.min_value is not None or field.max_value is not None):
                rule_id = self._generate_rule_id(ValidationType.RANGE_CHECK)

                # Build description
                desc_parts = []
                if field.min_value is not None and field.max_value is not None:
                    range_desc = f"between {field.min_value} and {field.max_value}"
                elif field.min_value is not None:
                    range_desc = f"greater than or equal to {field.min_value}"
                else:
                    range_desc = f"less than or equal to {field.max_value}"

                description = f"Check that {field.field_label} ({field.field_name}) is {range_desc}"
                if field.units:
                    description += f" {field.units}"

                # Build query text
                query_text = f"Please verify {field.field_label}. The value is outside the expected range ({range_desc}"
                if field.units:
                    query_text += f" {field.units}"
                query_text += ")."

                rule = ValidationRule(
                    rule_id=rule_id,
                    description=description,
                    severity=Severity.MAJOR,
                    query_text=query_text,
                    validation_type=ValidationType.RANGE_CHECK,
                    form_name=field.form_name,
                    field_name=field.field_name,
                    details={
                        'min_value': field.min_value,
                        'max_value': field.max_value,
                        'units': field.units
                    }
                )
                rules.append(rule)

        return rules

    def generate_required_field_checks(self) -> List[ValidationRule]:
        """Generate required field validations"""
        rules = []

        for field in self.crf_fields:
            if field.required:
                rule_id = self._generate_rule_id(ValidationType.REQUIRED_FIELD)

                description = f"Check that {field.field_label} ({field.field_name}) in {field.form_name} is not missing"
                query_text = f"Please provide the missing value for {field.field_label}."

                rule = ValidationRule(
                    rule_id=rule_id,
                    description=description,
                    severity=Severity.CRITICAL,
                    query_text=query_text,
                    validation_type=ValidationType.REQUIRED_FIELD,
                    form_name=field.form_name,
                    field_name=field.field_name
                )
                rules.append(rule)

        return rules

    def generate_date_consistency_checks(self) -> List[ValidationRule]:
        """Generate date consistency validations"""
        rules = []

        # Find date fields
        date_fields = [f for f in self.crf_fields if f.data_type == 'date']

        # Common date consistency checks
        date_checks = [
            {
                'field_pattern': r'inform.*consent.*date',
                'reference': 'study_start_date',
                'check': 'on_or_before',
                'description': 'Informed consent date should be on or before the study start date'
            },
            {
                'field_pattern': r'visit.*date',
                'reference': 'inform_consent_date',
                'check': 'on_or_after',
                'description': 'Visit date should be on or after informed consent date'
            },
            {
                'field_pattern': r'ae.*start.*date',
                'reference': 'ae_end_date',
                'check': 'on_or_before',
                'description': 'Adverse event start date should be on or before end date'
            }
        ]

        for field in date_fields:
            field_lower = field.field_name.lower()

            for check_config in date_checks:
                if re.search(check_config['field_pattern'], field_lower):
                    rule_id = self._generate_rule_id(ValidationType.DATE_CONSISTENCY)

                    query_text = f"Please verify the date for {field.field_label}. {check_config['description']}."

                    rule = ValidationRule(
                        rule_id=rule_id,
                        description=check_config['description'],
                        severity=Severity.MAJOR,
                        query_text=query_text,
                        validation_type=ValidationType.DATE_CONSISTENCY,
                        form_name=field.form_name,
                        field_name=field.field_name,
                        details={
                            'reference_field': check_config['reference'],
                            'check_type': check_config['check']
                        }
                    )
                    rules.append(rule)

        return rules

    def generate_logical_checks(self) -> List[ValidationRule]:
        """Generate logical consistency checks"""
        rules = []

        # Example: If answer is "Yes", details must be provided
        for field in self.crf_fields:
            field_lower = field.field_name.lower()

            # Check for Yes/No fields that might require follow-up
            if 'indicate' in field_lower or 'occurred' in field_lower:
                if field.valid_values and 'Yes' in field.valid_values:
                    rule_id = self._generate_rule_id(ValidationType.LOGICAL_CHECK)

                    description = f"If {field.field_label} is 'Yes', associated details must be provided"
                    query_text = f"Please provide details for {field.field_label} as it is marked as 'Yes'."

                    rule = ValidationRule(
                        rule_id=rule_id,
                        description=description,
                        severity=Severity.MAJOR,
                        query_text=query_text,
                        validation_type=ValidationType.LOGICAL_CHECK,
                        form_name=field.form_name,
                        field_name=field.field_name,
                        details={'condition': "value = 'Yes'"}
                    )
                    rules.append(rule)

        return rules

    def generate_cross_form_validations(self) -> List[ValidationRule]:
        """Generate cross-form validation checks"""
        rules = []

        # Group fields by form
        forms = {}
        for field in self.crf_fields:
            if field.form_name not in forms:
                forms[field.form_name] = []
            forms[field.form_name].append(field)

        # Example: Subject ID consistency across forms
        subject_id_fields = [f for f in self.crf_fields if 'subject' in f.field_name.lower() and 'id' in f.field_name.lower()]

        if len(subject_id_fields) > 1:
            rule_id = self._generate_rule_id(ValidationType.CROSS_FORM)

            form_names = ", ".join(set(f.form_name for f in subject_id_fields))
            description = f"Check that Subject ID is consistent across all forms: {form_names}"
            query_text = "Please verify the Subject ID. It appears to be inconsistent across forms."

            rule = ValidationRule(
                rule_id=rule_id,
                description=description,
                severity=Severity.CRITICAL,
                query_text=query_text,
                validation_type=ValidationType.CROSS_FORM,
                details={'forms': form_names}
            )
            rules.append(rule)

        return rules

    def generate_protocol_deviation_checks(self) -> List[ValidationRule]:
        """Generate protocol deviation checks"""
        rules = []

        # Example protocol deviation checks
        deviation_checks = [
            {
                'description': 'Check that inclusion criteria are met at screening',
                'query': 'Please verify that all inclusion criteria were met. A protocol deviation may have occurred.',
                'severity': Severity.CRITICAL
            },
            {
                'description': 'Check that exclusion criteria are not violated',
                'query': 'Please verify that no exclusion criteria were violated. A protocol deviation may have occurred.',
                'severity': Severity.CRITICAL
            },
            {
                'description': 'Check that visit windows are within protocol-specified ranges',
                'query': 'Please verify the visit date. It appears to be outside the protocol-specified window.',
                'severity': Severity.MAJOR
            },
            {
                'description': 'Check that prohibited medications were not taken',
                'query': 'Please verify the concomitant medication. It may be prohibited per protocol.',
                'severity': Severity.MAJOR
            }
        ]

        for check in deviation_checks:
            rule_id = self._generate_rule_id(ValidationType.PROTOCOL_DEVIATION)

            rule = ValidationRule(
                rule_id=rule_id,
                description=check['description'],
                severity=check['severity'],
                query_text=check['query'],
                validation_type=ValidationType.PROTOCOL_DEVIATION
            )
            rules.append(rule)

        return rules

    def add_custom_rule(self,
                       description: str,
                       query_text: str,
                       severity: Severity,
                       validation_type: ValidationType = ValidationType.CUSTOM,
                       form_name: Optional[str] = None,
                       field_name: Optional[str] = None,
                       details: Optional[Dict[str, Any]] = None) -> ValidationRule:
        """
        Add a custom validation rule

        Args:
            description: Rule description
            query_text: Query text to display when rule is triggered
            severity: Rule severity level
            validation_type: Type of validation
            form_name: Associated form name
            field_name: Associated field name
            details: Additional details

        Returns:
            Created ValidationRule
        """
        rule_id = self._generate_rule_id(validation_type)

        rule = ValidationRule(
            rule_id=rule_id,
            description=description,
            severity=severity,
            query_text=query_text,
            validation_type=validation_type,
            form_name=form_name,
            field_name=field_name,
            details=details or {}
        )

        self.validation_rules.append(rule)
        return rule

    def generate_all_rules(self) -> List[ValidationRule]:
        """Generate all standard validation rules"""
        all_rules = []

        all_rules.extend(self.generate_required_field_checks())
        all_rules.extend(self.generate_range_checks())
        all_rules.extend(self.generate_date_consistency_checks())
        all_rules.extend(self.generate_logical_checks())
        all_rules.extend(self.generate_cross_form_validations())
        all_rules.extend(self.generate_protocol_deviation_checks())

        self.validation_rules.extend(all_rules)
        return all_rules

    def _set_cell_background(self, cell, color: str):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _add_title_page(self, doc: Document):
        """Add title page to document"""
        # Title
        title = doc.add_heading('Data Validation Plan', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Protocol information
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ('Protocol Number:', self.protocol_info.protocol_number),
            ('Protocol Title:', self.protocol_info.protocol_title),
            ('Sponsor:', self.protocol_info.sponsor),
            ('Indication:', self.protocol_info.indication),
            ('Phase:', self.protocol_info.phase),
            ('DVP Version:', self.protocol_info.version)
        ]

        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            # Make label bold
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(info_table.rows[i].cells[0], 'D9E2F3')

        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_para.add_run(f'Date: {self.protocol_info.date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_page_break()

    def _add_introduction(self, doc: Document):
        """Add introduction section"""
        doc.add_heading('1. Introduction', 1)

        intro_text = f"""
This Data Validation Plan (DVP) describes the validation checks to be implemented for the clinical trial
"{self.protocol_info.protocol_title}" (Protocol {self.protocol_info.protocol_number}).

The purpose of this DVP is to ensure data quality and integrity throughout the study by defining:
• Required field validations
• Range checks for numeric and date fields
• Logical consistency checks
• Cross-form validations
• Protocol deviation checks

All validation rules defined in this document will be implemented in the Electronic Data Capture (EDC) system
and will generate queries when triggered.
        """.strip()

        doc.add_paragraph(intro_text)
        doc.add_paragraph()

    def _add_validation_rules_section(self, doc: Document):
        """Add validation rules section with table"""
        doc.add_heading('2. Validation Rules', 1)

        if not self.validation_rules:
            doc.add_paragraph('No validation rules defined.')
            return

        # Group rules by type
        rules_by_type = {}
        for rule in self.validation_rules:
            type_name = rule.validation_type.value
            if type_name not in rules_by_type:
                rules_by_type[type_name] = []
            rules_by_type[type_name].append(rule)

        # Add summary
        doc.add_heading('2.1 Summary', 2)
        summary_table = doc.add_table(rows=len(rules_by_type) + 1, cols=2)
        summary_table.style = 'Light Grid Accent 1'

        # Header
        summary_table.rows[0].cells[0].text = 'Validation Type'
        summary_table.rows[0].cells[1].text = 'Count'
        for cell in summary_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(cell, '4472C4')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Data
        for i, (type_name, rules) in enumerate(sorted(rules_by_type.items()), 1):
            summary_table.rows[i].cells[0].text = type_name
            summary_table.rows[i].cells[1].text = str(len(rules))

        doc.add_paragraph()

        # Add detailed rules for each type
        doc.add_heading('2.2 Detailed Validation Rules', 2)

        for type_name, rules in sorted(rules_by_type.items()):
            doc.add_heading(f'2.2.{list(rules_by_type.keys()).index(type_name) + 1} {type_name}', 3)

            # Create table for rules
            table = doc.add_table(rows=len(rules) + 1, cols=6)
            table.style = 'Light Grid Accent 1'

            # Set column widths
            widths = [Inches(0.8), Inches(1.2), Inches(1.0), Inches(2.5), Inches(0.8), Inches(1.5)]
            for i, width in enumerate(widths):
                for row in table.rows:
                    row.cells[i].width = width

            # Header row
            headers = ['Rule ID', 'Form', 'Field', 'Description', 'Severity', 'Query Text']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_background(cell, '4472C4')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

            # Data rows
            for i, rule in enumerate(rules, 1):
                table.rows[i].cells[0].text = rule.rule_id
                table.rows[i].cells[1].text = rule.form_name or 'N/A'
                table.rows[i].cells[2].text = rule.field_name or 'N/A'
                table.rows[i].cells[3].text = rule.description
                table.rows[i].cells[4].text = rule.severity.value
                table.rows[i].cells[5].text = rule.query_text

                # Color code severity
                severity_cell = table.rows[i].cells[4]
                if rule.severity == Severity.CRITICAL:
                    self._set_cell_background(severity_cell, 'FF0000')
                    severity_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                elif rule.severity == Severity.MAJOR:
                    self._set_cell_background(severity_cell, 'FFA500')
                else:
                    self._set_cell_background(severity_cell, 'FFFF00')

            doc.add_paragraph()

    def _add_appendix(self, doc: Document):
        """Add appendix section"""
        doc.add_heading('3. Appendix', 1)

        doc.add_heading('3.1 Severity Definitions', 2)

        severity_table = doc.add_table(rows=4, cols=2)
        severity_table.style = 'Light Grid Accent 1'

        # Header
        severity_table.rows[0].cells[0].text = 'Severity'
        severity_table.rows[0].cells[1].text = 'Definition'
        for cell in severity_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(cell, '4472C4')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Data
        definitions = [
            ('Critical', 'Issues that must be resolved immediately. May impact subject safety or data integrity.'),
            ('Major', 'Significant issues that should be addressed promptly. May affect data quality.'),
            ('Minor', 'Issues that should be reviewed but do not significantly impact data quality.')
        ]

        for i, (severity, definition) in enumerate(definitions, 1):
            severity_table.rows[i].cells[0].text = severity
            severity_table.rows[i].cells[1].text = definition

    def generate_dvp_document(self, output_path: str) -> str:
        """
        Generate complete DVP document

        Args:
            output_path: Path to save the generated .docx file

        Returns:
            Path to the generated document
        """
        # Create new document
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add sections
        self._add_title_page(doc)
        self._add_introduction(doc)
        self._add_validation_rules_section(doc)
        self._add_appendix(doc)

        # Save document
        doc.save(output_path)

        return output_path

    def export_rules_to_dict(self) -> List[Dict[str, Any]]:
        """Export validation rules to dictionary format"""
        return [
            {
                'rule_id': rule.rule_id,
                'description': rule.description,
                'severity': rule.severity.value,
                'query_text': rule.query_text,
                'validation_type': rule.validation_type.value,
                'form_name': rule.form_name,
                'field_name': rule.field_name,
                'details': rule.details
            }
            for rule in self.validation_rules
        ]

    def get_rules_summary(self) -> Dict[str, int]:
        """Get summary of rules by type"""
        summary = {}
        for rule in self.validation_rules:
            type_name = rule.validation_type.value
            summary[type_name] = summary.get(type_name, 0) + 1
        return summary


# Convenience function
def create_dvp(protocol_info: ProtocolInfo,
               crf_fields: List[CRFField],
               output_path: str,
               custom_rules: Optional[List[Dict[str, Any]]] = None) -> str:
    """
    Convenience function to create a DVP document

    Args:
        protocol_info: Protocol information
        crf_fields: List of CRF field definitions
        output_path: Path to save the DVP document
        custom_rules: Optional list of custom rules

    Returns:
        Path to the generated document
    """
    generator = DVPGenerator(protocol_info)
    generator.add_crf_fields(crf_fields)
    generator.generate_all_rules()

    # Add custom rules if provided
    if custom_rules:
        for rule_data in custom_rules:
            generator.add_custom_rule(
                description=rule_data['description'],
                query_text=rule_data['query_text'],
                severity=Severity[rule_data.get('severity', 'MAJOR')],
                validation_type=ValidationType[rule_data.get('validation_type', 'CUSTOM')],
                form_name=rule_data.get('form_name'),
                field_name=rule_data.get('field_name'),
                details=rule_data.get('details')
            )

    return generator.generate_dvp_document(output_path)
