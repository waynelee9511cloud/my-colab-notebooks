#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Bestat樣式分析器演示腳本

這個腳本展示了Bestat樣式分析器的核心功能，
包括建立文件、套用樣式、驗證規範等。

運行方式：
    python demo_bestat.py

作者：Clinical Document Automation Team
版本：1.0.0
"""

import os
import sys
from pathlib import Path

# 添加modules路徑
sys.path.insert(0, str(Path(__file__).parent / "modules"))

from docx import Document
from modules.bestat_style_analyzer import BestatStyleAnalyzer
from modules.word_formatter import WordFormatter


def print_section(title):
    """列印章節標題"""
    print("\n" + "=" * 70)
    print(f" {title}")
    print("=" * 70)


def demo_01_basic_style_application():
    """演示1：基本樣式套用"""
    print_section("演示1：基本Bestat樣式套用")

    # 建立輸出目錄
    output_dir = Path(__file__).parent / "output" / "demo"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 初始化分析器
    print("\n1. 初始化Bestat樣式分析器...")
    analyzer = BestatStyleAnalyzer()
    print(f"   ✓ 公司名稱: {analyzer.style_config['company_info']['name']}")
    print(f"   ✓ 主色調: RGB{tuple(analyzer.style_config['colors']['primary'].values())}")

    # 建立文件
    print("\n2. 建立Word文件...")
    doc = Document()
    doc.add_paragraph("第一章 研究背景")
    doc.add_paragraph("本臨床試驗旨在評估新藥的安全性和有效性。")
    doc.add_paragraph("第二章 研究設計")
    doc.add_paragraph("這是一項隨機、雙盲、安慰劑對照的臨床試驗。")
    print("   ✓ 已添加內容")

    # 套用Bestat樣式
    print("\n3. 套用Bestat標準樣式...")
    styled_doc = analyzer.apply_bestat_style(
        doc,
        document_title="第三期臨床試驗Protocol",
        protocol_number="PRO-2025-DEMO-001",
        version="1.0"
    )
    print("   ✓ 已套用Bestat樣式")

    # 儲存文件
    output_file = output_dir / "demo_01_basic.docx"
    styled_doc.save(str(output_file))
    print(f"\n4. 文件已儲存: {output_file}")


def demo_02_validation():
    """演示2：規範驗證"""
    print_section("演示2：文件規範驗證")

    output_dir = Path(__file__).parent / "output" / "demo"
    output_dir.mkdir(parents=True, exist_ok=True)

    analyzer = BestatStyleAnalyzer()

    # 建立符合規範的文件
    print("\n1. 建立符合Bestat規範的文件...")
    doc1 = Document()
    doc1.add_paragraph("符合規範的測試文件")
    styled_doc1 = analyzer.apply_bestat_style(
        doc1,
        document_title="規範文件",
        protocol_number="PRO-2025-COMPLIANT",
        version="1.0"
    )

    # 驗證
    print("\n2. 驗證文件規範...")
    validation = analyzer.validate_bestat_compliance(styled_doc1)

    print(f"\n   驗證結果:")
    print(f"   ✓ 符合規範: {validation['compliant']}")
    print(f"   ✓ 問題數: {validation['total_issues']}")
    print(f"   ✓ 警告數: {validation['total_warnings']}")

    if validation['issues']:
        print("\n   發現的問題:")
        for issue in validation['issues']:
            print(f"     - {issue}")

    if validation['warnings']:
        print("\n   發現的警告:")
        for warning in validation['warnings']:
            print(f"     - {warning}")

    # 儲存文件
    output_file = output_dir / "demo_02_validated.docx"
    styled_doc1.save(str(output_file))
    print(f"\n3. 已儲存: {output_file}")


def demo_03_word_formatter_integration():
    """演示3：Word Formatter整合"""
    print_section("演示3：與Word Formatter整合")

    output_dir = Path(__file__).parent / "output" / "demo"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 使用Word Formatter
    print("\n1. 建立Word Formatter實例...")
    formatter = WordFormatter()
    formatter.create_document()
    print("   ✓ Word Formatter已初始化")

    # 添加結構化內容
    print("\n2. 添加結構化內容...")
    formatter.apply_title_style("臨床試驗Protocol", level=1)
    formatter.apply_title_style("1. 試驗目的", level=2)
    formatter.apply_body_style(
        "評估新藥XYZ在治療疾病ABC中的安全性和有效性。",
        alignment='justify'
    )

    formatter.apply_title_style("2. 試驗設計", level=2)
    formatter.apply_body_style(
        "這是一項多中心、隨機、雙盲、安慰劑對照的第三期臨床試驗。",
        alignment='justify'
    )
    print("   ✓ 內容已添加")

    # 套用Bestat樣式
    print("\n3. 套用Bestat樣式...")
    formatter.apply_bestat_style(
        document_title="第三期臨床試驗Protocol",
        protocol_number="PRO-2025-INTEGRATED",
        version="2.0"
    )
    print("   ✓ Bestat樣式已套用")

    # 驗證
    print("\n4. 驗證規範...")
    validation = formatter.validate_bestat_compliance()
    print(f"   ✓ 符合規範: {validation['compliant']}")

    # 儲存
    output_file = output_dir / "demo_03_integrated.docx"
    formatter.save_document(str(output_file))
    print(f"\n5. 已儲存: {output_file}")


def demo_04_custom_config():
    """演示4：自訂配置"""
    print_section("演示4：使用自訂Bestat配置")

    output_dir = Path(__file__).parent / "output" / "demo"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 定義自訂配置
    print("\n1. 定義自訂配置（Taiwan分公司）...")
    custom_config = {
        "company_info": {
            "name": "Bestat Taiwan Branch",
            "full_name": "Bestat Clinical Research - Taiwan Office"
        },
        "fonts": {
            "title": {
                "size": 20,  # 更大的標題
                "color": {"r": 0, "g": 102, "b": 204}  # 亮藍色
            }
        },
        "colors": {
            "primary": {"r": 0, "g": 102, "b": 204}  # 使用亮藍作為主色
        }
    }
    print("   ✓ 自訂配置已定義")

    # 初始化分析器
    print("\n2. 使用自訂配置初始化...")
    analyzer = BestatStyleAnalyzer(config=custom_config)
    print(f"   ✓ 公司名稱: {analyzer.style_config['company_info']['name']}")
    print(f"   ✓ 標題大小: {analyzer.style_config['fonts']['title']['size']}pt")

    # 建立文件
    print("\n3. 建立文件並套用自訂樣式...")
    doc = Document()
    doc.add_paragraph("台灣臨床試驗Protocol")
    doc.add_paragraph("本Protocol適用於台灣地區的臨床試驗。")

    styled_doc = analyzer.apply_bestat_style(
        doc,
        document_title="台灣第三期臨床試驗",
        protocol_number="TW-PRO-2025-001",
        version="1.0-TW"
    )

    # 儲存
    output_file = output_dir / "demo_04_custom.docx"
    styled_doc.save(str(output_file))
    print(f"\n4. 已儲存: {output_file}")


def demo_05_config_management():
    """演示5：配置管理"""
    print_section("演示5：樣式配置管理")

    output_dir = Path(__file__).parent / "output" / "demo"
    output_dir.mkdir(parents=True, exist_ok=True)

    analyzer = BestatStyleAnalyzer()

    # 儲存預設配置
    print("\n1. 儲存Bestat預設配置...")
    config_file = output_dir / "bestat_config.json"
    analyzer.save_style_config(str(config_file))
    print(f"   ✓ 配置已儲存: {config_file}")

    # 載入配置
    print("\n2. 載入配置...")
    new_analyzer = BestatStyleAnalyzer()
    new_analyzer.load_style_config(str(config_file))
    print(f"   ✓ 配置已載入")

    # 驗證配置
    print(f"\n3. 驗證配置內容...")
    print(f"   ✓ 公司名稱: {new_analyzer.style_config['company_info']['name']}")
    print(f"   ✓ 頁面大小: {new_analyzer.style_config['page_setup']['size']}")
    print(f"   ✓ 標題字體: {new_analyzer.style_config['fonts']['title']['name']}")


def demo_summary():
    """顯示演示總結"""
    print_section("演示總結")

    output_dir = Path(__file__).parent / "output" / "demo"

    print("\n演示完成！以下功能已展示：")
    print("\n1. ✓ 基本Bestat樣式套用")
    print("2. ✓ 文件規範驗證")
    print("3. ✓ Word Formatter整合")
    print("4. ✓ 自訂配置使用")
    print("5. ✓ 配置管理功能")

    print(f"\n產生的檔案位於: {output_dir}")
    print("\n產生的檔案清單:")
    if output_dir.exists():
        for file in sorted(output_dir.glob("demo_*.docx")):
            print(f"  - {file.name}")
        for file in sorted(output_dir.glob("*.json")):
            print(f"  - {file.name}")

    print("\n更多資訊請參考:")
    print("  - BESTAT_STYLE_GUIDE.md      (完整使用指南)")
    print("  - BESTAT_QUICK_REFERENCE.md  (快速參考)")
    print("  - examples/bestat_style_example.py  (詳細範例)")


def main():
    """主函數"""
    print("=" * 70)
    print(" Bestat樣式分析器演示")
    print("=" * 70)
    print("\n本演示將展示Bestat樣式分析器的主要功能。")
    print("所有產生的檔案將儲存在 output/demo/ 目錄。")

    try:
        # 執行所有演示
        demo_01_basic_style_application()
        demo_02_validation()
        demo_03_word_formatter_integration()
        demo_04_custom_config()
        demo_05_config_management()

        # 顯示總結
        demo_summary()

    except Exception as e:
        print(f"\n錯誤: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
