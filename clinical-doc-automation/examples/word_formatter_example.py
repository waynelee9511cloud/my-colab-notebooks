"""
Word文件格式化引擎使用範例
展示如何使用WordFormatter類別建立符合公司規範的臨床試驗文件

使用前請先安裝依賴：
    pip install python-docx pillow
"""

import sys
import os

# 添加模組路徑
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

from modules.word_formatter import WordFormatter, create_clinical_document
from docx.shared import RGBColor, Pt, Inches
from datetime import datetime


def example_1_basic_usage():
    """
    範例1: 基本使用 - 建立簡單的臨床試驗文件
    """
    print("\n" + "=" * 60)
    print("範例1: 基本使用")
    print("=" * 60)

    # 建立格式化引擎實例
    formatter = WordFormatter()

    # 建立新文件
    formatter.create_document()

    # 設定頁面格式
    formatter.set_page_format(
        page_size='A4',
        orientation='portrait'
    )

    # 設定頁首
    formatter.set_header(
        document_title="Clinical Study Protocol",
        protocol_number="PRO-2025-001",
        version="1.0"
    )

    # 設定頁尾
    formatter.set_footer(
        confidential=True,
        include_page_numbers=True,
        include_date=True
    )

    # 添加標題
    formatter.apply_title_style("Protocol Synopsis", level=1)

    # 添加副標題
    formatter.apply_title_style("1. Study Overview", level=2)

    # 添加內文
    formatter.apply_body_style(
        "This is a Phase III, randomized, double-blind, placebo-controlled study "
        "to evaluate the efficacy and safety of the investigational product.",
        alignment='justify'
    )

    # 添加表格
    table = formatter.doc.add_table(rows=4, cols=2)
    formatter.apply_table_style(table)

    # 填入表格內容
    headers = ['Item', 'Description']
    data = [
        ['Study Phase', 'Phase III'],
        ['Study Design', 'Randomized, Double-Blind, Placebo-Controlled'],
        ['Number of Subjects', 'Approximately 300']
    ]

    # 設定表頭
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # 填入資料
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    # 儲存文件
    output_path = os.path.join(os.path.dirname(__file__), '..', 'output', 'example_1_basic.docx')
    formatter.save_document(output_path)

    print(f"範例1完成！文件已儲存至: {output_path}")


def example_2_complete_protocol():
    """
    範例2: 完整的Protocol文件 - 使用便利函數
    """
    print("\n" + "=" * 60)
    print("範例2: 完整的Protocol文件")
    print("=" * 60)

    # 使用便利函數快速建立文件
    formatter = create_clinical_document(
        document_title="Clinical Study Protocol",
        protocol_number="ABC-2025-001",
        version="2.0",
        sponsor="ABC Pharmaceutical Company",
        indication="Type 2 Diabetes Mellitus",
        output_path=os.path.join(os.path.dirname(__file__), '..', 'output', 'example_2_complete.docx')
    )

    # 添加目錄頁（手動添加，實際使用時可以用Word的內建目錄功能）
    formatter.apply_title_style("TABLE OF CONTENTS", level=1)
    formatter.apply_body_style("1. Protocol Synopsis .................. 3")
    formatter.apply_body_style("2. Study Objectives ................... 4")
    formatter.apply_body_style("3. Study Design ....................... 5")
    formatter.apply_body_style("4. Study Population ................... 6")
    formatter.apply_body_style("5. Treatment Plan ..................... 7")
    formatter.doc.add_page_break()

    # 第一章：Protocol Synopsis
    formatter.apply_title_style("1. PROTOCOL SYNOPSIS", level=1)

    # 建立Synopsis表格
    synopsis_table = formatter.doc.add_table(rows=11, cols=2)
    formatter.apply_table_style(synopsis_table, style_name='Light Grid Accent 1')

    synopsis_data = [
        ['Protocol Title', 'A Phase III Study to Evaluate the Efficacy and Safety of ABC-001'],
        ['Protocol Number', 'ABC-2025-001'],
        ['Study Phase', 'Phase III'],
        ['Sponsor', 'ABC Pharmaceutical Company'],
        ['Indication', 'Type 2 Diabetes Mellitus'],
        ['Study Design', 'Randomized, Double-Blind, Placebo-Controlled, Parallel Group'],
        ['Sample Size', 'Approximately 300 subjects (150 per arm)'],
        ['Treatment Duration', '24 weeks'],
        ['Primary Endpoint', 'Change from baseline in HbA1c at Week 24'],
        ['Secondary Endpoints', 'Change in fasting plasma glucose, body weight, and safety parameters']
    ]

    # 設定表頭
    synopsis_table.rows[0].cells[0].text = 'Item'
    synopsis_table.rows[0].cells[1].text = 'Description'

    # 填入資料
    for idx, (item, description) in enumerate(synopsis_data, start=1):
        synopsis_table.rows[idx].cells[0].text = item
        synopsis_table.rows[idx].cells[1].text = description

    formatter.doc.add_page_break()

    # 第二章：Study Objectives
    formatter.apply_title_style("2. STUDY OBJECTIVES", level=1)

    formatter.apply_title_style("2.1 Primary Objective", level=2)
    formatter.apply_body_style(
        "To demonstrate the superiority of ABC-001 compared to placebo in reducing "
        "HbA1c levels in subjects with Type 2 Diabetes Mellitus after 24 weeks of treatment.",
        alignment='justify'
    )

    formatter.apply_title_style("2.2 Secondary Objectives", level=2)
    formatter.apply_body_style("The secondary objectives of this study are:", alignment='justify')

    # 添加項目列表
    objectives = [
        "To evaluate the effect of ABC-001 on fasting plasma glucose levels",
        "To assess the safety and tolerability of ABC-001",
        "To evaluate the effect of ABC-001 on body weight",
        "To assess the quality of life in subjects treated with ABC-001"
    ]

    for obj in objectives:
        paragraph = formatter.doc.add_paragraph(obj, style='List Bullet')
        paragraph.paragraph_format.left_indent = Inches(0.5)

    formatter.doc.add_page_break()

    # 第三章：Study Design
    formatter.apply_title_style("3. STUDY DESIGN", level=1)

    formatter.apply_title_style("3.1 Overall Study Design", level=2)
    formatter.apply_body_style(
        "This is a Phase III, multicenter, randomized, double-blind, placebo-controlled, "
        "parallel-group study. Approximately 300 subjects will be randomized in a 1:1 ratio "
        "to receive either ABC-001 or placebo for 24 weeks.",
        alignment='justify'
    )

    formatter.apply_title_style("3.2 Study Duration", level=2)

    # 建立Study Duration表格
    duration_table = formatter.doc.add_table(rows=5, cols=2)
    formatter.apply_table_style(duration_table)

    duration_data = [
        ['Phase', 'Duration'],
        ['Screening Period', 'Up to 4 weeks'],
        ['Treatment Period', '24 weeks'],
        ['Follow-up Period', '4 weeks'],
        ['Total Study Duration', 'Up to 32 weeks']
    ]

    for row_idx, (phase, duration) in enumerate(duration_data):
        duration_table.rows[row_idx].cells[0].text = phase
        duration_table.rows[row_idx].cells[1].text = duration

    formatter.doc.add_page_break()

    # 第四章：Study Population
    formatter.apply_title_style("4. STUDY POPULATION", level=1)

    formatter.apply_title_style("4.1 Inclusion Criteria", level=2)
    formatter.apply_body_style("Subjects must meet all of the following criteria:", alignment='justify')

    inclusion_criteria = [
        "Male or female subjects aged 18-75 years",
        "Diagnosis of Type 2 Diabetes Mellitus for at least 6 months",
        "HbA1c ≥7.5% and ≤10.5% at screening",
        "BMI ≥18.5 and ≤40 kg/m²",
        "Willing and able to provide informed consent"
    ]

    for criterion in inclusion_criteria:
        paragraph = formatter.doc.add_paragraph(criterion, style='List Number')
        paragraph.paragraph_format.left_indent = Inches(0.5)

    formatter.doc.add_paragraph()

    formatter.apply_title_style("4.2 Exclusion Criteria", level=2)
    formatter.apply_body_style("Subjects meeting any of the following criteria will be excluded:", alignment='justify')

    exclusion_criteria = [
        "Type 1 Diabetes Mellitus or secondary forms of diabetes",
        "History of diabetic ketoacidosis or hyperosmolar hyperglycemic state",
        "Significant cardiovascular disease within 6 months",
        "Severe renal impairment (eGFR <30 mL/min/1.73m²)",
        "Active liver disease or ALT/AST >3× upper limit of normal",
        "Pregnancy or breastfeeding"
    ]

    for criterion in exclusion_criteria:
        paragraph = formatter.doc.add_paragraph(criterion, style='List Number')
        paragraph.paragraph_format.left_indent = Inches(0.5)

    # 儲存文件
    output_path = os.path.join(os.path.dirname(__file__), '..', 'output', 'example_2_complete.docx')
    formatter.save_document(output_path)

    print(f"範例2完成！文件已儲存至: {output_path}")


def example_3_custom_configuration():
    """
    範例3: 自訂配置 - 使用公司專屬的格式設定
    """
    print("\n" + "=" * 60)
    print("範例3: 自訂配置")
    print("=" * 60)

    # 自訂配置
    custom_config = {
        # 使用Arial字體而非Calibri
        'title_font': 'Arial',
        'body_font': 'Arial',
        'heading1_font': 'Arial',
        'heading2_font': 'Arial',

        # 調整字體大小
        'title_size': 18,
        'body_size': 12,

        # 自訂顏色（使用公司品牌色）
        'title_color': RGBColor(0, 102, 204),  # 藍色
        'heading1_color': RGBColor(0, 102, 204),

        # 調整邊距
        'margin_top': Inches(1.25),
        'margin_bottom': Inches(1.25),
        'margin_left': Inches(1.5),
        'margin_right': Inches(1.0),

        # 調整行距
        'line_spacing': 1.5,

        # 自訂公司資訊
        'company_name': 'XYZ Clinical Research',
        'confidential_text': 'CONFIDENTIAL & PROPRIETARY',
    }

    # 建立使用自訂配置的格式化引擎
    formatter = WordFormatter(config=custom_config)
    formatter.create_document()

    # 套用臨床試驗範本
    formatter.apply_clinical_trial_template(
        document_title="Investigator's Brochure",
        protocol_number="XYZ-2025-IB-001",
        version="3.0",
        sponsor="XYZ Clinical Research Organization"
    )

    # 添加內容
    formatter.apply_title_style("1. INTRODUCTION", level=1)
    formatter.apply_body_style(
        "This Investigator's Brochure (IB) presents a comprehensive summary of the "
        "nonclinical and clinical data available for the investigational product XYZ-001.",
        alignment='justify'
    )

    formatter.apply_title_style("1.1 Chemical Name and Structure", level=2)
    formatter.apply_body_style(
        "Chemical Name: (2S,3R,4S,5S,6R)-2-(4-chloro-3-(4-ethoxybenzyl)phenyl)-6-(hydroxymethyl)tetrahydro-2H-pyran-3,4,5-triol",
        alignment='left'
    )

    # 顯示當前配置
    print("\n當前配置：")
    config = formatter.get_config()
    for key, value in config.items():
        if not key.startswith('_'):
            print(f"  {key}: {value}")

    # 儲存文件
    output_path = os.path.join(os.path.dirname(__file__), '..', 'output', 'example_3_custom.docx')
    formatter.save_document(output_path)

    print(f"\n範例3完成！文件已儲存至: {output_path}")


def example_4_with_logo():
    """
    範例4: 包含Logo的文件
    注意：需要準備一個logo.png文件
    """
    print("\n" + "=" * 60)
    print("範例4: 包含Logo的文件")
    print("=" * 60)

    # 檢查是否有logo文件
    logo_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'company_logo.png')

    if not os.path.exists(logo_path):
        print(f"警告：找不到Logo文件: {logo_path}")
        print("將建立不含Logo的文件")
        logo_path = None

    # 建立文件
    formatter = create_clinical_document(
        document_title="Clinical Study Report",
        protocol_number="CSR-2025-001",
        version="Final 1.0",
        sponsor="Global Pharma Inc.",
        indication="Advanced Non-Small Cell Lung Cancer",
        logo_path=logo_path,
        output_path=os.path.join(os.path.dirname(__file__), '..', 'output', 'example_4_with_logo.docx')
    )

    # 添加執行摘要
    formatter.apply_title_style("EXECUTIVE SUMMARY", level=1)

    formatter.apply_title_style("Study Objectives", level=2)
    formatter.apply_body_style(
        "The primary objective of this study was to evaluate the overall survival (OS) "
        "in subjects with advanced non-small cell lung cancer (NSCLC) treated with "
        "the investigational product compared to standard of care.",
        alignment='justify'
    )

    formatter.apply_title_style("Study Design", level=2)
    formatter.apply_body_style(
        "This was a Phase III, randomized, open-label, multicenter study conducted "
        "in approximately 50 centers across 10 countries.",
        alignment='justify'
    )

    # 儲存
    output_path = os.path.join(os.path.dirname(__file__), '..', 'output', 'example_4_with_logo.docx')
    formatter.save_document(output_path)

    print(f"範例4完成！文件已儲存至: {output_path}")


def example_5_load_from_template():
    """
    範例5: 從現有範本載入並套用格式
    """
    print("\n" + "=" * 60)
    print("範例5: 從現有範本載入")
    print("=" * 60)

    # 檢查是否有範本文件
    template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'protocol_template.docx')

    formatter = WordFormatter()

    if os.path.exists(template_path):
        # 從範本載入
        formatter.create_document(template_path=template_path)
        print(f"已從範本載入: {template_path}")

        # 載入範本樣式
        styles = formatter.load_styles_from_template(template_path)
        print(f"載入了 {len(styles)} 個樣式")
    else:
        print(f"找不到範本文件: {template_path}")
        print("將建立新文件")
        formatter.create_document()

    # 設定基本格式
    formatter.set_page_format()
    formatter.set_header(
        document_title="Informed Consent Form",
        protocol_number="ICF-2025-001",
        version="2.0"
    )
    formatter.set_footer()

    # 添加內容
    formatter.apply_title_style("INFORMED CONSENT FORM", level=1)

    formatter.apply_title_style("Study Title", level=2)
    formatter.apply_body_style(
        "A Phase III Study to Evaluate the Efficacy and Safety of ABC-001 in Subjects with Type 2 Diabetes Mellitus",
        alignment='left'
    )

    formatter.apply_title_style("1. Introduction", level=2)
    formatter.apply_body_style(
        "You are being asked to participate in a research study. Before you decide whether "
        "to participate, it is important that you understand why the research is being done "
        "and what it will involve. Please take the time to read the following information carefully.",
        alignment='justify'
    )

    # 儲存
    output_path = os.path.join(os.path.dirname(__file__), '..', 'output', 'example_5_from_template.docx')
    formatter.save_document(output_path)

    print(f"範例5完成！文件已儲存至: {output_path}")


def example_6_complex_table():
    """
    範例6: 複雜表格格式化
    """
    print("\n" + "=" * 60)
    print("範例6: 複雜表格格式化")
    print("=" * 60)

    formatter = WordFormatter()
    formatter.create_document()
    formatter.set_page_format()

    # 設定頁首頁尾
    formatter.set_header(
        document_title="Study Visit Schedule",
        protocol_number="SVS-2025-001",
        version="1.0"
    )
    formatter.set_footer()

    # 標題
    formatter.apply_title_style("SCHEDULE OF ASSESSMENTS", level=1)

    # 建立複雜的訪視時程表
    table = formatter.doc.add_table(rows=11, cols=8)
    formatter.apply_table_style(table, style_name='Light Grid Accent 1')

    # 設定表頭
    headers = ['Assessment', 'Screening', 'Baseline', 'Week 4', 'Week 8', 'Week 12', 'Week 24', 'Follow-up']
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        # 設定表頭背景色（需要額外設定）

    # 填入資料
    assessments = [
        ['Informed Consent', 'X', '', '', '', '', '', ''],
        ['Demographics', 'X', '', '', '', '', '', ''],
        ['Medical History', 'X', '', '', '', '', '', ''],
        ['Physical Exam', 'X', 'X', '', 'X', '', 'X', 'X'],
        ['Vital Signs', 'X', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['ECG', 'X', 'X', '', '', '', 'X', ''],
        ['Laboratory Tests', 'X', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['HbA1c', 'X', 'X', '', '', '', 'X', ''],
        ['AE Assessment', '', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['Drug Dispensing', '', 'X', 'X', 'X', 'X', '', '']
    ]

    for row_idx, row_data in enumerate(assessments, start=1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    # 添加註解
    formatter.doc.add_paragraph()
    formatter.apply_body_style("Note: X indicates that the assessment should be performed at that visit.", alignment='left')
    formatter.apply_body_style("AE = Adverse Event; ECG = Electrocardiogram; HbA1c = Glycated Hemoglobin", alignment='left')

    # 儲存
    output_path = os.path.join(os.path.dirname(__file__), '..', 'output', 'example_6_complex_table.docx')
    formatter.save_document(output_path)

    print(f"範例6完成！文件已儲存至: {output_path}")


def run_all_examples():
    """
    執行所有範例
    """
    print("\n" + "=" * 60)
    print("Word文件格式化引擎 - 使用範例")
    print("=" * 60)

    # 確保輸出目錄存在
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'output')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"已建立輸出目錄: {output_dir}")

    try:
        example_1_basic_usage()
        example_2_complete_protocol()
        example_3_custom_configuration()
        example_4_with_logo()
        example_5_load_from_template()
        example_6_complex_table()

        print("\n" + "=" * 60)
        print("所有範例執行完成！")
        print("=" * 60)
        print(f"\n生成的文件位於: {output_dir}")

    except Exception as e:
        print(f"\n錯誤：{str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='Word文件格式化引擎使用範例')
    parser.add_argument('--example', type=int, choices=range(1, 7),
                       help='執行特定範例 (1-6)')
    parser.add_argument('--all', action='store_true',
                       help='執行所有範例')

    args = parser.parse_args()

    if args.all:
        run_all_examples()
    elif args.example:
        example_functions = {
            1: example_1_basic_usage,
            2: example_2_complete_protocol,
            3: example_3_custom_configuration,
            4: example_4_with_logo,
            5: example_5_load_from_template,
            6: example_6_complex_table
        }
        example_functions[args.example]()
    else:
        # 預設執行所有範例
        run_all_examples()
