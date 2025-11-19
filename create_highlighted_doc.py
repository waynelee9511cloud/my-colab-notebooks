#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成帶有語法高亮的 MedGemma 訓練教學文件
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pygments import lex
from pygments.lexers import PythonLexer
from pygments.token import Token
import os

def get_token_color(token_type):
    """根據 token 類型返回對應的顏色"""
    # 使用類似 VS Code Dark+ 主題的顏色方案
    color_map = {
        Token.Keyword: RGBColor(197, 134, 192),          # 紫色 - 關鍵字 (if, for, def, import)
        Token.Keyword.Namespace: RGBColor(197, 134, 192), # 紫色 - import, from
        Token.Name.Builtin: RGBColor(78, 201, 176),      # 青色 - 內建函數 (print, len, range)
        Token.Name.Function: RGBColor(220, 220, 170),    # 淺黃色 - 函數名
        Token.Name.Class: RGBColor(78, 201, 176),        # 青色 - 類名
        Token.String: RGBColor(206, 145, 120),           # 橘色 - 字串
        Token.String.Doc: RGBColor(106, 153, 85),        # 綠色 - 文檔字串
        Token.Comment: RGBColor(106, 153, 85),           # 綠色 - 註解
        Token.Number: RGBColor(181, 206, 168),           # 淺綠色 - 數字
        Token.Operator: RGBColor(212, 212, 212),         # 淺灰色 - 運算符
        Token.Name: RGBColor(156, 220, 254),             # 淺藍色 - 變數名
        Token.Punctuation: RGBColor(212, 212, 212),      # 淺灰色 - 標點符號
    }

    # 遞迴查找最接近的 token 類型
    while token_type not in color_map and token_type.parent:
        token_type = token_type.parent

    return color_map.get(token_type, RGBColor(212, 212, 212))  # 預設淺灰色

def add_highlighted_code(doc, code):
    """添加有語法高亮的程式碼"""
    # 創建段落
    p = doc.add_paragraph()
    p.style = 'Normal'

    # 設定段落背景色和邊框（模擬程式碼區塊）
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 使用 Pygments 解析程式碼
    lexer = PythonLexer()
    tokens = lex(code, lexer)

    # 為每個 token 創建不同顏色的 run
    for token_type, token_value in tokens:
        run = p.add_run(token_value)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = get_token_color(token_type)

    return p

def add_heading_custom(doc, text, level):
    """添加自訂標題"""
    h = doc.add_heading(text, level=level)
    if level == 1:
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_text(doc, text, bullet=False):
    """添加文字"""
    if bullet:
        p = doc.add_paragraph(text, style='List Bullet')
    else:
        p = doc.add_paragraph(text)
    return p

def create_document():
    """創建完整的訓練文件"""
    doc = Document()

    # 設定預設字體
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ========== 封面 ==========
    title = doc.add_heading('MedGemma 醫療術語校正模型', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Python 機器學習完整訓練教學文件', 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph('Notebook 20 個 Cells 完整講解')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(18)
        run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph('帶有語法高亮的程式碼範例')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ========== 目錄 ==========
    doc.add_heading('目錄', 1)
    cells_info = [
        ('Cell 1', '安裝套件'),
        ('Cell 2', '導入函式庫'),
        ('Cell 3', 'GPU 檢查'),
        ('Cell 4', '掛載 Google Drive'),
        ('Cell 5', '上傳檔案'),
        ('Cell 6', '讀取台灣藥物名稱'),
        ('Cell 7', '讀取 Excel 資料'),
        ('Cell 8', '整合台灣藥物名稱'),
        ('Cell 9', '資料平衡'),
        ('Cell 10', '設定訓練參數'),
        ('Cell 11', '準備訓練資料格式並初始化交叉驗證'),
        ('Cell 12', 'Fold 1 訓練（完整流程）'),
        ('Cell 13-16', 'Fold 2-5 訓練'),
        ('Cell 17', '統計分析'),
        ('Cell 18', '混淆矩陣視覺化'),
        ('Cell 19', '性能指標視覺化'),
        ('Cell 20', '生成詳細報告'),
    ]

    for cell_num, cell_desc in cells_info:
        add_text(doc, f'{cell_num}: {cell_desc}', bullet=True)

    doc.add_page_break()

    # ========== Cell 1 ==========
    add_heading_custom(doc, 'Cell 1: 安裝套件', 1)

    doc.add_heading('1.1 完整程式碼', 2)
    code1 = """# ==================== CELL 1: 安裝套件 ====================
print('📦 開始安裝必要套件...')
!pip install -q transformers datasets accelerate bitsandbytes peft openpyxl scikit-learn matplotlib seaborn
print('✅ 套件安裝完成！')"""

    add_highlighted_code(doc, code1)

    doc.add_heading('1.2 程式碼講解', 2)
    add_text(doc, '這個 Cell 安裝所有訓練所需的 Python 套件。')
    add_text(doc, '')
    add_text(doc, '【!pip install 命令】')
    add_text(doc, '• ! 符號：在 Jupyter/Colab 中執行 Shell 命令', bullet=True)
    add_text(doc, '• pip：Python 的套件管理工具', bullet=True)
    add_text(doc, '• install：安裝指令', bullet=True)
    add_text(doc, '• -q：安靜模式（quiet），減少輸出訊息', bullet=True)

    doc.add_heading('1.3 套件說明', 2)
    add_text(doc, '1. transformers：HuggingFace 的 Transformer 模型庫')
    add_text(doc, '2. datasets：資料集處理工具')
    add_text(doc, '3. accelerate：加速訓練的工具')
    add_text(doc, '4. bitsandbytes：模型量化工具（節省記憶體）')
    add_text(doc, '5. peft：參數高效微調工具（LoRA）')
    add_text(doc, '6. openpyxl：Excel 檔案讀取工具')
    add_text(doc, '7. scikit-learn：機器學習評估工具')
    add_text(doc, '8. matplotlib：資料視覺化工具')
    add_text(doc, '9. seaborn：進階視覺化工具')

    doc.add_heading('1.4 重要概念', 2)
    add_text(doc, '【量化（Quantization）】')
    add_text(doc, '將模型參數從 32 位元壓縮到 4 位元，可以節省 8 倍的記憶體。這讓大型模型（如 4B 參數的 MedGemma）能在有限的 GPU 記憶體上訓練。')
    add_text(doc, '')
    add_text(doc, '【比喻】')
    add_text(doc, '安裝套件就像購買工具：如果要修車，你不會自己打造扳手和螺絲起子，而是直接買現成的工具。Python 套件就是程式設計的「工具」。')

    doc.add_page_break()

    # ========== Cell 2 ==========
    add_heading_custom(doc, 'Cell 2: 導入函式庫', 1)

    doc.add_heading('2.1 完整程式碼', 2)
    code2 = """# ==================== CELL 2: 導入函式庫 ====================
import pandas as pd
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, TrainingArguments, Trainer, BitsAndBytesConfig
from datasets import Dataset
from peft import LoraConfig, get_peft_model, prepare_model_for_kbit_training
from sklearn.metrics import f1_score, precision_score, recall_score, accuracy_score, confusion_matrix
from sklearn.model_selection import KFold
from sklearn.utils import resample
from google.colab import files, drive
import os
import json
import numpy as np
import gc
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
import warnings
warnings.filterwarnings('ignore')
print('✅ 所有函式庫導入完成！')"""

    add_highlighted_code(doc, code2)

    doc.add_heading('2.2 主要函式庫分類', 2)
    add_text(doc, '• 資料處理：pandas (pd), numpy (np)', bullet=True)
    add_text(doc, '• 深度學習框架：torch (PyTorch)', bullet=True)
    add_text(doc, '• 模型相關：transformers, peft, datasets', bullet=True)
    add_text(doc, '• 評估指標：sklearn.metrics', bullet=True)
    add_text(doc, '• 交叉驗證：sklearn.model_selection', bullet=True)
    add_text(doc, '• 視覺化：matplotlib.pyplot (plt), seaborn (sns)', bullet=True)

    doc.add_heading('2.3 重要概念', 2)
    add_text(doc, '【模組別名（Alias）】')
    add_text(doc, '使用 as 給模組取簡短的別名，讓程式碼更簡潔。例如：pd.DataFrame() 比 pandas.DataFrame() 更短，也是業界慣例。')

    doc.add_page_break()

    # ========== Cell 3 ==========
    add_heading_custom(doc, 'Cell 3: GPU 檢查', 1)

    doc.add_heading('3.1 完整程式碼', 2)
    code3 = """# ==================== CELL 3: GPU 檢查 ====================
print('🔍 檢查 GPU 狀態...')
if not torch.cuda.is_available():
    raise RuntimeError('❌ 錯誤：需要 GPU 才能執行此程式！')

gpu_name = torch.cuda.get_device_name(0)
gpu_memory = torch.cuda.get_device_properties(0).total_memory / 1024**3

print(f'✅ GPU 已就緒！')
print(f'  GPU 型號: {gpu_name}')
print(f'  GPU 記憶體: {gpu_memory:.2f} GB')

# 清空 GPU 快取
torch.cuda.empty_cache()
gc.collect()
print('✅ GPU 快取已清空')"""

    add_highlighted_code(doc, code3)

    doc.add_heading('3.2 程式碼講解', 2)
    add_text(doc, '【torch.cuda.is_available()】')
    add_text(doc, '檢查 CUDA（NVIDIA GPU 運算平台）是否可用。如果沒有 GPU 或 GPU 未啟用，返回 False。')
    add_text(doc, '')
    add_text(doc, '【if not 條件】')
    add_text(doc, '• not：邏輯反轉', bullet=True)
    add_text(doc, '• 如果 GPU 不可用（False），則執行 if 區塊', bullet=True)
    add_text(doc, '')
    add_text(doc, '【f-string 格式化】')
    add_text(doc, "f'{變數:.2f}' 表示顯示 2 位小數。例如：f'{3.14159:.2f}' 輸出 '3.14'")

    doc.add_heading('3.3 重要概念', 2)
    add_text(doc, '【GPU vs CPU】')
    add_text(doc, 'GPU（圖形處理器）專門處理平行運算，訓練深度學習模型的速度比 CPU 快 10-100 倍。')

    doc.add_page_break()

    # ========== Cell 9: 資料平衡 ==========
    add_heading_custom(doc, 'Cell 9: 資料平衡', 1)

    doc.add_heading('9.1 為什麼需要資料平衡？', 2)
    add_text(doc, '如果訓練資料中「正確術語」有 900 筆，「錯誤術語」只有 100 筆（比例 9:1），模型會傾向預測「都是正確的」，因為這樣準確率有 90%！')
    add_text(doc, '')
    add_text(doc, '但實際上模型並沒有學會辨識錯誤，這就是「資料不平衡」問題。')

    doc.add_heading('9.2 完整程式碼', 2)
    code9 = """# ==================== CELL 9: 資料平衡 ====================
TARGET_ERROR_RATIO = 0.40  # 目標錯誤比例 40%
MAX_TOTAL_SAMPLES = 3500   # 資料上限

current_error_count = len(df[df['Correct_Output'] != 'No issues found.'])
current_correct_count = len(df[df['Correct_Output'] == 'No issues found.'])
current_ratio = current_error_count / len(df)

if current_ratio < TARGET_ERROR_RATIO:
    df_correct = df[df['Correct_Output'] == 'No issues found.'].copy()
    df_error = df[df['Correct_Output'] != 'No issues found.'].copy()

    # 上採樣錯誤資料
    needed_error_count = int(TARGET_ERROR_RATIO * len(df_correct) / (1 - TARGET_ERROR_RATIO))
    df_error_upsampled = resample(df_error, replace=True, n_samples=needed_error_count, random_state=42)

    # 合併並打亂
    df = pd.concat([df_correct, df_error_upsampled], ignore_index=True)
    df = df.sample(frac=1, random_state=42).reset_index(drop=True)

    print('🎉 資料平衡完成！')"""

    add_highlighted_code(doc, code9)

    doc.add_heading('9.3 重點概念', 2)
    add_text(doc, '【上採樣（Oversampling）】')
    add_text(doc, '• replace=True：重複抽樣，同一筆資料可以被抽到多次', bullet=True)
    add_text(doc, '• 從少量資料中「複製」出更多樣本', bullet=True)
    add_text(doc, '')
    add_text(doc, '【下採樣（Undersampling）】')
    add_text(doc, '• replace=False：不重複抽樣，每筆資料最多只會被抽到一次', bullet=True)
    add_text(doc, '• 從大量資料中隨機選取部分樣本', bullet=True)
    add_text(doc, '')
    add_text(doc, '【random_state=42】')
    add_text(doc, '固定隨機種子，讓「隨機」變得可重現。每次執行結果都一樣。')

    doc.add_heading('9.4 比喻', 2)
    add_text(doc, '就像準備考試題庫，你希望「簡單題：難題 = 6:4」，這樣才能平衡訓練。如果 99% 都是簡單題，考試時遇到難題就不會做了。')

    doc.add_page_break()

    # ========== Cell 10: 訓練參數 ==========
    add_heading_custom(doc, 'Cell 10: 設定訓練參數', 1)

    doc.add_heading('10.1 完整程式碼', 2)
    code10 = """# ==================== CELL 10: 設定訓練參數 ====================
BASE_MODEL_ID = 'google/medgemma-4b-it'
LEARNING_RATE = 1e-4      # 0.0001
NUM_EPOCHS = 3
BATCH_SIZE = 2
MAX_LENGTH = 384
GRADIENT_ACCUMULATION_STEPS = 8
N_SPLITS = 5

print(f'⚙️ 訓練參數設定')
print(f'  模型: {BASE_MODEL_ID}')
print(f'  有效批次大小: {BATCH_SIZE * GRADIENT_ACCUMULATION_STEPS}')  # 2 × 8 = 16"""

    add_highlighted_code(doc, code10)

    doc.add_heading('10.2 參數說明', 2)
    add_text(doc, '【BASE_MODEL_ID】預訓練模型名稱（Google 的醫療專用模型）')
    add_text(doc, '【LEARNING_RATE】學習率 = 0.0001，控制每次更新的步伐大小')
    add_text(doc, '【NUM_EPOCHS】訓練輪數 = 3，所有資料看 3 遍')
    add_text(doc, '【BATCH_SIZE】批次大小 = 2，一次處理 2 筆資料')
    add_text(doc, '【GRADIENT_ACCUMULATION_STEPS】梯度累積 = 8 次')
    add_text(doc, '【N_SPLITS】交叉驗證折數 = 5')

    doc.add_heading('10.3 重要概念', 2)
    add_text(doc, '【學習率】')
    add_text(doc, '• 太大：可能跨過最佳點，在山谷兩邊跳來跳去', bullet=True)
    add_text(doc, '• 太小：訓練太慢，可能永遠到不了谷底', bullet=True)
    add_text(doc, '• 0.0001 適合微調預訓練模型', bullet=True)
    add_text(doc, '')
    add_text(doc, '【梯度累積】')
    add_text(doc, '• 實際 batch size = 2（GPU 記憶體限制）', bullet=True)
    add_text(doc, '• 累積 8 次後才更新 = 有效 batch size = 16', bullet=True)
    add_text(doc, '• 比喻：一次只能搬 2 塊磚，但搬 8 趟再一起搬上樓', bullet=True)

    doc.add_page_break()

    # ========== Cell 12: 訓練流程 ==========
    add_heading_custom(doc, 'Cell 12: Fold 1 訓練（核心流程）', 1)

    doc.add_heading('12.1 訓練流程概覽', 2)
    add_text(doc, '1. 載入模型（4-bit 量化）')
    add_text(doc, '2. 配置 LoRA（只訓練 0.21% 參數）')
    add_text(doc, '3. Tokenize 資料（文字轉數字）')
    add_text(doc, '4. 設定訓練參數')
    add_text(doc, '5. 訓練 3 epochs')
    add_text(doc, '6. 評估測試集')
    add_text(doc, '7. 計算指標')
    add_text(doc, '8. 清理記憶體')

    doc.add_heading('12.2 模型量化配置', 2)
    code12_1 = """# 4-bit 量化配置
bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,                      # 使用 4-bit 量化
    bnb_4bit_quant_type='nf4',              # NF4 格式
    bnb_4bit_compute_dtype=torch.float16,   # 計算用 fp16
    bnb_4bit_use_double_quant=True          # 雙重量化
)"""

    add_highlighted_code(doc, code12_1)

    add_text(doc, '【量化效果】')
    add_text(doc, '• 原始：4B 參數 × 32 bits = 16 GB', bullet=True)
    add_text(doc, '• 量化：4B 參數 × 4 bits = 2 GB', bullet=True)
    add_text(doc, '• 節省 8 倍記憶體！', bullet=True)

    doc.add_heading('12.3 LoRA 配置', 2)
    code12_2 = """# LoRA 配置
lora_config = LoraConfig(
    r=16,                    # 秩（rank）
    lora_alpha=32,           # 縮放係數
    lora_dropout=0.1,        # Dropout 比例
    target_modules=['q_proj', 'k_proj', 'v_proj', 'o_proj'],
    task_type='CAUSAL_LM',
    bias='none'
)

model = get_peft_model(model, lora_config)
# 結果：只訓練 0.21% 的參數（840 萬 / 40 億）"""

    add_highlighted_code(doc, code12_2)

    add_text(doc, '【LoRA 比喻】')
    add_text(doc, '不用重做整套西裝（訓練 40 億參數），只調整袖子和腰圍（訓練 840 萬參數）就能完美合身。')

    doc.add_heading('12.4 Tokenization', 2)
    code12_3 = """# Tokenize 函數
def tokenize_function(examples):
    texts = examples['text']
    tokenized = tokenizer(texts, truncation=True, padding='max_length', max_length=384)
    tokenized['labels'] = tokenized['input_ids'].copy()
    return tokenized

# 應用到資料集
train_dataset = Dataset.from_dict(train_data).map(tokenize_function, batched=True)"""

    add_highlighted_code(doc, code12_3)

    add_text(doc, '【Tokenizer】')
    add_text(doc, '將文字轉成數字，電腦才能處理。例如："diabetes" → [1234, 5678]')

    doc.add_page_break()

    # ========== Cell 17: 統計分析 ==========
    add_heading_custom(doc, 'Cell 17: 統計分析', 1)

    doc.add_heading('17.1 計算平均值和標準差', 2)
    code17 = """# 5 個 fold 的 F1 score
values = [0.85, 0.87, 0.83, 0.86, 0.84]

# 計算統計量
mean = np.mean(values)              # 平均值: 0.85
std = np.std(values, ddof=1)        # 標準差: 0.0158

# 計算 95% 信賴區間
confidence_level = 0.95
degrees_freedom = len(values) - 1   # 自由度 = 4
confidence_interval = stats.t.interval(
    confidence_level,
    degrees_freedom,
    mean,
    stats.sem(values)
)
# 結果: [0.8304, 0.8696]"""

    add_highlighted_code(doc, code17)

    doc.add_heading('17.2 重要概念', 2)
    add_text(doc, '【平均值（Mean）】')
    add_text(doc, '中心趨勢，5 次測試的平均表現。計算方式：sum(values) / len(values)')
    add_text(doc, '')
    add_text(doc, '【標準差（Standard Deviation）】')
    add_text(doc, '離散程度，數值越小表示結果越穩定。公式：sqrt(Σ(x - mean)² / (n - 1))')
    add_text(doc, '')
    add_text(doc, '【信賴區間（Confidence Interval）】')
    add_text(doc, '真實平均值有 95% 機率落在此範圍。如果重複實驗 100 次，約有 95 次的結果會落在這個區間內。')
    add_text(doc, '')
    add_text(doc, '【ddof=1】')
    add_text(doc, '自由度修正。因為我們只有 5 個樣本（不是全部資料），所以用樣本標準差（ddof=1）而不是母體標準差（ddof=0）。')

    doc.add_heading('17.3 比喻', 2)
    add_text(doc, '射箭 5 次，平均 85 分。信賴區間告訴你：如果再射 1000 箭，95% 的機會平均分數在 83-87 分之間。')

    doc.add_page_break()

    # ========== Cell 18: 視覺化 ==========
    add_heading_custom(doc, 'Cell 18: 混淆矩陣視覺化', 1)

    doc.add_heading('18.1 什麼是混淆矩陣？', 2)
    add_text(doc, '混淆矩陣顯示模型預測的對錯情況：')
    add_text(doc, '')
    add_text(doc, '                  預測')
    add_text(doc, '              正確    錯誤')
    add_text(doc, '真實  正確     50      10    ← TN=50, FP=10')
    add_text(doc, '      錯誤      5      35    ← FN=5,  TP=35')
    add_text(doc, '')
    add_text(doc, '• TN (True Negative)：正確預測為正確 = 50', bullet=True)
    add_text(doc, '• TP (True Positive)：正確預測為錯誤 = 35', bullet=True)
    add_text(doc, '• FP (False Positive)：誤報（正確說成錯誤）= 5', bullet=True)
    add_text(doc, '• FN (False Negative)：漏報（錯誤說成正確）= 10', bullet=True)

    doc.add_heading('18.2 熱圖繪製', 2)
    code18 = """import seaborn as sns
import numpy as np

# 計算平均混淆矩陣
avg_cm = np.mean(fold_results['confusion_matrices'], axis=0)

# 繪製熱圖
sns.heatmap(
    avg_cm,
    annot=True,                          # 顯示數字
    fmt='.1f',                           # 格式化為 1 位小數
    cmap='Greens',                       # 綠色漸層
    xticklabels=['Correct', 'Error'],    # X 軸標籤
    yticklabels=['Correct', 'Error']     # Y 軸標籤
)"""

    add_highlighted_code(doc, code18)

    add_text(doc, '【熱圖（Heatmap）】')
    add_text(doc, '用顏色深淺表示數值大小，數字越大顏色越深。')
    add_text(doc, '')
    add_text(doc, '【annot=True】在格子裡顯示數字')
    add_text(doc, "【fmt='.1f'】格式化為 1 位小數（50.0）")
    add_text(doc, "【cmap='Greens'】使用綠色系漸層")

    doc.add_page_break()

    # ========== Cell 19: 折線圖 ==========
    add_heading_custom(doc, 'Cell 19: 性能指標視覺化', 1)

    doc.add_heading('19.1 折線圖繪製', 2)
    code19 = """import matplotlib.pyplot as plt

folds = [1, 2, 3, 4, 5]
values = [0.85, 0.87, 0.83, 0.86, 0.84]
mean = 0.85
ci_lower, ci_upper = 0.83, 0.87

# 繪製折線圖
ax.plot(folds, values, marker='o', linewidth=2, color='steelblue', label='Fold Results')

# 平均線（水平線）
ax.axhline(y=mean, color='red', linestyle='--', linewidth=2, label=f'Mean: {mean:.4f}')

# 信賴區間（陰影區域）
ax.axhspan(ci_lower, ci_upper, alpha=0.2, color='red', label='95% CI')

# 設定
ax.set_xlabel('Fold')
ax.set_ylabel('F1 Score')
ax.legend()"""

    add_highlighted_code(doc, code19)

    add_text(doc, '【plot】繪製折線圖')
    add_text(doc, '• marker="o"：在資料點加圓圈', bullet=True)
    add_text(doc, '• linewidth=2：線條寬度', bullet=True)
    add_text(doc, '• color="steelblue"：鋼藍色', bullet=True)
    add_text(doc, '')
    add_text(doc, '【axhline】繪製水平線（平均線）')
    add_text(doc, "• linestyle='--'：虛線", bullet=True)
    add_text(doc, '')
    add_text(doc, '【axhspan】繪製水平範圍（信賴區間陰影）')
    add_text(doc, '• alpha=0.2：透明度 20%，讓陰影不會擋住資料', bullet=True)

    doc.add_page_break()

    # ========== Cell 20: 報告 ==========
    add_heading_custom(doc, 'Cell 20: 生成詳細報告', 1)

    doc.add_heading('20.1 報告生成流程', 2)
    code20 = """# 建立報告行
report_lines = []
report_lines.append('=' * 80)
report_lines.append('MedGemma 訓練報告')
report_lines.append('=' * 80)
report_lines.append(f'平均 F1 Score: {mean:.4f} ± {std:.4f}')
report_lines.append(f'95% 信賴區間: [{ci_lower:.4f}, {ci_upper:.4f}]')

# 組合成完整文字
report_text = '\\n'.join(report_lines)

# 儲存到檔案
with open('cv_report.txt', 'w', encoding='utf-8') as f:
    f.write(report_text)

print('✅ 報告已保存')"""

    add_highlighted_code(doc, code20)

    doc.add_heading('20.2 重要概念', 2)
    add_text(doc, "【'=' * 80】")
    add_text(doc, "字串重複 80 次，產生分隔線：'===============...==============='")
    add_text(doc, '')
    add_text(doc, "【'\\n'.join()】")
    add_text(doc, '用換行符號連接 list 中的字串，組合成完整文字。')
    add_text(doc, '')
    add_text(doc, "【open('w')】")
    add_text(doc, '寫入模式，會覆蓋原有內容。如果要附加內容，用 "a" 模式。')
    add_text(doc, '')
    add_text(doc, "【encoding='utf-8'】")
    add_text(doc, '使用 UTF-8 編碼，支援中文、日文等多國語言。')

    doc.add_page_break()

    # ========== 總結 ==========
    add_heading_custom(doc, '完整訓練流程總結', 1)

    doc.add_heading('20 個 Cells 功能概覽', 2)
    add_text(doc, 'Cell 1-2：環境設定（安裝、導入）')
    add_text(doc, 'Cell 3-5：基礎準備（GPU、Drive、檔案）')
    add_text(doc, 'Cell 6-8：資料讀取與整合')
    add_text(doc, 'Cell 9：資料平衡（解決不平衡問題）⭐')
    add_text(doc, 'Cell 10-11：訓練準備（參數、格式轉換）')
    add_text(doc, 'Cell 12-16：模型訓練（5-fold 交叉驗證）⭐')
    add_text(doc, 'Cell 17：統計分析（平均、標準差、信賴區間）⭐')
    add_text(doc, 'Cell 18-19：視覺化（混淆矩陣、折線圖）')
    add_text(doc, 'Cell 20：生成報告')

    doc.add_heading('核心技術', 2)
    add_text(doc, '【量化】32-bit → 4-bit，節省 8 倍記憶體')
    add_text(doc, '【LoRA】只訓練 0.21% 參數，速度快 400 倍')
    add_text(doc, '【交叉驗證】5-fold，結果更可靠')
    add_text(doc, '【資料平衡】解決不平衡問題，提升模型表現')

    doc.add_heading('你學到了什麼？', 2)
    add_text(doc, '1. Python 基礎：變數、函數、迴圈、條件判斷')
    add_text(doc, '2. 資料處理：Pandas、Excel、資料平衡')
    add_text(doc, '3. 機器學習：訓練/測試集、交叉驗證、評估指標')
    add_text(doc, '4. 深度學習：量化、LoRA、訓練過程')
    add_text(doc, '5. 統計分析：平均值、標準差、信賴區間')
    add_text(doc, '6. 資料視覺化：熱圖、折線圖、Matplotlib/Seaborn')
    add_text(doc, '7. 檔案操作：讀取、寫入、路徑處理')

    add_text(doc, '')
    add_text(doc, '')
    p = doc.add_paragraph('恭喜你完成了完整的 MedGemma 訓練教學！🎉')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)

    return doc

# ========== 主程式 ==========
if __name__ == '__main__':
    print("=" * 80)
    print("正在生成帶有語法高亮的訓練文件...")
    print("=" * 80)
    print()

    doc = create_document()

    # 儲存 Word 文件
    output_path = '/home/user/my-colab-notebooks/MedGemma_訓練教學_語法高亮版.docx'
    doc.save(output_path)

    file_size = os.path.getsize(output_path) / 1024
    print(f"✅ Word 文件已生成: {output_path}")
    print(f"📄 檔案大小: {file_size:.2f} KB")
    print()
    print("=" * 80)
    print("文件特色：")
    print("• 程式碼有語法高亮（不同顏色）")
    print("• 關鍵字：紫色")
    print("• 字串：橘色")
    print("• 註解：綠色")
    print("• 函數名：淺黃色")
    print("• 數字：淺綠色")
    print("=" * 80)
