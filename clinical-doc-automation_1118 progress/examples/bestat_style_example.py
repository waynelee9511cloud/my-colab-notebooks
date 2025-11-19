"""
Bestat樣式分析器使用範例

展示所有主要功能：
1. 從範本提取樣式
2. 套用Bestat樣式到新文件
3. 驗證文件規範
4. 比較文件樣式
5. 生成樣式報告
6. 與Word Formatter整合

作者：Clinical Document Automation Team
版本：1.0.0
"""

import os
import sys
from pathlib import Path

# 添加modules路徑
sys.path.insert(0, str(Path(__file__).parent.parent / "modules"))

from docx import Document
from bestat_style_analyzer import (
    BestatStyleAnalyzer,
    analyze_document_style,
    apply_bestat_style_to_document,
    validate_document_compliance
)
from word_formatter import WordFormatter


def example_01_basic_usage():
    """範例1：基本使用"""
    print("\n" + "=" * 70)
    print("範例1：基本使用 - 建立符合Bestat規範的文件")
    print("=" * 70)

    # 建立輸出目錄
    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 初始化分析器
    analyzer = BestatStyleAnalyzer()

    # 建立新文件
    doc = Document()

    # 添加一些內容
    doc.add_paragraph("這是使用Bestat樣式的測試文件")
    doc.add_paragraph("第一節：簡介")
    doc.add_paragraph("本文件展示了Bestat公司的標準文件樣式。")

    # 套用Bestat樣式
    styled_doc = analyzer.apply_bestat_style(
        doc,
        document_title="臨床試驗Protocol",
        protocol_number="PRO-2025-001",
        version="1.0"
    )

    # 儲存文件
    output_path = output_dir / "example_01_basic.docx"
    styled_doc.save(str(output_path))

    print(f"✓ 文件已建立: {output_path}")
    print("✓ Bestat樣式已套用")


def example_02_extract_from_template():
    """範例2：從範本提取樣式"""
    print("\n" + "=" * 70)
    print("範例2：從現有文件提取樣式配置")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 先建立一個範本文件
    template_path = output_dir / "my_template.docx"
    doc = Document()
    doc.add_paragraph("這是自訂範本")

    # 設定一些樣式
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = \
        section.page_width.__class__(914400)  # 1 inch

    doc.save(str(template_path))
    print(f"✓ 範本已建立: {template_path}")

    # 提取樣式
    analyzer = BestatStyleAnalyzer()
    styles = analyzer.extract_styles_from_document(str(template_path))

    # 儲存為JSON配置
    config_path = output_dir / "extracted_style.json"
    analyzer.save_style_config(str(config_path), styles)

    print(f"✓ 樣式已提取並儲存至: {config_path}")
    print(f"✓ 提取了以下類別:")
    for category in styles.keys():
        if category != 'metadata':
            print(f"  - {category}")


def example_03_apply_with_custom_config():
    """範例3：使用自訂配置套用樣式"""
    print("\n" + "=" * 70)
    print("範例3：使用自訂Bestat配置")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 自訂配置
    custom_config = {
        "company_info": {
            "name": "Bestat Taiwan Branch",
            "full_name": "Bestat Clinical Research - Taiwan"
        },
        "fonts": {
            "title": {
                "name": "Arial",
                "size": 20,  # 更大的標題
                "color": {"r": 0, "g": 102, "b": 204}  # 亮藍色
            }
        },
        "colors": {
            "primary": {"r": 0, "g": 102, "b": 204}  # 使用亮藍色作為主色
        }
    }

    # 初始化分析器並載入自訂配置
    analyzer = BestatStyleAnalyzer(config=custom_config)

    # 建立文件
    doc = Document()
    doc.add_paragraph("使用自訂Bestat配置的文件")
    doc.add_paragraph("本文件使用了Taiwan分公司的樣式設定")

    # 套用樣式
    styled_doc = analyzer.apply_bestat_style(
        doc,
        document_title="台灣臨床試驗Protocol",
        protocol_number="TW-PRO-2025-001",
        version="1.0-TW"
    )

    # 儲存
    output_path = output_dir / "example_03_custom_config.docx"
    styled_doc.save(str(output_path))

    print(f"✓ 自訂配置已套用")
    print(f"✓ 文件已儲存: {output_path}")
    print(f"✓ 公司名稱: {analyzer.style_config['company_info']['name']}")


def example_04_validate_compliance():
    """範例4：驗證文件規範"""
    print("\n" + "=" * 70)
    print("範例4：驗證文件是否符合Bestat規範")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 建立兩個文件：一個符合規範，一個不符合

    # 1. 符合規範的文件
    analyzer = BestatStyleAnalyzer()
    doc1 = Document()
    doc1.add_paragraph("符合規範的文件")
    styled_doc1 = analyzer.apply_bestat_style(
        doc1,
        document_title="規範文件",
        protocol_number="PRO-2025-COMPLIANT",
        version="1.0"
    )
    compliant_path = output_dir / "compliant_doc.docx"
    styled_doc1.save(str(compliant_path))

    # 2. 不符合規範的文件
    doc2 = Document()
    doc2.add_paragraph("不符合規範的文件")
    # 故意使用錯誤的頁面設定
    section = doc2.sections[0]
    section.page_width = section.page_width.__class__(7772400)  # 錯誤的寬度
    non_compliant_path = output_dir / "non_compliant_doc.docx"
    doc2.save(str(non_compliant_path))

    # 驗證兩個文件
    print("\n驗證符合規範的文件:")
    validation1 = validate_document_compliance(str(compliant_path))
    print(f"  符合規範: {validation1['compliant']}")
    print(f"  問題數: {validation1['total_issues']}")
    print(f"  警告數: {validation1['total_warnings']}")

    print("\n驗證不符合規範的文件:")
    validation2 = validate_document_compliance(str(non_compliant_path))
    print(f"  符合規範: {validation2['compliant']}")
    print(f"  問題數: {validation2['total_issues']}")
    print(f"  警告數: {validation2['total_warnings']}")

    if validation2['issues']:
        print("\n  發現的問題:")
        for issue in validation2['issues']:
            print(f"    - {issue}")


def example_05_compare_documents():
    """範例5：比較兩個文件的樣式"""
    print("\n" + "=" * 70)
    print("範例5：比較兩個文件的樣式差異")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 建立兩個樣式不同的文件
    analyzer = BestatStyleAnalyzer()

    # 文件1：使用預設Bestat樣式
    doc1 = Document()
    doc1.add_paragraph("文件1：預設樣式")
    styled_doc1 = analyzer.apply_bestat_style(
        doc1,
        document_title="文件A",
        protocol_number="PRO-2025-A",
        version="1.0"
    )
    doc1_path = output_dir / "compare_doc1.docx"
    styled_doc1.save(str(doc1_path))

    # 文件2：使用自訂配置
    custom_config = {
        "page_setup": {
            "margins": {
                "top_inches": 1.5,  # 不同的邊距
                "bottom_inches": 1.5
            }
        }
    }
    analyzer2 = BestatStyleAnalyzer(config=custom_config)
    doc2 = Document()
    doc2.add_paragraph("文件2：自訂樣式")
    styled_doc2 = analyzer2.apply_bestat_style(
        doc2,
        document_title="文件B",
        protocol_number="PRO-2025-B",
        version="2.0"
    )
    doc2_path = output_dir / "compare_doc2.docx"
    styled_doc2.save(str(doc2_path))

    # 比較兩個文件
    comparison = analyzer.compare_styles(str(doc1_path), str(doc2_path))

    print(f"\n文件1: {doc1_path.name}")
    print(f"文件2: {doc2_path.name}")
    print(f"\n總差異數: {comparison['total_differences']}")

    if comparison['differences']['page_setup']:
        print("\n頁面設定差異:")
        for diff in comparison['differences']['page_setup']:
            print(f"  - {diff}")

    # 儲存比較報告
    report_path = output_dir / "comparison_report.json"
    analyzer.save_style_config(str(report_path), comparison)
    print(f"\n✓ 比較報告已儲存: {report_path}")


def example_06_generate_report():
    """範例6：生成詳細的樣式分析報告"""
    print("\n" + "=" * 70)
    print("範例6：生成完整的樣式分析報告")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 建立測試文件
    analyzer = BestatStyleAnalyzer()
    doc = Document()
    doc.add_paragraph("樣式分析測試文件")

    # 添加各種內容
    doc.add_paragraph("第一章 簡介")
    doc.add_paragraph("這是一個完整的測試文件，包含多種元素。")

    # 添加表格
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "數值"
    table.rows[0].cells[2].text = "備註"

    # 套用樣式
    styled_doc = analyzer.apply_bestat_style(
        doc,
        document_title="分析報告測試",
        protocol_number="PRO-2025-REPORT",
        version="1.0"
    )

    # 儲存文件
    doc_path = output_dir / "report_test_doc.docx"
    styled_doc.save(str(doc_path))

    # 生成報告
    report_path = output_dir / "detailed_style_report.json"
    analyzer.generate_style_report(str(doc_path), str(report_path))

    print(f"✓ 測試文件已建立: {doc_path}")
    print(f"✓ 詳細報告已生成: {report_path}")
    print("\n報告包含:")
    print("  - 提取的樣式設定")
    print("  - 規範驗證結果")
    print("  - 改進建議")


def example_07_word_formatter_integration():
    """範例7：與Word Formatter整合使用"""
    print("\n" + "=" * 70)
    print("範例7：與Word Formatter整合")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 使用Word Formatter建立文件
    formatter = WordFormatter()
    formatter.create_document()

    # 添加內容
    formatter.apply_title_style("臨床試驗Protocol", level=1)
    formatter.apply_body_style(
        "本Protocol描述了一項評估新藥療效的臨床試驗。",
        alignment='justify'
    )

    formatter.apply_title_style("1. 試驗目的", level=2)
    formatter.apply_body_style(
        "評估新藥在目標患者群體中的安全性和有效性。",
        alignment='justify'
    )

    # 套用Bestat樣式
    formatter.apply_bestat_style(
        document_title="第三期臨床試驗Protocol",
        protocol_number="PRO-2025-PHASE3",
        version="2.0"
    )

    # 驗證規範
    validation = formatter.validate_bestat_compliance()
    print(f"\n規範驗證:")
    print(f"  符合規範: {validation['compliant']}")
    print(f"  問題數: {validation['total_issues']}")
    print(f"  警告數: {validation['total_warnings']}")

    # 儲存文件
    output_path = output_dir / "example_07_integrated.docx"
    formatter.save_document(str(output_path))

    print(f"\n✓ 整合文件已建立: {output_path}")


def example_08_batch_processing():
    """範例8：批次處理多個文件"""
    print("\n" + "=" * 70)
    print("範例8：批次套用Bestat樣式")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples" / "batch"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 建立多個測試文件
    test_documents = [
        ("Protocol_001", "PRO-2025-001", "Phase I Trial"),
        ("Protocol_002", "PRO-2025-002", "Phase II Trial"),
        ("Protocol_003", "PRO-2025-003", "Phase III Trial"),
    ]

    analyzer = BestatStyleAnalyzer()
    processed_files = []

    for filename, protocol, title in test_documents:
        # 建立文件
        doc = Document()
        doc.add_paragraph(f"Protocol: {protocol}")
        doc.add_paragraph(f"Title: {title}")
        doc.add_paragraph("這是自動生成的Protocol文件。")

        # 套用Bestat樣式
        styled_doc = analyzer.apply_bestat_style(
            doc,
            document_title=title,
            protocol_number=protocol,
            version="1.0"
        )

        # 儲存
        output_path = output_dir / f"{filename}.docx"
        styled_doc.save(str(output_path))
        processed_files.append(output_path)

        print(f"  ✓ 已處理: {filename}")

    print(f"\n✓ 批次處理完成，共處理 {len(processed_files)} 個文件")
    print(f"✓ 輸出目錄: {output_dir}")


def example_09_convenience_functions():
    """範例9：使用便利函數"""
    print("\n" + "=" * 70)
    print("範例9：使用便利函數快速操作")
    print("=" * 70)

    output_dir = Path(__file__).parent.parent / "output" / "bestat_examples"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 建立測試文件
    test_doc = output_dir / "quick_test.docx"
    doc = Document()
    doc.add_paragraph("快速測試文件")
    doc.save(str(test_doc))

    # 1. 快速分析文件樣式
    print("\n1. 使用 analyze_document_style():")
    style_json = output_dir / "quick_analysis.json"
    styles = analyze_document_style(str(test_doc), str(style_json))
    print(f"  ✓ 已分析並儲存至: {style_json}")

    # 2. 快速套用Bestat樣式
    print("\n2. 使用 apply_bestat_style_to_document():")
    styled_doc = output_dir / "quick_styled.docx"
    apply_bestat_style_to_document(
        input_doc_path=str(test_doc),
        output_doc_path=str(styled_doc),
        document_title="快速樣式測試",
        protocol_number="PRO-2025-QUICK",
        version="1.0"
    )
    print(f"  ✓ 已套用樣式並儲存至: {styled_doc}")

    # 3. 快速驗證規範
    print("\n3. 使用 validate_document_compliance():")
    validation = validate_document_compliance(str(styled_doc))
    print(f"  ✓ 符合規範: {validation['compliant']}")
    print(f"  ✓ 問題數: {validation['total_issues']}")


def main():
    """執行所有範例"""
    print("=" * 70)
    print("Bestat樣式分析器使用範例")
    print("=" * 70)
    print("\n本範例將展示Bestat樣式分析器的所有主要功能")
    print("所有產生的檔案將儲存在 output/bestat_examples/ 目錄")

    try:
        # 執行所有範例
        example_01_basic_usage()
        example_02_extract_from_template()
        example_03_apply_with_custom_config()
        example_04_validate_compliance()
        example_05_compare_documents()
        example_06_generate_report()
        example_07_word_formatter_integration()
        example_08_batch_processing()
        example_09_convenience_functions()

        print("\n" + "=" * 70)
        print("所有範例執行完成！")
        print("=" * 70)
        print("\n請查看 output/bestat_examples/ 目錄中的產生檔案")

    except Exception as e:
        print(f"\n錯誤: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
