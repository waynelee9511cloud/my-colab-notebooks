"""
Data Management Plan (DMP) Generator Module

This module provides automated generation of Data Management Plans for clinical trials.
It creates comprehensive DMP documents based on Protocol information and follows
ICH GCP and FDA 21 CFR Part 11 requirements.

Author: Clinical Doc Automation Team
Date: 2025-11-18
Version: 1.0.0
"""

from typing import List, Dict, Optional, Any, Tuple
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from enum import Enum
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
except ImportError:
    raise ImportError(
        "python-docx is required. Install it with: pip install python-docx"
    )

# Try to import WordFormatter for consistent formatting
try:
    from .word_formatter import WordFormatter
    HAS_WORD_FORMATTER = True
except ImportError:
    HAS_WORD_FORMATTER = False
    print("Warning: WordFormatter not available. Using basic formatting.")


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class ProtocolInfo:
    """Protocol information for DMP generation"""
    protocol_number: str
    protocol_title: str
    sponsor: str
    indication: str
    phase: str
    study_design: str = ""
    sample_size: str = ""
    study_duration: str = ""
    version: str = "1.0"
    date: str = ""

    def __post_init__(self):
        if not self.date:
            self.date = datetime.now().strftime("%d-%b-%Y")


@dataclass
class DataManagementRole:
    """Data management role and responsibility"""
    role: str
    organization: str
    responsibilities: List[str]
    contact_person: Optional[str] = None
    contact_email: Optional[str] = None


@dataclass
class CRFDomain:
    """CRF domain/form information"""
    domain_name: str
    description: str
    visit_schedule: List[str]
    is_critical: bool = False
    validation_rules: int = 0


@dataclass
class ValidationCheck:
    """Data validation check configuration"""
    check_type: str
    description: str
    severity: str  # Critical, Major, Minor
    implementation: str  # Real-time, Batch, Manual


@dataclass
class Milestone:
    """Project milestone"""
    name: str
    description: str
    planned_date: str
    responsible: str
    status: str = "Planned"  # Planned, In Progress, Completed


@dataclass
class DMPSection:
    """Custom DMP section"""
    section_number: str
    title: str
    content: str
    subsections: List[Dict[str, str]] = field(default_factory=list)


# ============================================================================
# DMP Generator Class
# ============================================================================

class DMPGenerator:
    """
    Data Management Plan Generator

    Generates comprehensive DMP documents based on Protocol information.
    Follows ICH GCP and FDA 21 CFR Part 11 requirements.
    """

    # Standard DMP Sections
    STANDARD_SECTIONS = [
        "1. Introduction",
        "2. Study Overview",
        "3. Data Management Responsibilities",
        "4. Data Flow",
        "5. CRF Design",
        "6. Data Validation",
        "7. Data Quality Control",
        "8. Database Lock",
        "9. Data Security",
        "10. Archive"
    ]

    def __init__(self, protocol_info: ProtocolInfo):
        """
        Initialize DMP Generator

        Args:
            protocol_info: Protocol information for the study
        """
        self.protocol_info = protocol_info
        self.dm_roles: List[DataManagementRole] = []
        self.crf_domains: List[CRFDomain] = []
        self.validation_checks: List[ValidationCheck] = []
        self.milestones: List[Milestone] = []
        self.custom_sections: List[DMPSection] = []
        self.edc_system: str = "Electronic Data Capture (EDC) System"
        self.database_lock_criteria: List[str] = []

        # Initialize with default values
        self._set_default_values()

    def _set_default_values(self):
        """Set default values for DMP components"""

        # Default DM roles
        self.dm_roles = [
            DataManagementRole(
                role="Data Management Lead",
                organization=self.protocol_info.sponsor,
                responsibilities=[
                    "Overall oversight of data management activities",
                    "Review and approval of DMP and related documents",
                    "Database design review and approval",
                    "Quality control oversight"
                ]
            ),
            DataManagementRole(
                role="Clinical Data Manager",
                organization=self.protocol_info.sponsor,
                responsibilities=[
                    "CRF design and database setup",
                    "Data validation specification development",
                    "Query management",
                    "Database quality control",
                    "Database lock execution"
                ]
            ),
            DataManagementRole(
                role="Data Entry Personnel",
                organization="Clinical Site or CRO",
                responsibilities=[
                    "Data entry into EDC system",
                    "Query resolution",
                    "Source data verification support"
                ]
            )
        ]

        # Default database lock criteria
        self.database_lock_criteria = [
            "All CRFs completed and data entered into the database",
            "All data queries resolved or escalated to sponsor",
            "All monitoring visits completed and findings resolved",
            "SDV (Source Data Verification) completed as per monitoring plan",
            "All protocol deviations documented and reviewed",
            "Database quality control checks completed with acceptable error rate",
            "Medical coding completed and reviewed",
            "External data transfer completed and reconciled",
            "Database lock memo prepared and approved by relevant stakeholders"
        ]

        # Default validation checks
        self.validation_checks = [
            ValidationCheck(
                check_type="Required Field Check",
                description="Ensure all mandatory fields are completed",
                severity="Critical",
                implementation="Real-time"
            ),
            ValidationCheck(
                check_type="Range Check",
                description="Verify numeric values are within expected ranges",
                severity="Major",
                implementation="Real-time"
            ),
            ValidationCheck(
                check_type="Date Consistency",
                description="Verify logical date sequences",
                severity="Major",
                implementation="Real-time"
            ),
            ValidationCheck(
                check_type="Cross-form Validation",
                description="Check consistency across related forms",
                severity="Major",
                implementation="Batch"
            )
        ]

    def add_dm_role(self, role: DataManagementRole):
        """Add a data management role"""
        self.dm_roles.append(role)

    def add_crf_domain(self, domain: CRFDomain):
        """Add a CRF domain"""
        self.crf_domains.append(domain)

    def add_validation_check(self, check: ValidationCheck):
        """Add a validation check"""
        self.validation_checks.append(check)

    def add_milestone(self, milestone: Milestone):
        """Add a project milestone"""
        self.milestones.append(milestone)

    def add_custom_section(self, section: DMPSection):
        """Add a custom section to DMP"""
        self.custom_sections.append(section)

    def set_edc_system(self, system_name: str):
        """Set the EDC system name"""
        self.edc_system = system_name

    # ========================================================================
    # Document Generation Methods
    # ========================================================================

    def generate_dmp_document(self, output_path: str, use_word_formatter: bool = True) -> str:
        """
        Generate complete DMP document

        Args:
            output_path: Path to save the generated .docx file
            use_word_formatter: Use WordFormatter for consistent formatting

        Returns:
            Path to the generated document
        """
        # Create document
        if use_word_formatter and HAS_WORD_FORMATTER:
            formatter = WordFormatter()
            doc = formatter.create_document()
            formatter.apply_clinical_trial_template(
                document_title="Data Management Plan",
                protocol_number=self.protocol_info.protocol_number,
                version=self.protocol_info.version,
                sponsor=self.protocol_info.sponsor,
                indication=self.protocol_info.indication
            )
            self.doc = doc
            self.formatter = formatter
        else:
            self.doc = Document()
            self.formatter = None
            self._add_cover_page()

        # Add all sections
        self._add_introduction()
        self._add_study_overview()
        self._add_dm_responsibilities()
        self._add_data_flow()
        self._add_crf_design()
        self._add_data_validation()
        self._add_quality_control()
        self._add_database_lock()
        self._add_data_security()
        self._add_archive()

        # Add custom sections if any
        for custom_section in self.custom_sections:
            self._add_custom_section_to_doc(custom_section)

        # Add appendices
        self._add_appendices()

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))

        print(f"✓ DMP document generated: {output_path}")
        return str(output_path)

    def _add_cover_page(self):
        """Add cover page (used when WordFormatter is not available)"""
        # Title
        title = self.doc.add_heading('Data Management Plan', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Protocol info
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(
            f"{self.protocol_info.protocol_title}\n\n"
            f"Protocol Number: {self.protocol_info.protocol_number}\n"
            f"Sponsor: {self.protocol_info.sponsor}\n"
            f"Version: {self.protocol_info.version}\n\n"
            f"{self.protocol_info.date}"
        )
        info_run.font.size = Pt(12)

        self.doc.add_page_break()

    def _add_heading(self, text: str, level: int = 1):
        """Add formatted heading"""
        if self.formatter:
            self.formatter.apply_title_style(text, level=level+1)
        else:
            self.doc.add_heading(text, level=level)

    def _add_paragraph(self, text: str):
        """Add formatted paragraph"""
        if self.formatter:
            self.formatter.apply_body_style(text)
        else:
            self.doc.add_paragraph(text)

    def _add_introduction(self):
        """Add Section 1: Introduction"""
        self._add_heading("1. Introduction", level=1)

        intro_text = f"""
This Data Management Plan (DMP) outlines the data management strategy and procedures for the clinical study
"{self.protocol_info.protocol_title}" (Protocol {self.protocol_info.protocol_number}).

The purpose of this DMP is to ensure that all clinical data are:
• Collected accurately and completely
• Processed in a timely manner
• Validated according to predefined rules
• Stored securely with appropriate backup
• Maintained with full traceability and audit trail
• Compliant with ICH GCP, FDA 21 CFR Part 11, and applicable regulatory requirements

This document describes the data management responsibilities, processes, systems, and quality control
measures that will be implemented throughout the study lifecycle.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # Regulatory compliance statement
        self._add_heading("1.1 Regulatory Compliance", level=2)
        compliance_text = f"""
This Data Management Plan and all associated data management activities will comply with:

• ICH E6(R2) Good Clinical Practice (GCP)
• FDA 21 CFR Part 11 - Electronic Records; Electronic Signatures
• FDA 21 CFR Part 50 - Protection of Human Subjects
• FDA 21 CFR Part 56 - Institutional Review Boards
• GDPR (General Data Protection Regulation) where applicable
• Local regulatory requirements in participating countries

All data management systems and processes will be validated and maintained according to these standards.
        """.strip()

        self._add_paragraph(compliance_text)
        self.doc.add_paragraph()

    def _add_study_overview(self):
        """Add Section 2: Study Overview"""
        self._add_heading("2. Study Overview", level=1)

        # Study information table
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'

        study_info = [
            ("Protocol Number", self.protocol_info.protocol_number),
            ("Protocol Title", self.protocol_info.protocol_title),
            ("Sponsor", self.protocol_info.sponsor),
            ("Indication", self.protocol_info.indication),
            ("Phase", self.protocol_info.phase),
            ("Study Design", self.protocol_info.study_design or "To be specified"),
            ("Target Sample Size", self.protocol_info.sample_size or "To be specified"),
            ("Estimated Study Duration", self.protocol_info.study_duration or "To be specified")
        ]

        for i, (label, value) in enumerate(study_info):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            # Bold the label
            if table.rows[i].cells[0].paragraphs[0].runs:
                table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(table.rows[i].cells[0], 'D9E2F3')

        self.doc.add_paragraph()

    def _add_dm_responsibilities(self):
        """Add Section 3: Data Management Responsibilities"""
        self._add_heading("3. Data Management Responsibilities", level=1)

        overview_text = """
This section defines the roles and responsibilities of all parties involved in data management activities
for this clinical study. Clear definition of responsibilities ensures accountability and quality throughout
the data lifecycle.
        """.strip()

        self._add_paragraph(overview_text)
        self.doc.add_paragraph()

        # Roles and responsibilities table
        for i, role in enumerate(self.dm_roles, 1):
            self._add_heading(f"3.{i} {role.role}", level=2)

            self._add_paragraph(f"Organization: {role.organization}")

            if role.contact_person:
                self._add_paragraph(f"Contact: {role.contact_person}")
            if role.contact_email:
                self._add_paragraph(f"Email: {role.contact_email}")

            self._add_paragraph("Responsibilities:")

            for resp in role.responsibilities:
                p = self.doc.add_paragraph(resp, style='List Bullet')

            self.doc.add_paragraph()

    def _add_data_flow(self):
        """Add Section 4: Data Flow"""
        self._add_heading("4. Data Flow", level=1)

        intro_text = """
This section describes the flow of data from collection at clinical sites through to final database lock
and archival. Understanding the data flow is critical for ensuring data integrity and traceability.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # Data flow steps
        self._add_heading("4.1 Data Collection and Entry", level=2)

        collection_text = f"""
Data collection will be performed according to the protocol and study-specific procedures:

1. Source Data Collection: Clinical site staff will collect data from subjects and record in source documents
2. Data Entry: Authorized site personnel will enter data into the {self.edc_system}
3. Real-time Validation: The EDC system will perform immediate validation checks upon data entry
4. Edit Checks: Automated edit checks will identify potential errors or inconsistencies
5. Query Generation: System-generated or manual queries will be raised for data clarification
6. Query Resolution: Site staff will respond to queries with appropriate corrections or clarifications
7. Data Review: Data managers will review resolved queries and overall data quality
        """.strip()

        self._add_paragraph(collection_text)
        self.doc.add_paragraph()

        # Data flow diagram (represented as a table)
        self._add_heading("4.2 Data Flow Diagram", level=2)

        self._add_paragraph("The following diagram illustrates the data flow process:")
        self.doc.add_paragraph()

        # Create flow diagram table
        flow_table = self.doc.add_table(rows=11, cols=3)
        flow_table.style = 'Light List Accent 1'

        flow_steps = [
            ("Step", "Activity", "Responsible Party"),
            ("1", "Source Data Collection", "Clinical Site Staff"),
            ("2", "Data Entry into EDC", "Site Data Entry Personnel"),
            ("3", "Real-time Validation", "EDC System (Automated)"),
            ("4", "Edit Check Execution", "EDC System (Automated)"),
            ("5", "Query Generation", "Clinical Data Manager"),
            ("6", "Query Resolution", "Clinical Site"),
            ("7", "Data Review & QC", "Clinical Data Manager"),
            ("8", "Medical Coding", "Medical Coder"),
            ("9", "Database Quality Control", "QC Reviewer"),
            ("10", "Database Lock", "Data Management Lead"),
        ]

        for i, (step, activity, responsible) in enumerate(flow_steps):
            flow_table.rows[i].cells[0].text = step
            flow_table.rows[i].cells[1].text = activity
            flow_table.rows[i].cells[2].text = responsible

            # Header row formatting
            if i == 0:
                for j in range(3):
                    cell = flow_table.rows[i].cells[j]
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                    self._set_cell_background(cell, '4472C4')
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        self.doc.add_paragraph()

        # External data transfer
        self._add_heading("4.3 External Data Transfer", level=2)

        external_text = """
External data (e.g., central laboratory data, ECG data, imaging data) will be transferred to the clinical database:

• Data Transfer Format: Specified transfer specifications will define the format and content
• Data Transfer Frequency: According to the schedule defined in vendor contracts
• Data Reconciliation: All external data will be reconciled against expected subjects and visits
• Data Integration: External data will be integrated into the clinical database with full traceability
• Quality Control: All transferred data will undergo quality control review
        """.strip()

        self._add_paragraph(external_text)
        self.doc.add_paragraph()

    def _add_crf_design(self):
        """Add Section 5: CRF Design"""
        self._add_heading("5. CRF Design", level=1)

        intro_text = f"""
The Case Report Form (CRF) design is based on the protocol requirements and will capture all data
necessary to evaluate the study objectives. The CRF will be implemented in the {self.edc_system}.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # CRF development process
        self._add_heading("5.1 CRF Development Process", level=2)

        process_text = """
The CRF development follows a structured process:

1. Protocol Review: Thorough review of the protocol to identify all data collection requirements
2. Draft CRF Design: Create draft CRF based on protocol and study-specific requirements
3. Stakeholder Review: Circulate draft CRF to clinical team, biostatistics, and medical monitor
4. CRF Finalization: Incorporate feedback and finalize CRF design
5. Database Build: Implement CRF in EDC system
6. User Acceptance Testing (UAT): Conduct comprehensive testing of database and CRFs
7. CRF Approval: Obtain final approval from all stakeholders
        """.strip()

        self._add_paragraph(process_text)
        self.doc.add_paragraph()

        # CRF domains
        if self.crf_domains:
            self._add_heading("5.2 CRF Domains", level=2)

            self._add_paragraph("The following CRF domains have been identified for this study:")
            self.doc.add_paragraph()

            # Create CRF domains table
            domain_table = self.doc.add_table(rows=len(self.crf_domains) + 1, cols=4)
            domain_table.style = 'Light Grid Accent 1'

            # Header
            headers = ["Domain", "Description", "Critical", "Validation Rules"]
            for j, header in enumerate(headers):
                cell = domain_table.rows[0].cells[j]
                cell.text = header
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_background(cell, '4472C4')
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

            # Data rows
            for i, domain in enumerate(self.crf_domains, 1):
                domain_table.rows[i].cells[0].text = domain.domain_name
                domain_table.rows[i].cells[1].text = domain.description
                domain_table.rows[i].cells[2].text = "Yes" if domain.is_critical else "No"
                domain_table.rows[i].cells[3].text = str(domain.validation_rules)

                # Highlight critical domains
                if domain.is_critical:
                    self._set_cell_background(domain_table.rows[i].cells[2], 'FFE699')

            self.doc.add_paragraph()
        else:
            self._add_heading("5.2 CRF Domains", level=2)
            self._add_paragraph("CRF domains will be defined based on protocol requirements. "
                              "Typical domains include:")

            standard_domains = [
                "Demographics",
                "Medical History",
                "Inclusion/Exclusion Criteria",
                "Vital Signs",
                "Physical Examination",
                "Laboratory",
                "ECG",
                "Adverse Events",
                "Concomitant Medications",
                "Efficacy Assessments",
                "Study Drug Administration"
            ]

            for domain in standard_domains:
                self.doc.add_paragraph(domain, style='List Bullet')

            self.doc.add_paragraph()

        # CRF conventions
        self._add_heading("5.3 CRF Conventions", level=2)

        conventions_text = """
The following conventions will be used in CRF design:

• Date Format: DD-MMM-YYYY (e.g., 15-Jan-2025)
• Time Format: 24-hour format (HH:MM)
• Missing Data: Explicit reasons for missing data will be captured
• Units: Standard units will be specified for all measurements
• Coding: Medical coding will use MedDRA for adverse events and WHO Drug for medications
• Derived Variables: Automatically calculated fields will be clearly indicated
• Required Fields: Mandatory fields will be marked with an asterisk (*)
        """.strip()

        self._add_paragraph(conventions_text)
        self.doc.add_paragraph()

    def _add_data_validation(self):
        """Add Section 6: Data Validation"""
        self._add_heading("6. Data Validation", level=1)

        intro_text = """
Data validation is a critical component of data quality assurance. Validation checks are implemented
in the EDC system to identify potential errors and inconsistencies in real-time or through batch processes.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # Validation strategy
        self._add_heading("6.1 Validation Strategy", level=2)

        strategy_text = """
The validation strategy consists of multiple layers:

• Real-time Validation: Immediate checks performed during data entry
• Batch Validation: Periodic checks run across the entire database
• Manual Review: Targeted review by data managers and medical monitors
• Cross-form Validation: Consistency checks across related CRF pages
• Protocol Deviation Checks: Automated detection of potential protocol deviations
        """.strip()

        self._add_paragraph(strategy_text)
        self.doc.add_paragraph()

        # Validation checks table
        self._add_heading("6.2 Validation Checks", level=2)

        self._add_paragraph("The following validation checks will be implemented:")
        self.doc.add_paragraph()

        val_table = self.doc.add_table(rows=len(self.validation_checks) + 1, cols=4)
        val_table.style = 'Light Grid Accent 1'

        # Header
        headers = ["Check Type", "Description", "Severity", "Implementation"]
        for j, header in enumerate(headers):
            cell = val_table.rows[0].cells[j]
            cell.text = header
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(cell, '4472C4')
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Data rows
        for i, check in enumerate(self.validation_checks, 1):
            val_table.rows[i].cells[0].text = check.check_type
            val_table.rows[i].cells[1].text = check.description
            val_table.rows[i].cells[2].text = check.severity
            val_table.rows[i].cells[3].text = check.implementation

            # Color code severity
            severity_cell = val_table.rows[i].cells[2]
            if check.severity == "Critical":
                self._set_cell_background(severity_cell, 'FF6B6B')
            elif check.severity == "Major":
                self._set_cell_background(severity_cell, 'FFA500')
            else:
                self._set_cell_background(severity_cell, 'FFE699')

        self.doc.add_paragraph()

        # Query management
        self._add_heading("6.3 Query Management", level=2)

        query_text = """
Data queries are used to clarify or correct data issues identified through validation checks:

• Query Generation: Queries will be generated automatically by the EDC system or manually by data managers
• Query Assignment: Queries will be assigned to appropriate site personnel
• Query Response: Sites must respond to queries with corrections or explanations
• Query Review: Data managers will review query responses for adequacy
• Query Closure: Queries will be closed only after satisfactory resolution
• Query Metrics: Regular reports on query volume, aging, and resolution rates will be generated
        """.strip()

        self._add_paragraph(query_text)
        self.doc.add_paragraph()

    def _add_quality_control(self):
        """Add Section 7: Data Quality Control"""
        self._add_heading("7. Data Quality Control", level=1)

        intro_text = """
Data quality control (QC) activities ensure that data in the clinical database are accurate, complete,
and consistent with source documents. A comprehensive QC program will be implemented throughout the study.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # QC strategy
        self._add_heading("7.1 QC Strategy", level=2)

        strategy_text = """
The data quality control strategy includes:

• Database Design QC: Review and testing of database build before study initiation
• Ongoing Data Review: Regular review of data listings and query reports
• Data Cleaning: Systematic review and resolution of data inconsistencies
• Medical Coding QC: Quality control of medical and drug coding
• External Data QC: Verification of externally transferred data
• Pre-Lock QC: Comprehensive quality control review before database lock
        """.strip()

        self._add_paragraph(strategy_text)
        self.doc.add_paragraph()

        # QC activities table
        self._add_heading("7.2 QC Activities and Frequency", level=2)

        qc_table = self.doc.add_table(rows=8, cols=3)
        qc_table.style = 'Light Grid Accent 1'

        qc_activities = [
            ("QC Activity", "Frequency", "Responsible"),
            ("Data Entry Audit", "Monthly", "QC Reviewer"),
            ("Query Report Review", "Weekly", "Clinical Data Manager"),
            ("Medical Coding Review", "Ongoing", "Medical Coding Lead"),
            ("Data Listing Review", "Bi-weekly", "Clinical Data Manager"),
            ("External Data Reconciliation", "Per transfer", "Clinical Data Manager"),
            ("Protocol Deviation Review", "Monthly", "Data Management Lead"),
            ("Pre-Lock Database Review", "Before lock", "QC Team")
        ]

        for i, (activity, frequency, responsible) in enumerate(qc_activities):
            qc_table.rows[i].cells[0].text = activity
            qc_table.rows[i].cells[1].text = frequency
            qc_table.rows[i].cells[2].text = responsible

            if i == 0:
                for j in range(3):
                    cell = qc_table.rows[i].cells[j]
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                    self._set_cell_background(cell, '4472C4')
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        self.doc.add_paragraph()

        # QC metrics
        self._add_heading("7.3 Quality Metrics", level=2)

        metrics_text = """
The following quality metrics will be tracked and reported:

• Data Entry Accuracy Rate: Percentage of fields entered correctly without queries
• Query Response Time: Average time from query generation to resolution
• Open Query Rate: Percentage of open queries relative to total data points
• Critical Field Completion Rate: Percentage of critical fields completed
• Protocol Deviation Rate: Number of deviations per enrolled subject
• Database Error Rate: Errors identified per 1000 data points reviewed

Target: Maintain error rate below 0.5% before database lock.
        """.strip()

        self._add_paragraph(metrics_text)
        self.doc.add_paragraph()

    def _add_database_lock(self):
        """Add Section 8: Database Lock"""
        self._add_heading("8. Database Lock", level=1)

        intro_text = """
Database lock is the process of freezing the clinical database after all data have been entered,
validated, and quality control activities completed. Once locked, no further changes to data are permitted
without proper documentation and approval.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # Database lock criteria
        self._add_heading("8.1 Database Lock Criteria", level=2)

        self._add_paragraph("The database will be locked when the following criteria are met:")
        self.doc.add_paragraph()

        for criterion in self.database_lock_criteria:
            self.doc.add_paragraph(criterion, style='List Number')

        self.doc.add_paragraph()

        # Lock process
        self._add_heading("8.2 Database Lock Process", level=2)

        process_table = self.doc.add_table(rows=7, cols=3)
        process_table.style = 'Light Grid Accent 1'

        lock_steps = [
            ("Step", "Activity", "Responsible"),
            ("1", "Pre-lock database quality control review", "QC Team"),
            ("2", "Generate and review pre-lock data listings", "Clinical Data Manager"),
            ("3", "Prepare database lock memo documenting readiness", "Data Management Lead"),
            ("4", "Circulate lock memo for stakeholder review/approval", "Data Management Lead"),
            ("5", "Execute database lock in EDC system", "Clinical Data Manager"),
            ("6", "Generate final locked database extract", "Clinical Data Manager")
        ]

        for i, (step, activity, responsible) in enumerate(lock_steps):
            process_table.rows[i].cells[0].text = step
            process_table.rows[i].cells[1].text = activity
            process_table.rows[i].cells[2].text = responsible

            if i == 0:
                for j in range(3):
                    cell = process_table.rows[i].cells[j]
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                    self._set_cell_background(cell, '4472C4')
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        self.doc.add_paragraph()

        # Post-lock changes
        self._add_heading("8.3 Post-Lock Changes", level=2)

        post_lock_text = """
After database lock, any changes to data must be:

• Documented with clear justification for the change
• Approved by the Data Management Lead and Medical Monitor
• Tracked in a post-lock change log
• Implemented only after database unlock authorization
• Followed by re-lock of the database

Post-lock changes should be minimized through thorough pre-lock quality control.
        """.strip()

        self._add_paragraph(post_lock_text)
        self.doc.add_paragraph()

    def _add_data_security(self):
        """Add Section 9: Data Security"""
        self._add_heading("9. Data Security", level=1)

        intro_text = """
Data security measures ensure the confidentiality, integrity, and availability of clinical trial data
throughout the study lifecycle. All systems and processes comply with FDA 21 CFR Part 11 and GDPR requirements.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # Access control
        self._add_heading("9.1 Access Control", level=2)

        access_text = f"""
Access to the {self.edc_system} and clinical data is strictly controlled:

• User Authentication: All users must authenticate with unique username and password
• Role-Based Access: Access rights assigned based on user role and responsibilities
• Password Policy: Strong password requirements with regular password changes
• Account Deactivation: Immediate deactivation of accounts when users leave the study
• Access Logs: All system access is logged and periodically reviewed
• Training Requirements: All users must complete system training before access is granted
        """.strip()

        self._add_paragraph(access_text)
        self.doc.add_paragraph()

        # Audit trail
        self._add_heading("9.2 Audit Trail", level=2)

        audit_text = f"""
The {self.edc_system} maintains a complete audit trail for all data and system activities:

• All data entries, modifications, and deletions are logged
• Audit trail records user ID, date/time, old value, new value, and reason for change
• Audit trail cannot be modified or deleted by any user
• Audit trail is included in all database extracts and archives
• Regular audit trail reviews are performed to detect anomalies
        """.strip()

        self._add_paragraph(audit_text)
        self.doc.add_paragraph()

        # Data protection
        self._add_heading("9.3 Data Protection and Privacy", level=2)

        protection_text = """
Subject privacy and data protection measures include:

• Subject Identification: Subjects identified by unique subject ID, not by name
• Data De-identification: Personal identifiers removed from datasets for analysis
• Secure Transmission: All data transmissions encrypted using industry-standard protocols (SSL/TLS)
• Data Storage: Data stored in secure, access-controlled servers
• GDPR Compliance: Procedures for subject consent, data access, and right to be forgotten
• Data Retention: Data retained according to regulatory requirements and company policies
        """.strip()

        self._add_paragraph(protection_text)
        self.doc.add_paragraph()

        # Backup and disaster recovery
        self._add_heading("9.4 Backup and Disaster Recovery", level=2)

        backup_text = f"""
Backup and disaster recovery procedures ensure data availability:

• Regular Backups: Database backups performed daily
• Backup Storage: Backups stored in geographically separate secure location
• Backup Testing: Regular restoration tests to verify backup integrity
• Disaster Recovery Plan: Documented procedures for system recovery
• Recovery Time Objective (RTO): System restored within 24 hours of failure
• Business Continuity: Redundant systems to minimize service interruption
        """.strip()

        self._add_paragraph(backup_text)
        self.doc.add_paragraph()

    def _add_archive(self):
        """Add Section 10: Archive"""
        self._add_heading("10. Archive", level=1)

        intro_text = """
All clinical data, documentation, and related materials will be archived according to regulatory
requirements and company policies. Archived materials must remain accessible for regulatory inspections
and audits.
        """.strip()

        self._add_paragraph(intro_text)
        self.doc.add_paragraph()

        # Archival requirements
        self._add_heading("10.1 Archival Requirements", level=2)

        requirements_text = """
The following materials will be archived:

• Clinical Database: Final locked database with complete audit trail
• Database Specifications: Database design documents, validation specifications, edit check specifications
• CRF and eCRF: Blank CRFs, CRF completion guidelines, database setup specifications
• Data Management Plan: This document and any amendments
• Query Reports: All query reports and resolution documentation
• Data Transfer Specifications: External data transfer agreements and specifications
• Quality Control Documentation: QC reports, findings, and resolutions
• Database Lock Documentation: Database lock memo and approvals
• User Access Logs: System access logs for the study duration
• System Validation: EDC system validation documentation
        """.strip()

        self._add_paragraph(requirements_text)
        self.doc.add_paragraph()

        # Retention period
        self._add_heading("10.2 Retention Period", level=2)

        retention_table = self.doc.add_table(rows=4, cols=2)
        retention_table.style = 'Light Grid Accent 1'

        retention_info = [
            ("Document Type", "Retention Period"),
            ("Clinical Database and Documentation", "At least 25 years or per regulatory requirement"),
            ("Essential Documents", "25 years after study completion"),
            ("Audit Trail and Logs", "25 years after study completion")
        ]

        for i, (doc_type, period) in enumerate(retention_info):
            retention_table.rows[i].cells[0].text = doc_type
            retention_table.rows[i].cells[1].text = period

            if i == 0:
                for j in range(2):
                    cell = retention_table.rows[i].cells[j]
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                    self._set_cell_background(cell, '4472C4')
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        self.doc.add_paragraph()

        # Archive format
        self._add_heading("10.3 Archive Format and Storage", level=2)

        format_text = """
Archive format and storage specifications:

• Archive Format: PDF/A for documents, FDA-compliant datasets (e.g., CDISC SDTM/ADaM)
• Storage Media: Secure electronic archive system with appropriate backup
• Archive Location: Sponsor's clinical archive facility or contracted archive service
• Accessibility: Archived materials must be retrievable within 24 hours for regulatory inspection
• Archive Integrity: Regular integrity checks to ensure archived data remains readable
• Archive Index: Comprehensive index of all archived materials for easy retrieval
        """.strip()

        self._add_paragraph(format_text)
        self.doc.add_paragraph()

    def _add_custom_section_to_doc(self, section: DMPSection):
        """Add custom section to document"""
        self._add_heading(f"{section.section_number} {section.title}", level=1)
        self._add_paragraph(section.content)

        # Add subsections if any
        for i, subsection in enumerate(section.subsections, 1):
            self._add_heading(
                f"{section.section_number}.{i} {subsection.get('title', 'Subsection')}",
                level=2
            )
            self._add_paragraph(subsection.get('content', ''))

        self.doc.add_paragraph()

    def _add_appendices(self):
        """Add appendices"""
        self._add_heading("Appendix A: Abbreviations and Definitions", level=1)

        # Abbreviations table
        abbrev_table = self.doc.add_table(rows=16, cols=2)
        abbrev_table.style = 'Light Grid Accent 1'

        abbreviations = [
            ("Abbreviation", "Definition"),
            ("AE", "Adverse Event"),
            ("CDISC", "Clinical Data Interchange Standards Consortium"),
            ("CRF", "Case Report Form"),
            ("CRO", "Contract Research Organization"),
            ("DM", "Data Management"),
            ("DMP", "Data Management Plan"),
            ("DVP", "Data Validation Plan"),
            ("EDC", "Electronic Data Capture"),
            ("GCP", "Good Clinical Practice"),
            ("ICH", "International Council for Harmonisation"),
            ("QC", "Quality Control"),
            ("SAE", "Serious Adverse Event"),
            ("SDV", "Source Data Verification"),
            ("UAT", "User Acceptance Testing")
        ]

        for i, (abbr, definition) in enumerate(abbreviations):
            abbrev_table.rows[i].cells[0].text = abbr
            abbrev_table.rows[i].cells[1].text = definition

            if i == 0:
                for j in range(2):
                    cell = abbrev_table.rows[i].cells[j]
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                    self._set_cell_background(cell, '4472C4')
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        self.doc.add_paragraph()

        # Add timeline if milestones are defined
        if self.milestones:
            self._add_heading("Appendix B: Project Timeline", level=1)

            timeline_table = self.doc.add_table(rows=len(self.milestones) + 1, cols=4)
            timeline_table.style = 'Light Grid Accent 1'

            # Header
            headers = ["Milestone", "Description", "Planned Date", "Responsible"]
            for j, header in enumerate(headers):
                cell = timeline_table.rows[0].cells[j]
                cell.text = header
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_background(cell, '4472C4')
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

            # Data
            for i, milestone in enumerate(self.milestones, 1):
                timeline_table.rows[i].cells[0].text = milestone.name
                timeline_table.rows[i].cells[1].text = milestone.description
                timeline_table.rows[i].cells[2].text = milestone.planned_date
                timeline_table.rows[i].cells[3].text = milestone.responsible

    def _set_cell_background(self, cell, color: str):
        """Set cell background color"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except:
            pass  # Silently fail if background setting doesn't work

    # ========================================================================
    # Export Methods
    # ========================================================================

    def export_to_dict(self) -> Dict[str, Any]:
        """Export DMP configuration to dictionary"""
        return {
            'protocol_info': {
                'protocol_number': self.protocol_info.protocol_number,
                'protocol_title': self.protocol_info.protocol_title,
                'sponsor': self.protocol_info.sponsor,
                'indication': self.protocol_info.indication,
                'phase': self.protocol_info.phase,
                'version': self.protocol_info.version
            },
            'dm_roles': [
                {
                    'role': role.role,
                    'organization': role.organization,
                    'responsibilities': role.responsibilities,
                    'contact_person': role.contact_person,
                    'contact_email': role.contact_email
                }
                for role in self.dm_roles
            ],
            'crf_domains': [
                {
                    'domain_name': domain.domain_name,
                    'description': domain.description,
                    'is_critical': domain.is_critical,
                    'validation_rules': domain.validation_rules
                }
                for domain in self.crf_domains
            ],
            'validation_checks': [
                {
                    'check_type': check.check_type,
                    'description': check.description,
                    'severity': check.severity,
                    'implementation': check.implementation
                }
                for check in self.validation_checks
            ],
            'milestones': [
                {
                    'name': milestone.name,
                    'description': milestone.description,
                    'planned_date': milestone.planned_date,
                    'responsible': milestone.responsible,
                    'status': milestone.status
                }
                for milestone in self.milestones
            ]
        }


# ============================================================================
# Convenience Functions
# ============================================================================

def create_dmp(
    protocol_info: ProtocolInfo,
    output_path: str,
    crf_domains: Optional[List[CRFDomain]] = None,
    milestones: Optional[List[Milestone]] = None,
    edc_system: Optional[str] = None,
    use_word_formatter: bool = True
) -> str:
    """
    Convenience function to create a DMP document

    Args:
        protocol_info: Protocol information
        output_path: Path to save the DMP document
        crf_domains: Optional list of CRF domains
        milestones: Optional list of project milestones
        edc_system: Optional EDC system name
        use_word_formatter: Use WordFormatter for formatting

    Returns:
        Path to the generated document
    """
    generator = DMPGenerator(protocol_info)

    if crf_domains:
        for domain in crf_domains:
            generator.add_crf_domain(domain)

    if milestones:
        for milestone in milestones:
            generator.add_milestone(milestone)

    if edc_system:
        generator.set_edc_system(edc_system)

    return generator.generate_dmp_document(output_path, use_word_formatter=use_word_formatter)


def create_dmp_with_defaults(
    protocol_number: str,
    protocol_title: str,
    sponsor: str,
    indication: str,
    phase: str,
    output_path: str
) -> str:
    """
    Quick DMP creation with minimal parameters

    Args:
        protocol_number: Protocol number
        protocol_title: Protocol title
        sponsor: Sponsor name
        indication: Study indication
        phase: Study phase
        output_path: Path to save the DMP

    Returns:
        Path to the generated document
    """
    protocol_info = ProtocolInfo(
        protocol_number=protocol_number,
        protocol_title=protocol_title,
        sponsor=sponsor,
        indication=indication,
        phase=phase
    )

    return create_dmp(protocol_info, output_path)


# ============================================================================
# Example Usage and Testing
# ============================================================================

def example_usage():
    """
    Example usage of DMP Generator
    """
    print("=" * 80)
    print("Data Management Plan Generator - Example Usage")
    print("=" * 80)

    # 1. Create protocol information
    protocol_info = ProtocolInfo(
        protocol_number="PROTO-2025-001",
        protocol_title="A Phase III, Randomized, Double-Blind, Placebo-Controlled Study",
        sponsor="Example Pharmaceutical Company",
        indication="Type 2 Diabetes Mellitus",
        phase="Phase III",
        study_design="Randomized, Double-Blind, Placebo-Controlled",
        sample_size="300 subjects",
        study_duration="18 months",
        version="1.0"
    )

    # 2. Create DMP generator
    generator = DMPGenerator(protocol_info)

    # 3. Add CRF domains
    crf_domains = [
        CRFDomain(
            domain_name="Demographics",
            description="Subject demographic information",
            visit_schedule=["Screening"],
            is_critical=True,
            validation_rules=5
        ),
        CRFDomain(
            domain_name="Vital Signs",
            description="Blood pressure, heart rate, temperature",
            visit_schedule=["Screening", "Week 4", "Week 8", "Week 12"],
            is_critical=False,
            validation_rules=12
        ),
        CRFDomain(
            domain_name="Laboratory",
            description="Hematology, chemistry, HbA1c",
            visit_schedule=["Screening", "Week 4", "Week 8", "Week 12"],
            is_critical=True,
            validation_rules=25
        ),
        CRFDomain(
            domain_name="Adverse Events",
            description="All adverse events and serious adverse events",
            visit_schedule=["All visits"],
            is_critical=True,
            validation_rules=15
        )
    ]

    for domain in crf_domains:
        generator.add_crf_domain(domain)

    # 4. Add milestones
    milestones = [
        Milestone(
            name="Database Design",
            description="Complete CRF design and database build",
            planned_date="15-Feb-2025",
            responsible="Clinical Data Manager"
        ),
        Milestone(
            name="UAT Completion",
            description="User Acceptance Testing completed",
            planned_date="28-Feb-2025",
            responsible="Clinical Data Manager"
        ),
        Milestone(
            name="First Subject Enrolled",
            description="First subject enrolled in the study",
            planned_date="15-Mar-2025",
            responsible="Clinical Operations"
        ),
        Milestone(
            name="Last Subject Last Visit",
            description="Last subject completes final visit",
            planned_date="15-Sep-2026",
            responsible="Clinical Operations"
        ),
        Milestone(
            name="Database Lock",
            description="Clinical database locked for analysis",
            planned_date="30-Oct-2026",
            responsible="Data Management Lead"
        )
    ]

    for milestone in milestones:
        generator.add_milestone(milestone)

    # 5. Set EDC system
    generator.set_edc_system("Medidata Rave")

    # 6. Add custom validation checks
    generator.add_validation_check(
        ValidationCheck(
            check_type="HbA1c Protocol Deviation",
            description="Check if HbA1c values meet inclusion criteria",
            severity="Critical",
            implementation="Real-time"
        )
    )

    # 7. Add custom section
    custom_section = DMPSection(
        section_number="11",
        title="Study-Specific Considerations",
        content="This study has specific requirements for diabetes medications monitoring.",
        subsections=[
            {
                'title': 'Glucose Monitoring',
                'content': 'Subjects will perform daily glucose monitoring using provided glucometers.'
            },
            {
                'title': 'Insulin Dose Adjustments',
                'content': 'All insulin dose adjustments must be recorded in the eCRF within 24 hours.'
            }
        ]
    )

    generator.add_custom_section(custom_section)

    # 8. Generate DMP document
    output_path = "/tmp/DMP_PROTO-2025-001_v1.0.docx"

    try:
        result_path = generator.generate_dmp_document(output_path)
        print(f"\n✓ DMP document successfully generated: {result_path}")

        # 9. Export to dictionary
        dmp_dict = generator.export_to_dict()
        print(f"\n✓ DMP configuration exported to dictionary")
        print(f"  - DM Roles: {len(dmp_dict['dm_roles'])}")
        print(f"  - CRF Domains: {len(dmp_dict['crf_domains'])}")
        print(f"  - Validation Checks: {len(dmp_dict['validation_checks'])}")
        print(f"  - Milestones: {len(dmp_dict['milestones'])}")

    except Exception as e:
        print(f"\n✗ Error generating DMP: {str(e)}")
        import traceback
        traceback.print_exc()


def quick_demo():
    """Quick demo with minimal code"""
    print("=" * 80)
    print("Quick Demo - DMP Generator")
    print("=" * 80)

    # Create DMP with minimal parameters
    output_path = "/tmp/DMP_Quick_Demo.docx"

    result = create_dmp_with_defaults(
        protocol_number="DEMO-2025-001",
        protocol_title="A Phase II Study of Novel Diabetes Treatment",
        sponsor="Demo Pharma Inc.",
        indication="Type 2 Diabetes",
        phase="Phase II",
        output_path=output_path
    )

    print(f"\n✓ Quick DMP generated: {result}")


def test_dmp_generator():
    """
    Test suite for DMP Generator
    """
    print("=" * 80)
    print("DMP Generator - Test Suite")
    print("=" * 80)

    # Test 1: Basic DMP creation
    print("\nTest 1: Basic DMP Creation")
    try:
        protocol_info = ProtocolInfo(
            protocol_number="TEST-001",
            protocol_title="Test Protocol",
            sponsor="Test Sponsor",
            indication="Test Indication",
            phase="Phase I"
        )

        generator = DMPGenerator(protocol_info)
        assert generator.protocol_info.protocol_number == "TEST-001"
        print("✓ PASS: Basic DMP creation")
    except Exception as e:
        print(f"✗ FAIL: {str(e)}")

    # Test 2: Adding components
    print("\nTest 2: Adding Components")
    try:
        generator.add_crf_domain(
            CRFDomain(
                domain_name="Test Domain",
                description="Test Description",
                visit_schedule=["V1", "V2"]
            )
        )
        assert len(generator.crf_domains) == 1

        generator.add_milestone(
            Milestone(
                name="Test Milestone",
                description="Test",
                planned_date="01-Jan-2025",
                responsible="Tester"
            )
        )
        assert len(generator.milestones) == 1

        print("✓ PASS: Adding components")
    except Exception as e:
        print(f"✗ FAIL: {str(e)}")

    # Test 3: Document generation
    print("\nTest 3: Document Generation")
    try:
        output_path = "/tmp/test_dmp.docx"
        result = generator.generate_dmp_document(output_path, use_word_formatter=False)
        assert os.path.exists(result)
        print(f"✓ PASS: Document generated at {result}")
    except Exception as e:
        print(f"✗ FAIL: {str(e)}")

    # Test 4: Export to dictionary
    print("\nTest 4: Export to Dictionary")
    try:
        dmp_dict = generator.export_to_dict()
        assert 'protocol_info' in dmp_dict
        assert 'dm_roles' in dmp_dict
        assert 'crf_domains' in dmp_dict
        print("✓ PASS: Export to dictionary")
    except Exception as e:
        print(f"✗ FAIL: {str(e)}")

    print("\n" + "=" * 80)
    print("Test Suite Complete")
    print("=" * 80)


# ============================================================================
# Quick Start Guide
# ============================================================================

def quick_start_guide():
    """Display quick start guide"""
    guide = """
    ╔════════════════════════════════════════════════════════════════════════════╗
    ║              Data Management Plan Generator - Quick Start Guide            ║
    ╚════════════════════════════════════════════════════════════════════════════╝

    【Installation】
    pip install python-docx

    【Basic Usage】

    from modules.dmp_generator import DMPGenerator, ProtocolInfo, CRFDomain

    # 1. Create protocol information
    protocol_info = ProtocolInfo(
        protocol_number="PROTO-2025-001",
        protocol_title="My Clinical Study",
        sponsor="My Company",
        indication="Disease X",
        phase="Phase III"
    )

    # 2. Create DMP generator
    generator = DMPGenerator(protocol_info)

    # 3. Add CRF domains (optional)
    generator.add_crf_domain(CRFDomain(
        domain_name="Demographics",
        description="Subject demographics",
        visit_schedule=["Screening"],
        is_critical=True
    ))

    # 4. Generate DMP document
    generator.generate_dmp_document("output/DMP.docx")

    【Quick Create】

    from modules.dmp_generator import create_dmp_with_defaults

    create_dmp_with_defaults(
        protocol_number="PROTO-2025-001",
        protocol_title="My Study",
        sponsor="My Company",
        indication="Disease X",
        phase="Phase III",
        output_path="DMP.docx"
    )

    【Advanced Features】

    # Add milestones
    from modules.dmp_generator import Milestone

    generator.add_milestone(Milestone(
        name="Database Lock",
        description="Lock clinical database",
        planned_date="31-Dec-2025",
        responsible="Data Manager"
    ))

    # Add custom validation checks
    from modules.dmp_generator import ValidationCheck

    generator.add_validation_check(ValidationCheck(
        check_type="Custom Check",
        description="Study-specific validation",
        severity="Major",
        implementation="Real-time"
    ))

    # Add custom sections
    from modules.dmp_generator import DMPSection

    generator.add_custom_section(DMPSection(
        section_number="11",
        title="Study-Specific Requirements",
        content="Custom content here..."
    ))

    # Set EDC system
    generator.set_edc_system("Medidata Rave")

    【Standard DMP Sections】
    1. Introduction
    2. Study Overview
    3. Data Management Responsibilities
    4. Data Flow
    5. CRF Design
    6. Data Validation
    7. Data Quality Control
    8. Database Lock
    9. Data Security
    10. Archive

    【Compliance】
    ✓ ICH GCP E6(R2)
    ✓ FDA 21 CFR Part 11
    ✓ FDA 21 CFR Part 50 & 56
    ✓ GDPR considerations

    【Need Help?】
    Run: python -m modules.dmp_generator

    ╔════════════════════════════════════════════════════════════════════════════╗
    ║                    For detailed examples, see example_usage()              ║
    ╚════════════════════════════════════════════════════════════════════════════╝
    """
    print(guide)


# ============================================================================
# Main Entry Point
# ============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        if sys.argv[1] == "test":
            test_dmp_generator()
        elif sys.argv[1] == "demo":
            quick_demo()
        elif sys.argv[1] == "example":
            example_usage()
        else:
            print("Unknown command. Use: test, demo, or example")
    else:
        # Default: show quick start guide
        quick_start_guide()
        print("\n" + "=" * 80)
        print("Run with arguments:")
        print("  python -m modules.dmp_generator test     # Run test suite")
        print("  python -m modules.dmp_generator demo     # Quick demo")
        print("  python -m modules.dmp_generator example  # Detailed example")
        print("=" * 80)
