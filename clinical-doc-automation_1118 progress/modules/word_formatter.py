"""
Word文件格式化引擎 - Clinical Document Automation
用於確保生成的臨床試驗文件符合公司規範

依賴套件：
    pip install python-docx pillow

作者：Clinical Document Automation Team
版本：1.0.0
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, Dict, Any, Tuple
import os
from datetime import datetime

# 嘗試導入Bestat樣式分析器（如果可用）
try:
    from .bestat_style_analyzer import BestatStyleAnalyzer
    BESTAT_AVAILABLE = True
except ImportError:
    BESTAT_AVAILABLE = False


class WordFormatter:
    """
    Word文件格式化引擎，用於臨床試驗文件的標準化格式設定

    主要功能：
    - 頁面格式設定（A4、邊距等）
    - 字體樣式統一管理
    - 頁首頁尾自動生成
    - 從範本文件讀取樣式
    - 套用公司規範樣式模板
    """

    # 臨床試驗文件標準格式設定
    DEFAULT_CONFIG = {
        # 頁面設定
        'page_size': 'A4',
        'page_width': Inches(8.27),  # A4寬度
        'page_height': Inches(11.69),  # A4高度
        'margin_top': Inches(1.0),
        'margin_bottom': Inches(1.0),
        'margin_left': Inches(1.0),
        'margin_right': Inches(1.0),
        'header_distance': Inches(0.5),
        'footer_distance': Inches(0.5),

        # 字體設定
        'title_font': 'Calibri',
        'title_size': 16,
        'title_bold': True,
        'title_color': RGBColor(0, 51, 102),  # 深藍色

        'heading1_font': 'Calibri',
        'heading1_size': 14,
        'heading1_bold': True,
        'heading1_color': RGBColor(0, 51, 102),

        'heading2_font': 'Calibri',
        'heading2_size': 12,
        'heading2_bold': True,
        'heading2_color': RGBColor(0, 51, 102),

        'body_font': 'Calibri',
        'body_size': 11,
        'body_bold': False,
        'body_color': RGBColor(0, 0, 0),

        'table_font': 'Calibri',
        'table_size': 10,
        'table_bold': False,
        'table_header_bold': True,

        # 行距設定
        'line_spacing': 1.15,
        'paragraph_spacing_before': Pt(0),
        'paragraph_spacing_after': Pt(10),

        # 頁首頁尾設定
        'confidential_text': 'CONFIDENTIAL',
        'show_page_numbers': True,
        'page_number_format': 'Page {0} of {1}',

        # 公司資訊
        'company_name': 'Clinical Research Organization',
        'document_type': 'Clinical Study Protocol',
    }

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        初始化Word格式化引擎

        Args:
            config: 自訂配置字典，會覆蓋DEFAULT_CONFIG中的對應項目
        """
        self.config = self.DEFAULT_CONFIG.copy()
        if config:
            self.config.update(config)

        self.doc = None
        self.template_doc = None
        self.bestat_analyzer = None

        # 初始化Bestat樣式分析器（如果可用）
        if BESTAT_AVAILABLE:
            self.bestat_analyzer = BestatStyleAnalyzer()

    def create_document(self, template_path: Optional[str] = None) -> Document:
        """
        建立新文件或從範本載入

        Args:
            template_path: 範本文件路徑（可選）

        Returns:
            Document: python-docx Document物件
        """
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
            self.template_doc = template_path
            print(f"已從範本載入文件: {template_path}")
        else:
            self.doc = Document()
            print("已建立新文件")

        return self.doc

    def set_page_format(self,
                       page_size: str = 'A4',
                       orientation: str = 'portrait',
                       margin_top: Optional[Inches] = None,
                       margin_bottom: Optional[Inches] = None,
                       margin_left: Optional[Inches] = None,
                       margin_right: Optional[Inches] = None) -> None:
        """
        設定頁面格式

        Args:
            page_size: 頁面大小 ('A4', 'Letter')
            orientation: 頁面方向 ('portrait', 'landscape')
            margin_top: 上邊距
            margin_bottom: 下邊距
            margin_left: 左邊距
            margin_right: 右邊距
        """
        if not self.doc:
            raise ValueError("請先建立或載入文件")

        section = self.doc.sections[0]

        # 設定頁面大小
        if page_size == 'A4':
            if orientation == 'portrait':
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
            else:
                section.page_width = Inches(11.69)
                section.page_height = Inches(8.27)
        elif page_size == 'Letter':
            if orientation == 'portrait':
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
            else:
                section.page_width = Inches(11)
                section.page_height = Inches(8.5)

        # 設定邊距
        section.top_margin = margin_top or self.config['margin_top']
        section.bottom_margin = margin_bottom or self.config['margin_bottom']
        section.left_margin = margin_left or self.config['margin_left']
        section.right_margin = margin_right or self.config['margin_right']

        # 設定頁首頁尾距離
        section.header_distance = self.config['header_distance']
        section.footer_distance = self.config['footer_distance']

        print(f"頁面格式已設定: {page_size}, {orientation}")

    def apply_font_style(self,
                        paragraph,
                        font_name: Optional[str] = None,
                        font_size: Optional[int] = None,
                        bold: Optional[bool] = None,
                        italic: Optional[bool] = None,
                        color: Optional[RGBColor] = None,
                        alignment: Optional[WD_ALIGN_PARAGRAPH] = None) -> None:
        """
        套用字體樣式到段落

        Args:
            paragraph: 段落物件
            font_name: 字體名稱
            font_size: 字體大小（點數）
            bold: 是否粗體
            italic: 是否斜體
            color: 字體顏色
            alignment: 對齊方式
        """
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

        if font_name:
            run.font.name = font_name
            # 設定中文字體
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        if font_size:
            run.font.size = Pt(font_size)

        if bold is not None:
            run.font.bold = bold

        if italic is not None:
            run.font.italic = italic

        if color:
            run.font.color.rgb = color

        if alignment:
            paragraph.alignment = alignment

        # 設定行距
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = self.config['line_spacing']
        paragraph_format.space_before = self.config['paragraph_spacing_before']
        paragraph_format.space_after = self.config['paragraph_spacing_after']

    def set_header(self,
                   document_title: str,
                   protocol_number: str,
                   version: str,
                   logo_path: Optional[str] = None,
                   include_date: bool = True) -> None:
        """
        設定文件頁首

        Args:
            document_title: 文件標題
            protocol_number: Protocol編號
            version: 版本資訊
            logo_path: 公司Logo圖片路徑（可選）
            include_date: 是否包含日期
        """
        if not self.doc:
            raise ValueError("請先建立或載入文件")

        section = self.doc.sections[0]
        header = section.header

        # 清除現有頁首內容
        header.paragraphs[0].clear()

        # 建立表格來組織頁首內容（Logo | 文件資訊）
        if logo_path and os.path.exists(logo_path):
            table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            table.autofit = False

            # 左側：Logo
            logo_cell = table.rows[0].cells[0]
            logo_paragraph = logo_cell.paragraphs[0]
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Inches(1.5))

            # 右側：文件資訊
            info_cell = table.rows[0].cells[1]
            info_cell.width = Inches(5.0)
        else:
            # 只有文件資訊，不包含Logo
            table = header.add_table(rows=1, cols=1, width=Inches(6.5))
            info_cell = table.rows[0].cells[0]

        # 設定文件資訊
        info_paragraph = info_cell.paragraphs[0]
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 添加文件標題
        title_run = info_paragraph.add_run(f"{document_title}\n")
        title_run.font.name = self.config['title_font']
        title_run.font.size = Pt(10)
        title_run.font.bold = True

        # 添加Protocol編號
        protocol_run = info_paragraph.add_run(f"Protocol: {protocol_number}\n")
        protocol_run.font.name = self.config['body_font']
        protocol_run.font.size = Pt(9)

        # 添加版本資訊
        version_run = info_paragraph.add_run(f"Version: {version}")
        version_run.font.name = self.config['body_font']
        version_run.font.size = Pt(9)

        # 添加日期（如果需要）
        if include_date:
            date_run = info_paragraph.add_run(f"\n{datetime.now().strftime('%d-%b-%Y')}")
            date_run.font.name = self.config['body_font']
            date_run.font.size = Pt(9)

        # 添加分隔線
        header.add_paragraph('_' * 80)

        print("頁首已設定完成")

    def set_footer(self,
                   confidential: bool = True,
                   include_page_numbers: bool = True,
                   include_date: bool = True,
                   custom_text: Optional[str] = None) -> None:
        """
        設定文件頁尾

        Args:
            confidential: 是否顯示機密聲明
            include_page_numbers: 是否包含頁碼
            include_date: 是否包含日期
            custom_text: 自訂文字
        """
        if not self.doc:
            raise ValueError("請先建立或載入文件")

        section = self.doc.sections[0]
        footer = section.footer

        # 清除現有頁尾內容
        footer.paragraphs[0].clear()

        # 添加分隔線
        footer.add_paragraph('_' * 80)

        # 建立表格來組織頁尾內容（左：機密聲明 | 中：自訂文字 | 右：頁碼）
        table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        table.autofit = False

        # 左側：機密聲明
        left_cell = table.rows[0].cells[0]
        left_cell.width = Inches(2.0)
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if confidential:
            conf_run = left_paragraph.add_run(self.config['confidential_text'])
            conf_run.font.name = self.config['body_font']
            conf_run.font.size = Pt(8)
            conf_run.font.bold = True
            conf_run.font.color.rgb = RGBColor(255, 0, 0)  # 紅色

        # 中間：日期或自訂文字
        center_cell = table.rows[0].cells[1]
        center_cell.width = Inches(2.5)
        center_paragraph = center_cell.paragraphs[0]
        center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if custom_text:
            custom_run = center_paragraph.add_run(custom_text)
            custom_run.font.name = self.config['body_font']
            custom_run.font.size = Pt(8)
        elif include_date:
            date_run = center_paragraph.add_run(datetime.now().strftime('%d-%b-%Y'))
            date_run.font.name = self.config['body_font']
            date_run.font.size = Pt(8)

        # 右側：頁碼
        right_cell = table.rows[0].cells[2]
        right_cell.width = Inches(2.0)
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if include_page_numbers:
            self._add_page_number(right_paragraph)

        print("頁尾已設定完成")

    def _add_page_number(self, paragraph) -> None:
        """
        添加頁碼到段落

        Args:
            paragraph: 段落物件
        """
        run = paragraph.add_run()
        run.font.name = self.config['body_font']
        run.font.size = Pt(8)

        # 添加 "Page "
        run.text = "Page "

        # 添加當前頁碼
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # 添加 " of "
        run = paragraph.add_run(" of ")
        run.font.name = self.config['body_font']
        run.font.size = Pt(8)

        # 添加總頁數
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')

        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = "NUMPAGES"

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)

    def apply_title_style(self, text: str, level: int = 1) -> Any:
        """
        添加並套用標題樣式

        Args:
            text: 標題文字
            level: 標題層級（1-3）

        Returns:
            段落物件
        """
        if not self.doc:
            raise ValueError("請先建立或載入文件")

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)

        if level == 1:
            run.font.name = self.config['title_font']
            run.font.size = Pt(self.config['title_size'])
            run.font.bold = self.config['title_bold']
            run.font.color.rgb = self.config['title_color']
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)
        elif level == 2:
            run.font.name = self.config['heading1_font']
            run.font.size = Pt(self.config['heading1_size'])
            run.font.bold = self.config['heading1_bold']
            run.font.color.rgb = self.config['heading1_color']
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)
        elif level == 3:
            run.font.name = self.config['heading2_font']
            run.font.size = Pt(self.config['heading2_size'])
            run.font.bold = self.config['heading2_bold']
            run.font.color.rgb = self.config['heading2_color']
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)

        # 設定中文字體
        run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)

        return paragraph

    def apply_body_style(self, text: str, alignment: str = 'left') -> Any:
        """
        添加並套用內文樣式

        Args:
            text: 內文文字
            alignment: 對齊方式 ('left', 'center', 'right', 'justify')

        Returns:
            段落物件
        """
        if not self.doc:
            raise ValueError("請先建立或載入文件")

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)

        run.font.name = self.config['body_font']
        run.font.size = Pt(self.config['body_size'])
        run.font.bold = self.config['body_bold']
        run.font.color.rgb = self.config['body_color']

        # 設定中文字體
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['body_font'])

        # 設定對齊方式
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # 設定段落格式
        paragraph.paragraph_format.line_spacing = self.config['line_spacing']
        paragraph.paragraph_format.space_after = Pt(10)

        return paragraph

    def apply_table_style(self,
                         table,
                         style_name: str = 'Light Grid Accent 1',
                         header_row: bool = True,
                         auto_fit: bool = True) -> None:
        """
        套用表格樣式

        Args:
            table: 表格物件
            style_name: 表格樣式名稱
            header_row: 是否有標題列
            auto_fit: 是否自動調整欄寬
        """
        # 設定表格樣式
        try:
            table.style = style_name
        except:
            print(f"警告：找不到樣式 '{style_name}'，使用預設樣式")

        # 設定表格字體
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.config['table_font']
                        run.font.size = Pt(self.config['table_size'])
                        # 設定中文字體
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['table_font'])

        # 設定標題列為粗體
        if header_row and len(table.rows) > 0:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = self.config['table_header_bold']

        # 自動調整欄寬
        if auto_fit:
            table.autofit = True

    def load_styles_from_template(self, template_path: str) -> Dict[str, Any]:
        """
        從範本文件載入樣式設定

        Args:
            template_path: 範本文件路徑

        Returns:
            樣式設定字典
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"找不到範本文件: {template_path}")

        template_doc = Document(template_path)
        styles = {}

        # 載入段落樣式
        for style in template_doc.styles:
            if style.type == 1:  # 段落樣式
                style_info = {
                    'name': style.name,
                    'font_name': style.font.name if style.font.name else None,
                    'font_size': style.font.size.pt if style.font.size else None,
                    'font_bold': style.font.bold if style.font.bold else False,
                }
                styles[style.name] = style_info

        print(f"已從範本載入 {len(styles)} 個樣式")
        return styles

    def apply_clinical_trial_template(self,
                                     document_title: str,
                                     protocol_number: str,
                                     version: str,
                                     sponsor: str = "",
                                     indication: str = "",
                                     logo_path: Optional[str] = None) -> None:
        """
        套用臨床試驗文件標準範本

        Args:
            document_title: 文件標題
            protocol_number: Protocol編號
            version: 版本資訊
            sponsor: 贊助商名稱
            indication: 適應症
            logo_path: 公司Logo路徑
        """
        if not self.doc:
            self.create_document()

        # 設定頁面格式
        self.set_page_format()

        # 設定頁首
        self.set_header(
            document_title=document_title,
            protocol_number=protocol_number,
            version=version,
            logo_path=logo_path
        )

        # 設定頁尾
        self.set_footer(
            confidential=True,
            include_page_numbers=True,
            include_date=True
        )

        # 添加封面頁
        self._create_cover_page(
            document_title=document_title,
            protocol_number=protocol_number,
            version=version,
            sponsor=sponsor,
            indication=indication
        )

        print("臨床試驗文件範本已套用完成")

    def _create_cover_page(self,
                          document_title: str,
                          protocol_number: str,
                          version: str,
                          sponsor: str,
                          indication: str) -> None:
        """
        建立封面頁

        Args:
            document_title: 文件標題
            protocol_number: Protocol編號
            version: 版本資訊
            sponsor: 贊助商
            indication: 適應症
        """
        # 文件標題
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Inches(2)
        title_run = title_para.add_run(document_title.upper())
        title_run.font.name = self.config['title_font']
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = self.config['title_color']

        # 添加空行
        self.doc.add_paragraph()

        # Protocol編號
        protocol_para = self.doc.add_paragraph()
        protocol_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        protocol_run = protocol_para.add_run(f"Protocol Number: {protocol_number}")
        protocol_run.font.name = self.config['body_font']
        protocol_run.font.size = Pt(14)
        protocol_run.font.bold = True

        # 版本
        version_para = self.doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_run = version_para.add_run(f"Version: {version}")
        version_run.font.name = self.config['body_font']
        version_run.font.size = Pt(12)

        # 添加空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 贊助商
        if sponsor:
            sponsor_para = self.doc.add_paragraph()
            sponsor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sponsor_run = sponsor_para.add_run(f"Sponsor: {sponsor}")
            sponsor_run.font.name = self.config['body_font']
            sponsor_run.font.size = Pt(12)

        # 適應症
        if indication:
            indication_para = self.doc.add_paragraph()
            indication_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            indication_run = indication_para.add_run(f"Indication: {indication}")
            indication_run.font.name = self.config['body_font']
            indication_run.font.size = Pt(12)

        # 添加空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime('%d %B %Y'))
        date_run.font.name = self.config['body_font']
        date_run.font.size = Pt(12)

        # 機密聲明
        self.doc.add_paragraph()
        conf_para = self.doc.add_paragraph()
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run = conf_para.add_run("CONFIDENTIAL")
        conf_run.font.name = self.config['body_font']
        conf_run.font.size = Pt(14)
        conf_run.font.bold = True
        conf_run.font.color.rgb = RGBColor(255, 0, 0)

        # 添加分頁
        self.doc.add_page_break()

    def save_document(self, output_path: str) -> None:
        """
        儲存文件

        Args:
            output_path: 輸出文件路徑
        """
        if not self.doc:
            raise ValueError("沒有文件可儲存")

        # 確保輸出目錄存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        self.doc.save(output_path)
        print(f"文件已儲存至: {output_path}")

    def get_config(self) -> Dict[str, Any]:
        """
        取得當前配置

        Returns:
            配置字典
        """
        return self.config.copy()

    def update_config(self, config_updates: Dict[str, Any]) -> None:
        """
        更新配置

        Args:
            config_updates: 要更新的配置項目
        """
        self.config.update(config_updates)
        print(f"已更新 {len(config_updates)} 個配置項目")

    def apply_bestat_style(self,
                          document_title: str = "",
                          protocol_number: str = "",
                          version: str = "",
                          logo_path: Optional[str] = None,
                          bestat_config_path: Optional[str] = None) -> None:
        """
        套用Bestat公司標準樣式

        Args:
            document_title: 文件標題
            protocol_number: Protocol編號
            version: 版本號
            logo_path: Bestat Logo路徑
            bestat_config_path: Bestat樣式配置JSON路徑（可選）

        Raises:
            RuntimeError: 如果Bestat樣式分析器不可用
        """
        if not BESTAT_AVAILABLE or not self.bestat_analyzer:
            raise RuntimeError(
                "Bestat樣式分析器不可用。請確保bestat_style_analyzer.py在相同目錄下。"
            )

        if not self.doc:
            raise ValueError("請先建立或載入文件")

        # 載入自訂配置（如果提供）
        if bestat_config_path:
            self.bestat_analyzer.load_style_config(bestat_config_path)

        # 套用Bestat樣式
        self.doc = self.bestat_analyzer.apply_bestat_style(
            self.doc,
            document_title=document_title,
            protocol_number=protocol_number,
            version=version,
            logo_path=logo_path
        )

        print("✓ Bestat樣式已套用完成")

    def validate_bestat_compliance(self) -> Dict[str, Any]:
        """
        驗證當前文件是否符合Bestat規範

        Returns:
            驗證結果字典

        Raises:
            RuntimeError: 如果Bestat樣式分析器不可用
            ValueError: 如果尚未建立文件
        """
        if not BESTAT_AVAILABLE or not self.bestat_analyzer:
            raise RuntimeError(
                "Bestat樣式分析器不可用。請確保bestat_style_analyzer.py在相同目錄下。"
            )

        if not self.doc:
            raise ValueError("請先建立或載入文件")

        return self.bestat_analyzer.validate_bestat_compliance(self.doc)

    def extract_bestat_styles(self, template_path: str, output_json_path: str) -> Dict[str, Any]:
        """
        從範本文件提取Bestat樣式並儲存為JSON

        Args:
            template_path: 範本文件路徑
            output_json_path: 輸出JSON路徑

        Returns:
            提取的樣式字典

        Raises:
            RuntimeError: 如果Bestat樣式分析器不可用
        """
        if not BESTAT_AVAILABLE or not self.bestat_analyzer:
            raise RuntimeError(
                "Bestat樣式分析器不可用。請確保bestat_style_analyzer.py在相同目錄下。"
            )

        styles = self.bestat_analyzer.extract_styles_from_document(template_path)
        self.bestat_analyzer.save_style_config(output_json_path, styles)

        return styles

    def load_bestat_config(self, config_path: str) -> None:
        """
        載入Bestat樣式配置

        Args:
            config_path: JSON配置檔案路徑

        Raises:
            RuntimeError: 如果Bestat樣式分析器不可用
        """
        if not BESTAT_AVAILABLE or not self.bestat_analyzer:
            raise RuntimeError(
                "Bestat樣式分析器不可用。請確保bestat_style_analyzer.py在相同目錄下。"
            )

        self.bestat_analyzer.load_style_config(config_path)
        print(f"✓ 已載入Bestat配置: {config_path}")


# 便利函數
def create_clinical_document(document_title: str,
                            protocol_number: str,
                            version: str,
                            output_path: str,
                            sponsor: str = "",
                            indication: str = "",
                            logo_path: Optional[str] = None,
                            custom_config: Optional[Dict[str, Any]] = None) -> WordFormatter:
    """
    快速建立臨床試驗文件

    Args:
        document_title: 文件標題
        protocol_number: Protocol編號
        version: 版本資訊
        output_path: 輸出路徑
        sponsor: 贊助商名稱
        indication: 適應症
        logo_path: 公司Logo路徑
        custom_config: 自訂配置

    Returns:
        WordFormatter實例，可繼續添加內容
    """
    formatter = WordFormatter(config=custom_config)
    formatter.create_document()
    formatter.apply_clinical_trial_template(
        document_title=document_title,
        protocol_number=protocol_number,
        version=version,
        sponsor=sponsor,
        indication=indication,
        logo_path=logo_path
    )

    return formatter


if __name__ == "__main__":
    # 測試代碼
    print("Word文件格式化引擎測試")
    print("=" * 60)

    # 建立簡單的測試文件
    formatter = WordFormatter()
    formatter.create_document()
    formatter.set_page_format()
    formatter.set_header(
        document_title="Clinical Study Protocol",
        protocol_number="PRO-2025-001",
        version="1.0"
    )
    formatter.set_footer()

    # 添加標題
    formatter.apply_title_style("1. Introduction", level=2)
    formatter.apply_body_style("This is a sample clinical trial document created using the WordFormatter engine.")

    print("\n測試完成！")
