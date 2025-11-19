"""
Bestat公司文件樣式分析器和範本生成器 - Clinical Document Automation
用於提取、分析、套用和驗證Bestat公司的文件樣式規範

依賴套件：
    pip install python-docx pillow

功能：
    - 從範本Word文件提取完整樣式設定
    - 生成JSON格式的樣式配置檔案
    - 套用Bestat樣式到新文件
    - 驗證文件是否符合Bestat規範
    - 比較兩個文件的樣式差異

作者：Clinical Document Automation Team
版本：1.0.0
日期：2025-11-18
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from typing import Dict, Any, List, Optional, Tuple
import json
import os
from datetime import datetime
from pathlib import Path


class BestatStyleAnalyzer:
    """
    Bestat公司文件樣式分析器

    主要功能：
    - 分析Word文件並提取所有樣式設定
    - 生成標準化的樣式配置檔案
    - 套用Bestat公司標準樣式
    - 驗證文件是否符合Bestat規範
    - 比較文件樣式差異
    """

    # Bestat公司預設樣式規範
    BESTAT_DEFAULT_STYLE = {
        "company_info": {
            "name": "Bestat Inc.",
            "full_name": "Bestat Clinical Research Organization",
            "website": "www.bestat.com",
            "established": "2010"
        },
        "page_setup": {
            "size": "A4",
            "width_inches": 8.27,
            "height_inches": 11.69,
            "orientation": "portrait",
            "margins": {
                "top_inches": 1.0,
                "bottom_inches": 1.0,
                "left_inches": 1.0,
                "right_inches": 1.0
            },
            "header_distance_inches": 0.5,
            "footer_distance_inches": 0.5
        },
        "fonts": {
            "title": {
                "name": "Arial",
                "size": 18,
                "bold": True,
                "italic": False,
                "color": {"r": 0, "g": 51, "b": 153},  # Bestat藍色
                "alignment": "center"
            },
            "heading1": {
                "name": "Arial",
                "size": 16,
                "bold": True,
                "italic": False,
                "color": {"r": 0, "g": 51, "b": 153},
                "alignment": "left"
            },
            "heading2": {
                "name": "Arial",
                "size": 14,
                "bold": True,
                "italic": False,
                "color": {"r": 0, "g": 102, "b": 204},
                "alignment": "left"
            },
            "heading3": {
                "name": "Arial",
                "size": 12,
                "bold": True,
                "italic": False,
                "color": {"r": 0, "g": 102, "b": 204},
                "alignment": "left"
            },
            "body": {
                "name": "Calibri",
                "size": 11,
                "bold": False,
                "italic": False,
                "color": {"r": 0, "g": 0, "b": 0},
                "alignment": "justify"
            },
            "table_header": {
                "name": "Arial",
                "size": 10,
                "bold": True,
                "italic": False,
                "color": {"r": 255, "g": 255, "b": 255},  # 白色
                "background_color": {"r": 0, "g": 51, "b": 153}  # Bestat藍色背景
            },
            "table_body": {
                "name": "Calibri",
                "size": 10,
                "bold": False,
                "italic": False,
                "color": {"r": 0, "g": 0, "b": 0}
            }
        },
        "colors": {
            "primary": {"r": 0, "g": 51, "b": 153},  # Bestat主色（深藍）
            "secondary": {"r": 0, "g": 102, "b": 204},  # Bestat輔色（亮藍）
            "accent": {"r": 255, "g": 153, "b": 0},  # 強調色（橙色）
            "text": {"r": 0, "g": 0, "b": 0},  # 文字色（黑色）
            "background": {"r": 255, "g": 255, "b": 255},  # 背景色（白色）
            "confidential": {"r": 255, "g": 0, "b": 0}  # 機密標記（紅色）
        },
        "paragraph_spacing": {
            "line_spacing": 1.15,
            "before_pt": 0,
            "after_pt": 10,
            "first_line_indent_inches": 0
        },
        "header": {
            "enabled": True,
            "logo_position": "left",
            "logo_width_inches": 1.5,
            "logo_height_inches": 0.6,
            "text_position": "right",
            "font_name": "Arial",
            "font_size": 9,
            "include_date": True,
            "include_protocol_number": True,
            "include_version": True,
            "separator_line": True
        },
        "footer": {
            "enabled": True,
            "font_name": "Arial",
            "font_size": 8,
            "show_confidential": True,
            "confidential_position": "left",
            "show_page_numbers": True,
            "page_number_position": "right",
            "page_number_format": "Page {page} of {total}",
            "show_date": True,
            "date_position": "center",
            "separator_line": True
        },
        "table_style": {
            "default_style": "Light Grid Accent 1",
            "border_color": {"r": 0, "g": 51, "b": 153},
            "border_width": 0.5,
            "cell_padding_inches": 0.05,
            "auto_fit": True
        }
    }

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        初始化Bestat樣式分析器

        Args:
            config: 自訂配置，會與預設配置合併
        """
        self.style_config = self._deep_merge(
            self.BESTAT_DEFAULT_STYLE.copy(),
            config or {}
        )

    def _deep_merge(self, base: Dict, overlay: Dict) -> Dict:
        """
        深度合併兩個字典

        Args:
            base: 基礎字典
            overlay: 覆蓋字典

        Returns:
            合併後的字典
        """
        result = base.copy()
        for key, value in overlay.items():
            if key in result and isinstance(result[key], dict) and isinstance(value, dict):
                result[key] = self._deep_merge(result[key], value)
            else:
                result[key] = value
        return result

    def extract_styles_from_document(self, document_path: str) -> Dict[str, Any]:
        """
        從Word文件提取完整的樣式設定

        Args:
            document_path: Word文件路徑

        Returns:
            包含所有樣式資訊的字典
        """
        if not os.path.exists(document_path):
            raise FileNotFoundError(f"找不到文件: {document_path}")

        doc = Document(document_path)
        extracted_styles = {}

        # 1. 提取頁面設定
        extracted_styles['page_setup'] = self._extract_page_setup(doc)

        # 2. 提取字體樣式
        extracted_styles['fonts'] = self._extract_font_styles(doc)

        # 3. 提取顏色方案
        extracted_styles['colors'] = self._extract_color_scheme(doc)

        # 4. 提取段落樣式
        extracted_styles['paragraph_spacing'] = self._extract_paragraph_styles(doc)

        # 5. 提取頁首設定
        extracted_styles['header'] = self._extract_header_info(doc)

        # 6. 提取頁尾設定
        extracted_styles['footer'] = self._extract_footer_info(doc)

        # 7. 提取表格樣式
        extracted_styles['table_style'] = self._extract_table_styles(doc)

        # 8. 提取Logo資訊
        extracted_styles['logo_info'] = self._extract_logo_info(doc)

        # 9. 添加元數據
        extracted_styles['metadata'] = {
            'extracted_from': document_path,
            'extracted_date': datetime.now().isoformat(),
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'total_sections': len(doc.sections)
        }

        print(f"✓ 已從文件提取樣式: {document_path}")
        return extracted_styles

    def _extract_page_setup(self, doc: Document) -> Dict[str, Any]:
        """提取頁面設定"""
        section = doc.sections[0]

        page_setup = {
            "width_inches": round(section.page_width.inches, 2),
            "height_inches": round(section.page_height.inches, 2),
            "orientation": "portrait" if section.page_width < section.page_height else "landscape",
            "margins": {
                "top_inches": round(section.top_margin.inches, 2),
                "bottom_inches": round(section.bottom_margin.inches, 2),
                "left_inches": round(section.left_margin.inches, 2),
                "right_inches": round(section.right_margin.inches, 2)
            },
            "header_distance_inches": round(section.header_distance.inches, 2),
            "footer_distance_inches": round(section.footer_distance.inches, 2)
        }

        # 判斷頁面大小
        if abs(page_setup['width_inches'] - 8.27) < 0.1:
            page_setup['size'] = 'A4'
        elif abs(page_setup['width_inches'] - 8.5) < 0.1:
            page_setup['size'] = 'Letter'
        else:
            page_setup['size'] = 'Custom'

        return page_setup

    def _extract_font_styles(self, doc: Document) -> Dict[str, Any]:
        """提取字體樣式"""
        font_styles = {}

        # 從文件樣式中提取
        for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Normal']:
            try:
                style = doc.styles[style_name]
                if hasattr(style, 'font'):
                    font_info = {
                        "name": style.font.name or "Calibri",
                        "size": int(style.font.size.pt) if style.font.size else 11,
                        "bold": bool(style.font.bold),
                        "italic": bool(style.font.italic),
                    }

                    # 提取顏色
                    if style.font.color and style.font.color.rgb:
                        rgb = style.font.color.rgb
                        font_info["color"] = {
                            "r": rgb[0] if len(rgb) > 0 else 0,
                            "g": rgb[1] if len(rgb) > 1 else 0,
                            "b": rgb[2] if len(rgb) > 2 else 0
                        }
                    else:
                        font_info["color"] = {"r": 0, "g": 0, "b": 0}

                    # 轉換樣式名稱
                    key = style_name.lower().replace(' ', '_')
                    if key == 'title':
                        font_styles['title'] = font_info
                    elif key == 'heading_1':
                        font_styles['heading1'] = font_info
                    elif key == 'heading_2':
                        font_styles['heading2'] = font_info
                    elif key == 'heading_3':
                        font_styles['heading3'] = font_info
                    elif key == 'normal':
                        font_styles['body'] = font_info
            except:
                continue

        return font_styles

    def _extract_color_scheme(self, doc: Document) -> Dict[str, Any]:
        """提取顏色方案"""
        colors = {
            "primary": {"r": 0, "g": 51, "b": 153},
            "secondary": {"r": 0, "g": 102, "b": 204},
            "text": {"r": 0, "g": 0, "b": 0},
            "background": {"r": 255, "g": 255, "b": 255}
        }

        # 從段落中提取顏色
        color_counts = {}
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    color_key = f"{rgb[0]},{rgb[1]},{rgb[2]}"
                    color_counts[color_key] = color_counts.get(color_key, 0) + 1

        # 找出最常用的顏色作為主色
        if color_counts:
            sorted_colors = sorted(color_counts.items(), key=lambda x: x[1], reverse=True)
            if len(sorted_colors) > 0:
                r, g, b = map(int, sorted_colors[0][0].split(','))
                if not (r == 0 and g == 0 and b == 0):  # 排除黑色
                    colors['primary'] = {"r": r, "g": g, "b": b}

        return colors

    def _extract_paragraph_styles(self, doc: Document) -> Dict[str, Any]:
        """提取段落樣式"""
        paragraph_styles = {
            "line_spacing": 1.15,
            "before_pt": 0,
            "after_pt": 10,
            "first_line_indent_inches": 0
        }

        # 從第一個非空段落提取設定
        for para in doc.paragraphs:
            if para.text.strip():
                if para.paragraph_format.line_spacing:
                    paragraph_styles['line_spacing'] = round(para.paragraph_format.line_spacing, 2)
                if para.paragraph_format.space_before:
                    paragraph_styles['before_pt'] = round(para.paragraph_format.space_before.pt, 1)
                if para.paragraph_format.space_after:
                    paragraph_styles['after_pt'] = round(para.paragraph_format.space_after.pt, 1)
                if para.paragraph_format.first_line_indent:
                    paragraph_styles['first_line_indent_inches'] = round(
                        para.paragraph_format.first_line_indent.inches, 2
                    )
                break

        return paragraph_styles

    def _extract_header_info(self, doc: Document) -> Dict[str, Any]:
        """提取頁首資訊"""
        header_info = {
            "enabled": False,
            "font_name": "Arial",
            "font_size": 9,
            "include_date": False,
            "include_protocol_number": False,
            "include_version": False,
            "separator_line": False
        }

        if doc.sections:
            section = doc.sections[0]
            if section.header and section.header.paragraphs:
                header_info['enabled'] = True

                # 檢查頁首內容
                header_text = ' '.join([p.text for p in section.header.paragraphs])
                header_info['include_date'] = any(
                    date_format in header_text.lower()
                    for date_format in ['date', 'dated', '202', '2025']
                )
                header_info['include_protocol_number'] = 'protocol' in header_text.lower()
                header_info['include_version'] = 'version' in header_text.lower()
                header_info['separator_line'] = '_' in header_text or '-' * 5 in header_text

                # 提取字體資訊
                for para in section.header.paragraphs:
                    if para.runs:
                        run = para.runs[0]
                        if run.font.name:
                            header_info['font_name'] = run.font.name
                        if run.font.size:
                            header_info['font_size'] = int(run.font.size.pt)
                        break

        return header_info

    def _extract_footer_info(self, doc: Document) -> Dict[str, Any]:
        """提取頁尾資訊"""
        footer_info = {
            "enabled": False,
            "font_name": "Arial",
            "font_size": 8,
            "show_confidential": False,
            "show_page_numbers": False,
            "show_date": False,
            "separator_line": False
        }

        if doc.sections:
            section = doc.sections[0]
            if section.footer and section.footer.paragraphs:
                footer_info['enabled'] = True

                # 檢查頁尾內容
                footer_text = ' '.join([p.text for p in section.footer.paragraphs])
                footer_info['show_confidential'] = 'confidential' in footer_text.lower()
                footer_info['show_page_numbers'] = 'page' in footer_text.lower()
                footer_info['show_date'] = any(
                    date_format in footer_text.lower()
                    for date_format in ['date', '202', '2025']
                )
                footer_info['separator_line'] = '_' in footer_text or '-' * 5 in footer_text

                # 提取字體資訊
                for para in section.footer.paragraphs:
                    if para.runs:
                        run = para.runs[0]
                        if run.font.name:
                            footer_info['font_name'] = run.font.name
                        if run.font.size:
                            footer_info['font_size'] = int(run.font.size.pt)
                        break

        return footer_info

    def _extract_table_styles(self, doc: Document) -> Dict[str, Any]:
        """提取表格樣式"""
        table_styles = {
            "default_style": "Light Grid Accent 1",
            "border_width": 0.5,
            "cell_padding_inches": 0.05,
            "auto_fit": True
        }

        if doc.tables:
            table = doc.tables[0]
            if table.style:
                table_styles['default_style'] = str(table.style.name)

        return table_styles

    def _extract_logo_info(self, doc: Document) -> Dict[str, Any]:
        """提取Logo資訊"""
        logo_info = {
            "has_logo": False,
            "position": "left",
            "width_inches": 1.5,
            "height_inches": 0.6
        }

        # 檢查頁首中的圖片
        if doc.sections:
            section = doc.sections[0]
            if section.header:
                for para in section.header.paragraphs:
                    for run in para.runs:
                        if hasattr(run, '_element'):
                            drawings = run._element.findall('.//{*}drawing')
                            if drawings:
                                logo_info['has_logo'] = True
                                # 可以進一步提取圖片尺寸
                                break

        return logo_info

    def save_style_config(self, output_path: str, style_config: Optional[Dict[str, Any]] = None) -> None:
        """
        儲存樣式配置到JSON檔案

        Args:
            output_path: 輸出JSON檔案路徑
            style_config: 樣式配置字典（None則使用當前配置）
        """
        config_to_save = style_config or self.style_config

        # 確保輸出目錄存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 轉換RGBColor物件為可序列化的字典
        serializable_config = self._make_serializable(config_to_save)

        # 儲存為格式化的JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(serializable_config, f, indent=2, ensure_ascii=False)

        print(f"✓ 樣式配置已儲存至: {output_path}")

    def _make_serializable(self, obj: Any) -> Any:
        """將物件轉換為可JSON序列化的格式"""
        if isinstance(obj, dict):
            return {k: self._make_serializable(v) for k, v in obj.items()}
        elif isinstance(obj, (list, tuple)):
            return [self._make_serializable(item) for item in obj]
        elif isinstance(obj, RGBColor):
            return {"r": obj[0], "g": obj[1], "b": obj[2]}
        elif isinstance(obj, (Inches, Cm, Pt)):
            return float(obj)
        else:
            return obj

    def load_style_config(self, config_path: str) -> Dict[str, Any]:
        """
        從JSON檔案載入樣式配置

        Args:
            config_path: JSON配置檔案路徑

        Returns:
            樣式配置字典
        """
        if not os.path.exists(config_path):
            raise FileNotFoundError(f"找不到配置檔案: {config_path}")

        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)

        self.style_config = self._deep_merge(self.BESTAT_DEFAULT_STYLE.copy(), config)
        print(f"✓ 已載入樣式配置: {config_path}")

        return self.style_config

    def apply_bestat_style(self, doc: Document,
                          document_title: str = "",
                          protocol_number: str = "",
                          version: str = "",
                          logo_path: Optional[str] = None) -> Document:
        """
        套用Bestat樣式到文件

        Args:
            doc: python-docx Document物件
            document_title: 文件標題
            protocol_number: Protocol編號
            version: 版本號
            logo_path: Logo圖片路徑

        Returns:
            套用樣式後的Document物件
        """
        print("開始套用Bestat樣式...")

        # 1. 套用頁面設定
        self._apply_page_setup(doc)

        # 2. 套用頁首
        if self.style_config['header']['enabled']:
            self._apply_header(doc, document_title, protocol_number, version, logo_path)

        # 3. 套用頁尾
        if self.style_config['footer']['enabled']:
            self._apply_footer(doc)

        # 4. 套用段落樣式
        self._apply_paragraph_styles(doc)

        # 5. 套用表格樣式
        self._apply_table_styles(doc)

        print("✓ Bestat樣式套用完成")
        return doc

    def _apply_page_setup(self, doc: Document) -> None:
        """套用頁面設定"""
        section = doc.sections[0]
        page_setup = self.style_config['page_setup']

        section.page_width = Inches(page_setup['width_inches'])
        section.page_height = Inches(page_setup['height_inches'])
        section.top_margin = Inches(page_setup['margins']['top_inches'])
        section.bottom_margin = Inches(page_setup['margins']['bottom_inches'])
        section.left_margin = Inches(page_setup['margins']['left_inches'])
        section.right_margin = Inches(page_setup['margins']['right_inches'])
        section.header_distance = Inches(page_setup['header_distance_inches'])
        section.footer_distance = Inches(page_setup['footer_distance_inches'])

        print("  ✓ 頁面設定已套用")

    def _apply_header(self, doc: Document, title: str, protocol: str, version: str, logo_path: Optional[str]) -> None:
        """套用頁首"""
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].clear()

        header_config = self.style_config['header']

        # 建立表格佈局
        if logo_path and os.path.exists(logo_path):
            table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            logo_cell = table.rows[0].cells[0]
            logo_para = logo_cell.paragraphs[0]
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Inches(header_config['logo_width_inches']))
            info_cell = table.rows[0].cells[1]
        else:
            table = header.add_table(rows=1, cols=1, width=Inches(6.5))
            info_cell = table.rows[0].cells[0]

        # 設定文件資訊
        info_para = info_cell.paragraphs[0]
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if title:
            title_run = info_para.add_run(f"{title}\n")
            title_run.font.name = header_config['font_name']
            title_run.font.size = Pt(header_config['font_size'] + 1)
            title_run.font.bold = True

        if protocol and header_config['include_protocol_number']:
            protocol_run = info_para.add_run(f"Protocol: {protocol}\n")
            protocol_run.font.name = header_config['font_name']
            protocol_run.font.size = Pt(header_config['font_size'])

        if version and header_config['include_version']:
            version_run = info_para.add_run(f"Version: {version}\n")
            version_run.font.name = header_config['font_name']
            version_run.font.size = Pt(header_config['font_size'])

        if header_config['include_date']:
            date_run = info_para.add_run(datetime.now().strftime('%d-%b-%Y'))
            date_run.font.name = header_config['font_name']
            date_run.font.size = Pt(header_config['font_size'])

        if header_config['separator_line']:
            sep_para = header.add_paragraph('_' * 80)
            sep_para.runs[0].font.size = Pt(8)

        print("  ✓ 頁首已套用")

    def _apply_footer(self, doc: Document) -> None:
        """套用頁尾"""
        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].clear()

        footer_config = self.style_config['footer']

        if footer_config['separator_line']:
            sep_para = footer.add_paragraph('_' * 80)
            sep_para.runs[0].font.size = Pt(8)

        # 建立表格佈局
        table = footer.add_table(rows=1, cols=3, width=Inches(6.5))

        # 左側：機密聲明
        left_cell = table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if footer_config['show_confidential']:
            conf_run = left_para.add_run("CONFIDENTIAL")
            conf_run.font.name = footer_config['font_name']
            conf_run.font.size = Pt(footer_config['font_size'])
            conf_run.font.bold = True
            color = self.style_config['colors']['confidential']
            conf_run.font.color.rgb = RGBColor(color['r'], color['g'], color['b'])

        # 中間：日期
        center_cell = table.rows[0].cells[1]
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if footer_config['show_date']:
            date_run = center_para.add_run(datetime.now().strftime('%d-%b-%Y'))
            date_run.font.name = footer_config['font_name']
            date_run.font.size = Pt(footer_config['font_size'])

        # 右側：頁碼
        right_cell = table.rows[0].cells[2]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if footer_config['show_page_numbers']:
            self._add_page_number_field(right_para, footer_config)

        print("  ✓ 頁尾已套用")

    def _add_page_number_field(self, paragraph, footer_config: Dict) -> None:
        """添加頁碼欄位"""
        from docx.oxml import OxmlElement

        run = paragraph.add_run("Page ")
        run.font.name = footer_config['font_name']
        run.font.size = Pt(footer_config['font_size'])

        # 當前頁碼
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # " of "
        run = paragraph.add_run(" of ")
        run.font.name = footer_config['font_name']
        run.font.size = Pt(footer_config['font_size'])

        # 總頁數
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)

    def _apply_paragraph_styles(self, doc: Document) -> None:
        """套用段落樣式"""
        spacing = self.style_config['paragraph_spacing']

        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = spacing['line_spacing']
            para.paragraph_format.space_before = Pt(spacing['before_pt'])
            para.paragraph_format.space_after = Pt(spacing['after_pt'])

            if spacing['first_line_indent_inches'] > 0:
                para.paragraph_format.first_line_indent = Inches(spacing['first_line_indent_inches'])

        print("  ✓ 段落樣式已套用")

    def _apply_table_styles(self, doc: Document) -> None:
        """套用表格樣式"""
        table_config = self.style_config['table_style']

        for table in doc.tables:
            try:
                table.style = table_config['default_style']
            except:
                pass

            if table_config['auto_fit']:
                table.autofit = True

        if doc.tables:
            print(f"  ✓ 已套用表格樣式到 {len(doc.tables)} 個表格")

    def validate_bestat_compliance(self, doc: Document) -> Dict[str, Any]:
        """
        驗證文件是否符合Bestat規範

        Args:
            doc: python-docx Document物件

        Returns:
            驗證結果字典，包含是否通過和詳細問題清單
        """
        print("開始驗證Bestat規範...")

        issues = []
        warnings = []

        # 1. 檢查頁面設定
        page_issues = self._validate_page_setup(doc)
        issues.extend(page_issues)

        # 2. 檢查頁首
        header_issues = self._validate_header(doc)
        issues.extend(header_issues)

        # 3. 檢查頁尾
        footer_issues = self._validate_footer(doc)
        issues.extend(footer_issues)

        # 4. 檢查字體
        font_warnings = self._validate_fonts(doc)
        warnings.extend(font_warnings)

        # 5. 檢查表格
        table_warnings = self._validate_tables(doc)
        warnings.extend(table_warnings)

        result = {
            "compliant": len(issues) == 0,
            "total_issues": len(issues),
            "total_warnings": len(warnings),
            "issues": issues,
            "warnings": warnings,
            "validation_date": datetime.now().isoformat()
        }

        if result['compliant']:
            print("✓ 文件符合Bestat規範")
        else:
            print(f"✗ 發現 {len(issues)} 個問題，{len(warnings)} 個警告")

        return result

    def _validate_page_setup(self, doc: Document) -> List[str]:
        """驗證頁面設定"""
        issues = []
        section = doc.sections[0]
        expected = self.style_config['page_setup']

        # 檢查頁面尺寸（允許0.1英寸誤差）
        if abs(section.page_width.inches - expected['width_inches']) > 0.1:
            issues.append(f"頁面寬度不符: 期望 {expected['width_inches']}\"，實際 {section.page_width.inches:.2f}\"")

        if abs(section.page_height.inches - expected['height_inches']) > 0.1:
            issues.append(f"頁面高度不符: 期望 {expected['height_inches']}\"，實際 {section.page_height.inches:.2f}\"")

        # 檢查邊距（允許0.2英寸誤差）
        margins = expected['margins']
        if abs(section.top_margin.inches - margins['top_inches']) > 0.2:
            issues.append(f"上邊距不符: 期望 {margins['top_inches']}\"，實際 {section.top_margin.inches:.2f}\"")

        return issues

    def _validate_header(self, doc: Document) -> List[str]:
        """驗證頁首"""
        issues = []

        if not doc.sections:
            issues.append("文件缺少section")
            return issues

        section = doc.sections[0]
        if not section.header or not section.header.paragraphs:
            if self.style_config['header']['enabled']:
                issues.append("缺少頁首")

        return issues

    def _validate_footer(self, doc: Document) -> List[str]:
        """驗證頁尾"""
        issues = []

        if not doc.sections:
            issues.append("文件缺少section")
            return issues

        section = doc.sections[0]
        if not section.footer or not section.footer.paragraphs:
            if self.style_config['footer']['enabled']:
                issues.append("缺少頁尾")

        return issues

    def _validate_fonts(self, doc: Document) -> List[str]:
        """驗證字體（警告級別）"""
        warnings = []
        expected_fonts = set()

        for font_type in ['title', 'heading1', 'heading2', 'heading3', 'body']:
            if font_type in self.style_config['fonts']:
                expected_fonts.add(self.style_config['fonts'][font_type]['name'])

        # 收集實際使用的字體
        actual_fonts = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    actual_fonts.add(run.font.name)

        # 檢查非標準字體
        unexpected_fonts = actual_fonts - expected_fonts
        if unexpected_fonts:
            warnings.append(f"使用了非標準字體: {', '.join(unexpected_fonts)}")

        return warnings

    def _validate_tables(self, doc: Document) -> List[str]:
        """驗證表格（警告級別）"""
        warnings = []

        if not doc.tables:
            return warnings

        expected_style = self.style_config['table_style']['default_style']

        for i, table in enumerate(doc.tables):
            if table.style and str(table.style.name) != expected_style:
                warnings.append(f"表格 {i+1} 樣式不符: 期望 '{expected_style}'，實際 '{table.style.name}'")

        return warnings

    def compare_styles(self, doc1_path: str, doc2_path: str) -> Dict[str, Any]:
        """
        比較兩個文件的樣式差異

        Args:
            doc1_path: 第一個文件路徑
            doc2_path: 第二個文件路徑

        Returns:
            包含差異資訊的字典
        """
        print(f"比較樣式差異: {os.path.basename(doc1_path)} vs {os.path.basename(doc2_path)}")

        # 提取兩個文件的樣式
        style1 = self.extract_styles_from_document(doc1_path)
        style2 = self.extract_styles_from_document(doc2_path)

        differences = {
            "page_setup": self._compare_dicts(
                style1.get('page_setup', {}),
                style2.get('page_setup', {})
            ),
            "fonts": self._compare_dicts(
                style1.get('fonts', {}),
                style2.get('fonts', {})
            ),
            "colors": self._compare_dicts(
                style1.get('colors', {}),
                style2.get('colors', {})
            ),
            "header": self._compare_dicts(
                style1.get('header', {}),
                style2.get('header', {})
            ),
            "footer": self._compare_dicts(
                style1.get('footer', {}),
                style2.get('footer', {})
            )
        }

        # 計算差異統計
        total_differences = sum(
            len(diff) for diff in differences.values() if isinstance(diff, list)
        )

        result = {
            "doc1": doc1_path,
            "doc2": doc2_path,
            "total_differences": total_differences,
            "differences": differences,
            "comparison_date": datetime.now().isoformat()
        }

        print(f"✓ 發現 {total_differences} 個樣式差異")
        return result

    def _compare_dicts(self, dict1: Dict, dict2: Dict, path: str = "") -> List[str]:
        """遞迴比較兩個字典的差異"""
        differences = []

        all_keys = set(dict1.keys()) | set(dict2.keys())

        for key in all_keys:
            current_path = f"{path}.{key}" if path else key

            if key not in dict1:
                differences.append(f"{current_path}: 僅存在於文件2")
            elif key not in dict2:
                differences.append(f"{current_path}: 僅存在於文件1")
            elif isinstance(dict1[key], dict) and isinstance(dict2[key], dict):
                differences.extend(self._compare_dicts(dict1[key], dict2[key], current_path))
            elif dict1[key] != dict2[key]:
                differences.append(
                    f"{current_path}: 文件1={dict1[key]}, 文件2={dict2[key]}"
                )

        return differences

    def generate_style_report(self, doc_path: str, output_path: str) -> None:
        """
        生成詳細的樣式分析報告

        Args:
            doc_path: 要分析的文件路徑
            output_path: 報告輸出路徑（JSON格式）
        """
        print(f"生成樣式分析報告: {doc_path}")

        # 提取樣式
        styles = self.extract_styles_from_document(doc_path)

        # 載入文件進行驗證
        doc = Document(doc_path)
        validation = self.validate_bestat_compliance(doc)

        # 組合報告
        report = {
            "document": doc_path,
            "analysis_date": datetime.now().isoformat(),
            "extracted_styles": styles,
            "validation_results": validation,
            "recommendations": self._generate_recommendations(validation)
        }

        # 儲存報告
        self.save_style_config(output_path, report)
        print(f"✓ 報告已儲存至: {output_path}")

    def _generate_recommendations(self, validation: Dict[str, Any]) -> List[str]:
        """根據驗證結果生成建議"""
        recommendations = []

        if not validation['compliant']:
            recommendations.append("建議修正所有發現的問題以符合Bestat規範")

            if validation['issues']:
                recommendations.append(f"優先處理 {len(validation['issues'])} 個關鍵問題")

        if validation['warnings']:
            recommendations.append(f"檢視 {len(validation['warnings'])} 個警告項目")

        if validation['compliant'] and not validation['warnings']:
            recommendations.append("文件完全符合Bestat規範，無需修改")

        return recommendations


# 便利函數
def analyze_document_style(document_path: str, output_json_path: str) -> Dict[str, Any]:
    """
    快速分析文件樣式並儲存為JSON

    Args:
        document_path: 要分析的Word文件路徑
        output_json_path: 輸出JSON路徑

    Returns:
        提取的樣式字典
    """
    analyzer = BestatStyleAnalyzer()
    styles = analyzer.extract_styles_from_document(document_path)
    analyzer.save_style_config(output_json_path, styles)
    return styles


def apply_bestat_style_to_document(input_doc_path: str,
                                   output_doc_path: str,
                                   config_path: Optional[str] = None,
                                   document_title: str = "",
                                   protocol_number: str = "",
                                   version: str = "",
                                   logo_path: Optional[str] = None) -> None:
    """
    套用Bestat樣式到文件

    Args:
        input_doc_path: 輸入文件路徑
        output_doc_path: 輸出文件路徑
        config_path: 樣式配置JSON路徑（可選）
        document_title: 文件標題
        protocol_number: Protocol編號
        version: 版本號
        logo_path: Logo路徑
    """
    analyzer = BestatStyleAnalyzer()

    if config_path:
        analyzer.load_style_config(config_path)

    doc = Document(input_doc_path)
    doc = analyzer.apply_bestat_style(doc, document_title, protocol_number, version, logo_path)
    doc.save(output_doc_path)

    print(f"✓ 已套用Bestat樣式並儲存至: {output_doc_path}")


def validate_document_compliance(document_path: str,
                                config_path: Optional[str] = None) -> Dict[str, Any]:
    """
    驗證文件是否符合Bestat規範

    Args:
        document_path: 要驗證的文件路徑
        config_path: 樣式配置JSON路徑（可選）

    Returns:
        驗證結果字典
    """
    analyzer = BestatStyleAnalyzer()

    if config_path:
        analyzer.load_style_config(config_path)

    doc = Document(document_path)
    return analyzer.validate_bestat_compliance(doc)


if __name__ == "__main__":
    print("=" * 70)
    print("Bestat公司文件樣式分析器測試")
    print("=" * 70)

    # 建立測試實例
    analyzer = BestatStyleAnalyzer()

    # 顯示預設配置
    print("\n1. Bestat預設樣式配置:")
    print(f"   - 公司名稱: {analyzer.style_config['company_info']['name']}")
    print(f"   - 頁面大小: {analyzer.style_config['page_setup']['size']}")
    print(f"   - 主要顏色: RGB{tuple(analyzer.style_config['colors']['primary'].values())}")
    print(f"   - 標題字體: {analyzer.style_config['fonts']['title']['name']}")

    # 儲存預設配置
    output_dir = "/home/user/my-colab-notebooks/clinical-doc-automation/output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    config_file = os.path.join(output_dir, "bestat_default_style.json")
    analyzer.save_style_config(config_file)

    print(f"\n✓ 測試完成！預設樣式已儲存至: {config_file}")
    print("\n使用範例:")
    print("  from modules.bestat_style_analyzer import BestatStyleAnalyzer")
    print("  analyzer = BestatStyleAnalyzer()")
    print("  styles = analyzer.extract_styles_from_document('template.docx')")
    print("  analyzer.save_style_config('my_style.json', styles)")
