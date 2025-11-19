"""
CRF (Case Report Form) Generator Module

This module provides functionality to automatically generate CRF documents
based on clinical trial protocol information.

Author: Clinical Documentation Automation Team
Date: 2025-11-18
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Optional, Any
from datetime import datetime
import os


class CRFDomain:
    """Represents a CRF domain with its fields and metadata."""

    def __init__(self, name: str, description: str, fields: List[Dict[str, Any]]):
        """
        Initialize a CRF Domain.

        Args:
            name: Domain name (e.g., "Demographics", "Vital Signs")
            description: Description of the domain
            fields: List of field definitions, each containing:
                - name: Field name
                - label: Field label
                - type: Field type (text, numeric, date, checkbox, dropdown)
                - required: Whether field is required (True/False)
                - options: List of options (for dropdown/checkbox)
                - coding_instruction: Instructions for coding this field
                - unit: Unit of measurement (optional)
        """
        self.name = name
        self.description = description
        self.fields = fields

    def validate(self) -> bool:
        """Validate domain structure."""
        if not self.name or not self.fields:
            return False

        required_keys = ['name', 'label', 'type']
        for field in self.fields:
            if not all(key in field for key in required_keys):
                return False

        return True


class CRFGenerator:
    """Main CRF Generator class."""

    # Standard CRF Domains with predefined fields
    STANDARD_DOMAINS = {
        'demographics': {
            'name': 'Demographics',
            'description': 'Subject demographic information',
            'fields': [
                {
                    'name': 'subject_id',
                    'label': 'Subject ID',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Unique identifier assigned to the subject'
                },
                {
                    'name': 'initials',
                    'label': 'Subject Initials',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'First and last name initials (e.g., JD for John Doe)'
                },
                {
                    'name': 'date_of_birth',
                    'label': 'Date of Birth',
                    'type': 'date',
                    'required': True,
                    'coding_instruction': 'Format: DD-MMM-YYYY (e.g., 15-JAN-1980)'
                },
                {
                    'name': 'age',
                    'label': 'Age',
                    'type': 'numeric',
                    'required': True,
                    'unit': 'years',
                    'coding_instruction': 'Age at time of informed consent'
                },
                {
                    'name': 'sex',
                    'label': 'Sex',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Male', 'Female'],
                    'coding_instruction': 'Biological sex assigned at birth'
                },
                {
                    'name': 'race',
                    'label': 'Race',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['White', 'Black or African American', 'Asian', 'American Indian or Alaska Native', 'Native Hawaiian or Other Pacific Islander', 'Other'],
                    'coding_instruction': 'Select one or more applicable categories'
                },
                {
                    'name': 'ethnicity',
                    'label': 'Ethnicity',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Hispanic or Latino', 'Not Hispanic or Latino', 'Unknown'],
                    'coding_instruction': 'Ethnic background per FDA guidelines'
                }
            ]
        },
        'medical_history': {
            'name': 'Medical History',
            'description': 'Subject medical history',
            'fields': [
                {
                    'name': 'condition',
                    'label': 'Medical Condition',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Describe the medical condition. Use MedDRA preferred terms when possible.'
                },
                {
                    'name': 'start_date',
                    'label': 'Start Date',
                    'type': 'date',
                    'required': False,
                    'coding_instruction': 'Date condition was first diagnosed or observed. Use partial dates if exact date unknown (e.g., MMM-YYYY or YYYY)'
                },
                {
                    'name': 'end_date',
                    'label': 'End Date',
                    'type': 'date',
                    'required': False,
                    'coding_instruction': 'Date condition resolved. Leave blank if ongoing.'
                },
                {
                    'name': 'ongoing',
                    'label': 'Ongoing',
                    'type': 'checkbox',
                    'required': True,
                    'options': ['Yes', 'No'],
                    'coding_instruction': 'Check "Yes" if condition is still present'
                },
                {
                    'name': 'severity',
                    'label': 'Severity',
                    'type': 'dropdown',
                    'required': False,
                    'options': ['Mild', 'Moderate', 'Severe'],
                    'coding_instruction': 'Clinical severity of the condition'
                }
            ]
        },
        'vital_signs': {
            'name': 'Vital Signs',
            'description': 'Vital signs measurements',
            'fields': [
                {
                    'name': 'assessment_date',
                    'label': 'Assessment Date',
                    'type': 'date',
                    'required': True,
                    'coding_instruction': 'Date of vital signs assessment'
                },
                {
                    'name': 'assessment_time',
                    'label': 'Assessment Time',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Time of assessment (24-hour format: HH:MM)'
                },
                {
                    'name': 'temperature',
                    'label': 'Temperature',
                    'type': 'numeric',
                    'required': True,
                    'unit': '°C',
                    'coding_instruction': 'Body temperature in Celsius. Record to one decimal place.'
                },
                {
                    'name': 'systolic_bp',
                    'label': 'Systolic Blood Pressure',
                    'type': 'numeric',
                    'required': True,
                    'unit': 'mmHg',
                    'coding_instruction': 'Systolic BP in mmHg. Subject should be seated and rested.'
                },
                {
                    'name': 'diastolic_bp',
                    'label': 'Diastolic Blood Pressure',
                    'type': 'numeric',
                    'required': True,
                    'unit': 'mmHg',
                    'coding_instruction': 'Diastolic BP in mmHg'
                },
                {
                    'name': 'heart_rate',
                    'label': 'Heart Rate',
                    'type': 'numeric',
                    'required': True,
                    'unit': 'bpm',
                    'coding_instruction': 'Heart rate in beats per minute'
                },
                {
                    'name': 'respiratory_rate',
                    'label': 'Respiratory Rate',
                    'type': 'numeric',
                    'required': True,
                    'unit': 'breaths/min',
                    'coding_instruction': 'Respiratory rate per minute'
                },
                {
                    'name': 'weight',
                    'label': 'Weight',
                    'type': 'numeric',
                    'required': True,
                    'unit': 'kg',
                    'coding_instruction': 'Body weight in kilograms. Record to one decimal place.'
                },
                {
                    'name': 'height',
                    'label': 'Height',
                    'type': 'numeric',
                    'required': False,
                    'unit': 'cm',
                    'coding_instruction': 'Height in centimeters. Required at baseline only.'
                }
            ]
        },
        'laboratory_tests': {
            'name': 'Laboratory Tests',
            'description': 'Laboratory test results',
            'fields': [
                {
                    'name': 'collection_date',
                    'label': 'Collection Date',
                    'type': 'date',
                    'required': True,
                    'coding_instruction': 'Date sample was collected'
                },
                {
                    'name': 'collection_time',
                    'label': 'Collection Time',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Time sample was collected (24-hour format)'
                },
                {
                    'name': 'test_name',
                    'label': 'Test Name',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Name of laboratory test (use standard LOINC codes when available)'
                },
                {
                    'name': 'result',
                    'label': 'Result',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Test result value with appropriate precision'
                },
                {
                    'name': 'unit',
                    'label': 'Unit',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Unit of measurement (e.g., mg/dL, mmol/L)'
                },
                {
                    'name': 'normal_range',
                    'label': 'Normal Range',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Laboratory normal reference range'
                },
                {
                    'name': 'clinically_significant',
                    'label': 'Clinically Significant',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Yes', 'No', 'Not Applicable'],
                    'coding_instruction': 'Investigator assessment of clinical significance'
                }
            ]
        },
        'adverse_events': {
            'name': 'Adverse Events',
            'description': 'Adverse event reporting',
            'fields': [
                {
                    'name': 'ae_term',
                    'label': 'Adverse Event Term',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Describe AE using MedDRA preferred term. Be specific and use medical terminology.'
                },
                {
                    'name': 'start_date',
                    'label': 'Start Date',
                    'type': 'date',
                    'required': True,
                    'coding_instruction': 'Date AE first occurred or was observed'
                },
                {
                    'name': 'end_date',
                    'label': 'End Date',
                    'type': 'date',
                    'required': False,
                    'coding_instruction': 'Date AE resolved. Leave blank if ongoing.'
                },
                {
                    'name': 'ongoing',
                    'label': 'Ongoing',
                    'type': 'checkbox',
                    'required': True,
                    'options': ['Yes', 'No'],
                    'coding_instruction': 'Check "Yes" if AE has not resolved'
                },
                {
                    'name': 'severity',
                    'label': 'Severity',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Mild', 'Moderate', 'Severe'],
                    'coding_instruction': 'Mild: Awareness of event but easily tolerated; Moderate: Discomfort enough to cause interference; Severe: Incapacitating'
                },
                {
                    'name': 'serious',
                    'label': 'Serious',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Yes', 'No'],
                    'coding_instruction': 'SAE criteria: death, life-threatening, hospitalization, disability, congenital anomaly, or medically important'
                },
                {
                    'name': 'relationship',
                    'label': 'Relationship to Study Drug',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Not Related', 'Unlikely', 'Possible', 'Probable', 'Definite'],
                    'coding_instruction': 'Investigator assessment of relationship to study intervention'
                },
                {
                    'name': 'action_taken',
                    'label': 'Action Taken',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['None', 'Dose Reduced', 'Dose Interrupted', 'Drug Withdrawn', 'Other'],
                    'coding_instruction': 'Action taken with study drug due to AE'
                },
                {
                    'name': 'outcome',
                    'label': 'Outcome',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Recovered/Resolved', 'Recovering/Resolving', 'Not Recovered/Not Resolved', 'Recovered with Sequelae', 'Fatal', 'Unknown'],
                    'coding_instruction': 'Final outcome of the adverse event'
                }
            ]
        },
        'concomitant_medications': {
            'name': 'Concomitant Medications',
            'description': 'Medications taken during the study',
            'fields': [
                {
                    'name': 'medication_name',
                    'label': 'Medication Name',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Generic name preferred. Include strength if known.'
                },
                {
                    'name': 'indication',
                    'label': 'Indication',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Reason for taking medication'
                },
                {
                    'name': 'dose',
                    'label': 'Dose',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Dose amount and unit (e.g., 10 mg, 1 tablet)'
                },
                {
                    'name': 'route',
                    'label': 'Route',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Oral', 'Intravenous', 'Intramuscular', 'Subcutaneous', 'Topical', 'Inhalation', 'Other'],
                    'coding_instruction': 'Route of administration'
                },
                {
                    'name': 'frequency',
                    'label': 'Frequency',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Dosing frequency (e.g., once daily, BID, PRN)'
                },
                {
                    'name': 'start_date',
                    'label': 'Start Date',
                    'type': 'date',
                    'required': True,
                    'coding_instruction': 'Date medication started'
                },
                {
                    'name': 'end_date',
                    'label': 'End Date',
                    'type': 'date',
                    'required': False,
                    'coding_instruction': 'Date medication stopped. Leave blank if ongoing.'
                },
                {
                    'name': 'ongoing',
                    'label': 'Ongoing',
                    'type': 'checkbox',
                    'required': True,
                    'options': ['Yes', 'No'],
                    'coding_instruction': 'Check "Yes" if still taking medication'
                }
            ]
        },
        'study_drug_administration': {
            'name': 'Study Drug Administration',
            'description': 'Study drug dosing information',
            'fields': [
                {
                    'name': 'administration_date',
                    'label': 'Administration Date',
                    'type': 'date',
                    'required': True,
                    'coding_instruction': 'Date study drug was administered'
                },
                {
                    'name': 'administration_time',
                    'label': 'Administration Time',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Time study drug was administered (24-hour format)'
                },
                {
                    'name': 'dose_amount',
                    'label': 'Dose Amount',
                    'type': 'numeric',
                    'required': True,
                    'coding_instruction': 'Amount of study drug administered'
                },
                {
                    'name': 'dose_unit',
                    'label': 'Dose Unit',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Unit of dose (e.g., mg, mL, tablets)'
                },
                {
                    'name': 'route',
                    'label': 'Route',
                    'type': 'dropdown',
                    'required': True,
                    'options': ['Oral', 'Intravenous', 'Intramuscular', 'Subcutaneous', 'Other'],
                    'coding_instruction': 'Route of administration per protocol'
                },
                {
                    'name': 'lot_number',
                    'label': 'Lot Number',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Study drug lot/batch number'
                },
                {
                    'name': 'administered_by',
                    'label': 'Administered By',
                    'type': 'text',
                    'required': True,
                    'coding_instruction': 'Name or initials of person who administered study drug'
                },
                {
                    'name': 'comments',
                    'label': 'Comments',
                    'type': 'text',
                    'required': False,
                    'coding_instruction': 'Any relevant comments or deviations'
                }
            ]
        }
    }

    def __init__(self, protocol_info: Optional[Dict[str, Any]] = None):
        """
        Initialize CRF Generator.

        Args:
            protocol_info: Dictionary containing protocol information
                - study_title: Study title
                - protocol_number: Protocol number
                - sponsor: Sponsor name
                - version: CRF version
                - custom_domains: List of custom domain definitions
        """
        self.protocol_info = protocol_info or {}
        self.document = None
        self.custom_domains = {}

        # Load custom domains if provided
        if 'custom_domains' in self.protocol_info:
            for domain_def in self.protocol_info['custom_domains']:
                domain = CRFDomain(**domain_def)
                if domain.validate():
                    self.custom_domains[domain.name.lower().replace(' ', '_')] = {
                        'name': domain.name,
                        'description': domain.description,
                        'fields': domain.fields
                    }

    def _set_cell_background(self, cell, color_hex: str):
        """Set cell background color."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_cell_border(self, cell):
        """Set cell borders."""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # Create border elements
        tcBorders = OxmlElement('w:tcBorders')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)

        tcPr.append(tcBorders)

    def _add_header(self):
        """Add CRF document header."""
        # Title
        title = self.document.add_heading('Case Report Form (CRF)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Protocol information table
        info_table = self.document.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ('Study Title:', self.protocol_info.get('study_title', 'N/A')),
            ('Protocol Number:', self.protocol_info.get('protocol_number', 'N/A')),
            ('Sponsor:', self.protocol_info.get('sponsor', 'N/A')),
            ('CRF Version:', self.protocol_info.get('version', '1.0'))
        ]

        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

            # Format label cell
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(row.cells[0], 'D3D3D3')

        self.document.add_paragraph()

    def _add_domain_section(self, domain_key: str, domain_data: Dict[str, Any]):
        """
        Add a CRF domain section to the document.

        Args:
            domain_key: Domain identifier key
            domain_data: Domain configuration data
        """
        # Domain heading
        heading = self.document.add_heading(domain_data['name'], 1)
        heading.style.font.color.rgb = RGBColor(0, 51, 102)

        # Domain description
        desc_para = self.document.add_paragraph(domain_data['description'])
        desc_para.style = 'Intense Quote'

        # Create domain table
        fields = domain_data['fields']
        table = self.document.add_table(rows=len(fields) + 1, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Field Name', 'Type', 'Required', 'Value/Response']

        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_background(cell, '4472C4')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Field rows
        for i, field in enumerate(fields, start=1):
            row = table.rows[i]

            # Field label
            row.cells[0].text = field['label']
            if field.get('unit'):
                row.cells[0].text += f" ({field['unit']})"

            # Field type
            field_type = field['type'].capitalize()
            if field['type'] == 'dropdown' and 'options' in field:
                field_type += f" ({len(field['options'])} options)"
            row.cells[1].text = field_type

            # Required
            row.cells[2].text = 'Yes' if field.get('required', False) else 'No'
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Value/Response cell - leave blank for data entry
            if field['type'] == 'dropdown' and 'options' in field:
                options_text = ' □ ' + '  □ '.join(field['options'])
                row.cells[3].text = options_text
            elif field['type'] == 'checkbox' and 'options' in field:
                options_text = ' □ ' + '  □ '.join(field['options'])
                row.cells[3].text = options_text
            else:
                row.cells[3].text = ''

            # Apply borders
            for cell in row.cells:
                self._set_cell_border(cell)

        # Add coding instructions section
        self.document.add_paragraph()
        inst_heading = self.document.add_heading(f'{domain_data["name"]} - Coding Instructions', 2)
        inst_heading.style.font.color.rgb = RGBColor(102, 102, 102)

        for field in fields:
            if 'coding_instruction' in field:
                p = self.document.add_paragraph(style='List Bullet')
                p.add_run(f"{field['label']}: ").bold = True
                p.add_run(field['coding_instruction'])

        # Add page break after each domain
        self.document.add_page_break()

    def generate_crf(
        self,
        domains: Optional[List[str]] = None,
        output_path: Optional[str] = None,
        include_all_standard: bool = False
    ) -> str:
        """
        Generate CRF document.

        Args:
            domains: List of domain keys to include. If None, includes all standard domains.
            output_path: Path to save the generated CRF. If None, auto-generates path.
            include_all_standard: If True, includes all standard domains regardless of domains parameter.

        Returns:
            Path to generated CRF document
        """
        # Initialize document
        self.document = Document()

        # Set document properties
        self.document.core_properties.title = 'Case Report Form'
        self.document.core_properties.author = self.protocol_info.get('sponsor', 'Clinical Trial')
        self.document.core_properties.created = datetime.now()

        # Add header
        self._add_header()

        # Determine which domains to include
        if include_all_standard:
            selected_domains = list(self.STANDARD_DOMAINS.keys())
        elif domains:
            selected_domains = domains
        else:
            selected_domains = list(self.STANDARD_DOMAINS.keys())

        # Add standard domains
        for domain_key in selected_domains:
            if domain_key in self.STANDARD_DOMAINS:
                self._add_domain_section(domain_key, self.STANDARD_DOMAINS[domain_key])
            elif domain_key in self.custom_domains:
                self._add_domain_section(domain_key, self.custom_domains[domain_key])
            else:
                print(f"Warning: Domain '{domain_key}' not found. Skipping.")

        # Add custom domains that weren't explicitly requested
        for domain_key, domain_data in self.custom_domains.items():
            if domain_key not in selected_domains:
                self._add_domain_section(domain_key, domain_data)

        # Generate output path if not provided
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            protocol_num = self.protocol_info.get('protocol_number', 'UNKNOWN').replace('/', '_')
            output_path = f'CRF_{protocol_num}_{timestamp}.docx'

        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Save document
        self.document.save(output_path)
        print(f"CRF generated successfully: {output_path}")

        return output_path

    def add_custom_domain(self, domain: CRFDomain) -> bool:
        """
        Add a custom domain to the generator.

        Args:
            domain: CRFDomain object

        Returns:
            True if successfully added, False otherwise
        """
        if not domain.validate():
            print(f"Error: Invalid domain structure for '{domain.name}'")
            return False

        domain_key = domain.name.lower().replace(' ', '_')
        self.custom_domains[domain_key] = {
            'name': domain.name,
            'description': domain.description,
            'fields': domain.fields
        }

        print(f"Custom domain '{domain.name}' added successfully")
        return True

    def get_available_domains(self) -> List[str]:
        """
        Get list of available domain keys.

        Returns:
            List of domain keys (both standard and custom)
        """
        standard = list(self.STANDARD_DOMAINS.keys())
        custom = list(self.custom_domains.keys())
        return standard + custom

    def export_domain_template(self, domain_key: str, output_path: str) -> bool:
        """
        Export a single domain as a template document.

        Args:
            domain_key: Domain identifier
            output_path: Path to save template

        Returns:
            True if successful, False otherwise
        """
        # Find domain
        if domain_key in self.STANDARD_DOMAINS:
            domain_data = self.STANDARD_DOMAINS[domain_key]
        elif domain_key in self.custom_domains:
            domain_data = self.custom_domains[domain_key]
        else:
            print(f"Error: Domain '{domain_key}' not found")
            return False

        # Create document
        self.document = Document()
        self.document.core_properties.title = f'{domain_data["name"]} Template'

        # Add domain section
        self._add_domain_section(domain_key, domain_data)

        # Save
        self.document.save(output_path)
        print(f"Domain template exported: {output_path}")

        return True


# Example usage and utility functions
def create_example_crf():
    """Create an example CRF with standard domains."""

    # Protocol information
    protocol_info = {
        'study_title': 'A Phase III Randomized, Double-Blind Study of Novel Drug X in Patients with Advanced Disease',
        'protocol_number': 'PROTO-2025-001',
        'sponsor': 'Example Pharmaceutical Company',
        'version': '1.0'
    }

    # Initialize generator
    generator = CRFGenerator(protocol_info)

    # Generate CRF with selected domains
    output_path = generator.generate_crf(
        domains=['demographics', 'vital_signs', 'adverse_events', 'study_drug_administration'],
        output_path='example_CRF.docx'
    )

    return output_path


def create_custom_domain_example():
    """Create an example CRF with custom domains."""

    # Protocol information
    protocol_info = {
        'study_title': 'Custom Study with Specialized Assessments',
        'protocol_number': 'CUSTOM-2025-001',
        'sponsor': 'Research Institute',
        'version': '1.0'
    }

    # Initialize generator
    generator = CRFGenerator(protocol_info)

    # Define custom domain
    custom_domain = CRFDomain(
        name='Quality of Life Assessment',
        description='Patient-reported quality of life measures',
        fields=[
            {
                'name': 'assessment_date',
                'label': 'Assessment Date',
                'type': 'date',
                'required': True,
                'coding_instruction': 'Date QoL questionnaire was completed'
            },
            {
                'name': 'physical_functioning',
                'label': 'Physical Functioning Score',
                'type': 'numeric',
                'required': True,
                'coding_instruction': 'Score range 0-100, higher score indicates better functioning'
            },
            {
                'name': 'emotional_wellbeing',
                'label': 'Emotional Well-being Score',
                'type': 'numeric',
                'required': True,
                'coding_instruction': 'Score range 0-100, higher score indicates better well-being'
            },
            {
                'name': 'pain_level',
                'label': 'Pain Level',
                'type': 'dropdown',
                'required': True,
                'options': ['None', 'Mild', 'Moderate', 'Severe', 'Very Severe'],
                'coding_instruction': 'Subject-reported pain level in past 24 hours'
            }
        ]
    )

    # Add custom domain
    generator.add_custom_domain(custom_domain)

    # Generate CRF
    output_path = generator.generate_crf(
        domains=['demographics', 'quality_of_life_assessment'],
        output_path='custom_domain_CRF.docx'
    )

    return output_path


if __name__ == '__main__':
    # Example 1: Create standard CRF
    print("=" * 80)
    print("Example 1: Creating CRF with standard domains")
    print("=" * 80)
    crf_path = create_example_crf()
    print(f"\nGenerated CRF: {crf_path}\n")

    # Example 2: Create CRF with custom domain
    print("=" * 80)
    print("Example 2: Creating CRF with custom domain")
    print("=" * 80)
    custom_crf_path = create_custom_domain_example()
    print(f"\nGenerated custom CRF: {custom_crf_path}\n")

    # Example 3: List available domains
    print("=" * 80)
    print("Example 3: Listing available domains")
    print("=" * 80)
    generator = CRFGenerator()
    domains = generator.get_available_domains()
    print(f"\nAvailable domains: {', '.join(domains)}\n")
